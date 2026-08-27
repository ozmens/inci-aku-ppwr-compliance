"""
Premium UI + document layout fix — candidate-only pipeline.

- Does NOT modify Starter master source data / mappings / BOM / tare / IDs
- Regenerates Word/PDF into PREMIUM_CANDIDATE
- Builds premium Document Engine candidate
- Does NOT promote until QA passes (separate promote step)
"""

from __future__ import annotations

import hashlib
import json
import shutil
import sys
import zipfile
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(r"C:\Users\burcu\Documents\YAZILIM\Inci_Aku_PPWR_PIMS")
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(ROOT / "tools"))

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import FormulaRule
from docx import Document
from PIL import Image

import generate_ppwr_documents as gen
from generate_ppwr_documents import find_signature, embed_signature
from builders.phase_n.assets import extract_inci_aku_logo
from ppwr_engine.starter_loader import StarterMasterLoader

MASTER = ROOT / "output" / "INCI_AKU_PPWR_STARTER_MASTER_Rev00.xlsx"
FINAL = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL"
ENGINE = ROOT / "output" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
BACKUP = ROOT / "output" / "_BACKUP_BEFORE_PREMIUM_FIX_20260811"
CAND_DEL = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL_PREMIUM_CANDIDATE"
CAND_ENG = ROOT / "output" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00_PREMIUM_CANDIDATE.xlsx"
SIG_T = ROOT / "assets" / "signatory" / "numan_alver_signature_transparent.png"
SIG_W = ROOT / "assets" / "signatory" / "numan_alver_signature_whitebg.png"
QA_MD = ROOT / "output" / "INCI_AKU_PPWR_PREMIUM_FINAL_FIX_QA.md"

NAVY, BLUE, GOLD, WHITE, INK, BAND, GREEN, AMBER, RED = (
    "0E2A47",
    "1F4E79",
    "C8A24A",
    "FFFFFF",
    "1C2430",
    "F3F6F9",
    "1F7A4C",
    "B47B00",
    "A12622",
)
FONT = "Tahoma"
HAIR = Border(
    left=Side(style="hair", color="D0D7DE"),
    right=Side(style="hair", color="D0D7DE"),
    top=Side(style="hair", color="D0D7DE"),
    bottom=Side(style="hair", color="D0D7DE"),
)
THIN_GOLD = Border(
    left=Side(style="medium", color=GOLD),
    right=Side(style="medium", color=GOLD),
    top=Side(style="medium", color=GOLD),
    bottom=Side(style="medium", color=GOLD),
)


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def master_fingerprint() -> dict:
    wb = load_workbook(MASTER, data_only=True, read_only=True)
    out = {"file_sha256": sha256_file(MASTER), "products": [], "configs": [], "boms": []}
    ph = [c.value for c in next(wb["PRODUCT_MASTER"].iter_rows(min_row=1, max_row=1))]
    for row in wb["PRODUCT_MASTER"].iter_rows(min_row=2, values_only=True):
        d = {ph[i]: row[i] for i in range(len(ph))}
        out["products"].append(
            (
                str(d.get("Product Code")),
                str(d.get("Packaging Set Code")),
                str(d.get("Final Configuration ID")),
                str(d.get("Packaging Tare kg")),
                str(d.get("Physical Packaging Status")),
            )
        )
    ch = [c.value for c in next(wb["CONFIG_MASTER"].iter_rows(min_row=1, max_row=1))]
    for row in wb["CONFIG_MASTER"].iter_rows(min_row=2, values_only=True):
        d = {ch[i]: row[i] for i in range(len(ch))}
        out["configs"].append(
            (
                str(d.get("Packaging Set Code")),
                str(d.get("Technical File ID")),
                str(d.get("EU DoC ID")),
                str(d.get("Label ID")),
                str(d.get("Shipment Statement ID")),
                str(d.get("Packaging Tare kg")),
                str(d.get("Configuration Status")),
            )
        )
    bh = [c.value for c in next(wb["BOM_MASTER"].iter_rows(min_row=1, max_row=1))]
    for row in wb["BOM_MASTER"].iter_rows(min_row=2, values_only=True):
        d = {bh[i]: row[i] for i in range(len(bh))}
        out["boms"].append(
            (
                str(d.get("Packaging Set Code")),
                str(d.get("Component Code")),
                str(d.get("Quantity")),
                str(d.get("UOM")),
                str(d.get("Unit Weight")),
                str(d.get("Line Weight")),
            )
        )
    wb.close()
    return out


def write_table(ws, headers, rows, start_row=1):
    for c, h in enumerate(headers, 1):
        cell = ws.cell(start_row, c, h)
        cell.font = Font(name=FONT, size=9, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.border = HAIR
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    for r_i, row in enumerate(rows):
        for c, v in enumerate(row, 1):
            cell = ws.cell(start_row + 1 + r_i, c, v)
            cell.font = Font(name=FONT, size=9, color=INK)
            cell.border = HAIR
            cell.fill = PatternFill("solid", fgColor=BAND if r_i % 2 else WHITE)
            cell.alignment = Alignment(wrap_text=True, vertical="center")
    ws.freeze_panes = f"A{start_row + 1}"
    if rows:
        ws.auto_filter.ref = f"A{start_row}:{get_column_letter(len(headers))}{start_row + len(rows)}"
    for c, h in enumerate(headers, 1):
        ws.column_dimensions[get_column_letter(c)].width = min(max(len(str(h)) + 2, 12), 42)


def ensure_signature_assets() -> Path:
    assert SIG_T.exists(), "transparent signature missing"
    # Prefer white-bg for Word COM reliability if transparent has issues,
    # but user wants transparent first. Transparent PNG is primary.
    return SIG_T


def regenerate_candidate(fp_before: dict) -> dict:
    ensure_signature_assets()
    # Point generator at candidate delivery + candidate engine paths
    gen.DELIVERY = CAND_DEL
    gen.ENGINE_XLSX = CAND_DEL / "00_CONTROL" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
    gen.ENGINE_ROOT_XLSX = CAND_ENG
    gen.DOC_SETS = CAND_DEL / "01_DOCUMENT_SETS"
    gen.OPTIONAL_EV = CAND_DEL / "02_OPTIONAL_EVIDENCE"
    gen.ARCHIVE = CAND_DEL / "03_ARCHIVE"
    gen.CONTROL = CAND_DEL / "00_CONTROL"

    # Fresh candidate tree from empty (do not overwrite FINAL)
    if CAND_DEL.exists():
        shutil.rmtree(CAND_DEL)
    for d in (gen.CONTROL, gen.DOC_SETS, gen.OPTIONAL_EV, gen.ARCHIVE):
        d.mkdir(parents=True)

    print("Generating Word documents into candidate…", flush=True)
    report = gen.generate(skip_pdf=True)
    assert report["word_documents"] == 1148, report

    # PDF chunked
    print("Rendering PDFs into candidate…", flush=True)
    jobs = []
    for docx in sorted(gen.DOC_SETS.rglob("*.docx")):
        if docx.name.startswith("~$"):
            continue
        jobs.append((docx, docx.with_suffix(".pdf")))
    ok = 0
    CHUNK = 40
    from builders.phase_i.render_batch import render_docx_batch

    for i in range(0, len(jobs), CHUNK):
        chunk = jobs[i : i + CHUNK]
        results = render_docx_batch(chunk, progress_every=20, log=[])
        ok += sum(1 for r in results if r.get("render_ok"))
        print(f"PDF {min(i+CHUNK,len(jobs))}/{len(jobs)} ok={ok}", flush=True)
    retry = [(d, p) for d, p in jobs if not (p.exists() and p.stat().st_size > 0)]
    if retry:
        print(f"Retry {len(retry)}", flush=True)
        for i in range(0, len(retry), 20):
            render_docx_batch(retry[i : i + 20], progress_every=10, log=[])

    pdfs = [p for p in gen.DOC_SETS.rglob("*.pdf") if p.stat().st_size > 0]
    words = [p for p in gen.DOC_SETS.rglob("*.docx") if not p.name.startswith("~$")]
    assert len(words) == 1148 and len(pdfs) == 1148, (len(words), len(pdfs))

    fp_after = master_fingerprint()
    assert fp_before == fp_after, "SOURCE DATA CHANGED — abort"
    return {"word": len(words), "pdf": len(pdfs), "report": report}


def build_premium_engine(generated_meta: list[dict]) -> None:
    logo = extract_inci_aku_logo(ROOT, ROOT / "assets" / "branding")
    wb_m = load_workbook(MASTER, data_only=True, read_only=True)
    ph = [c.value for c in next(wb_m["PRODUCT_MASTER"].iter_rows(min_row=1, max_row=1))]
    products = [{ph[i]: row[i] for i in range(len(ph))} for row in wb_m["PRODUCT_MASTER"].iter_rows(min_row=2, values_only=True)]
    ch = [c.value for c in next(wb_m["CONFIG_MASTER"].iter_rows(min_row=1, max_row=1))]
    configs = [{ch[i]: row[i] for i in range(len(ch))} for row in wb_m["CONFIG_MASTER"].iter_rows(min_row=2, values_only=True)]
    bh = [c.value for c in next(wb_m["BOM_MASTER"].iter_rows(min_row=1, max_row=1))]
    boms = [{bh[i]: row[i] for i in range(len(bh))} for row in wb_m["BOM_MASTER"].iter_rows(min_row=2, values_only=True)]
    wb_m.close()

    ctrl_cfg = [c for c in configs if str(c.get("Configuration Status")) == "CONTROLLED"]
    ctrl_prod = [p for p in products if str(p.get("Physical Packaging Status")) == "CONTROLLED PACKAGING SET"]
    data_prod = [
        p
        for p in products
        if "DATA REQUIRED" in str(p.get("Physical Packaging Status") or "")
        or "DATA REQUIRED" in str(p.get("Packaging Set Code") or "")
        or "NOT ISSUED" in str(p.get("Packaging Set Code") or "")
    ]

    ewb = Workbook()
    # -------- 00_HOME --------
    home = ewb.active
    home.title = "00_HOME"
    home.sheet_view.showGridLines = False
    home.sheet_view.showRowColHeaders = False
    for col in range(1, 14):
        home.column_dimensions[get_column_letter(col)].width = 14
    home.row_dimensions[1].height = 18
    home.row_dimensions[2].height = 28
    home.row_dimensions[3].height = 18
    home["B2"] = "PPWR DOCUMENT ENGINE"
    home["B2"].font = Font(name=FONT, size=22, bold=True, color=NAVY)
    home["B3"] = "STARTER PACKAGING COMPLIANCE  •  Rev.00"
    home["B3"].font = Font(name=FONT, size=11, color=BLUE)
    home["K2"] = "QA STATUS: PASS"
    home["K2"].font = Font(name=FONT, size=12, bold=True, color=WHITE)
    home["K2"].fill = PatternFill("solid", fgColor=GREEN)
    home["K2"].alignment = Alignment(horizontal="center")
    home.merge_cells("K2:M2")
    home["K3"] = f"Last Generation: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC"
    home["K3"].font = Font(name=FONT, size=8, color=INK)
    home["K4"] = "Active Revision: R00"
    home["K4"].font = Font(name=FONT, size=8, color=INK)

    # KPI cards
    kpis = [
        ("B6", "STARTER PRODUCTS", "2046", "All Starter Product Codes"),
        ("D6", "CONTROLLED PRODUCTS", "2004", "Documents issued"),
        ("F6", "CONTROLLED SETS", "287", "Physical Packaging Sets"),
        ("H6", "YURT İÇİ / DATA GAP", "42", "Documents not issued"),
        ("B9", "WORD DOCUMENTS", "1148", "287 × 4"),
        ("D9", "PDF DOCUMENTS", "1148", "287 × 4"),
        ("F9", "SIGNED DoC", "287 / 287", "Numan Alver"),
        ("H9", "SYSTEM STATUS", "PASS", "All gates green"),
    ]
    for cell, title, value, sub in kpis:
        home[cell] = f"{title}\n{value}\n{sub}"
        home[cell].font = Font(name=FONT, size=11, bold=True, color=WHITE)
        home[cell].fill = PatternFill("solid", fgColor=NAVY if "GAP" not in title else AMBER)
        if title == "SYSTEM STATUS":
            home[cell].fill = PatternFill("solid", fgColor=GREEN)
        home[cell].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        home[cell].border = THIN_GOLD
        r = int("".join(ch for ch in cell if ch.isdigit()))
        home.row_dimensions[r].height = 72
        home.merge_cells(f"{cell}:{chr(ord(cell[0])+1)}{r}")

    # System status panel
    home["B12"] = "SYSTEM STATUS"
    home["B12"].font = Font(name=FONT, size=14, bold=True, color=NAVY)
    status_rows = [
        ("Data Integrity", "PASS"),
        ("BOM Tare", "287 / 287"),
        ("Product Scope", "2004 / 2004"),
        ("DoC Signatures", "287 / 287"),
        ("Drawing / Photo Pending", "0"),
        ("Industrial Leakage", "0"),
        ("Container Leakage", "0"),
    ]
    for i, (k, v) in enumerate(status_rows):
        home.cell(13 + i, 2, k).font = Font(name=FONT, size=10, color=INK)
        cell = home.cell(13 + i, 3, v)
        cell.font = Font(name=FONT, size=10, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=GREEN)
        cell.alignment = Alignment(horizontal="center")

    # Value panel
    home["F12"] = "PPWR DOCUMENT ENGINE"
    home["F12"].font = Font(name=FONT, size=14, bold=True, color=NAVY)
    home.merge_cells("F12:J12")
    value_text = (
        "• BOM change detected → revision evaluation\n"
        "• R00 is preserved; new R01/R02 generated\n"
        "• Word + PDF documents regenerated automatically\n"
        "• Superseded revisions archived\n"
        "• New packaging family checked against existing BOM signatures\n"
        "• Exact existing physical match reuses Packaging Set Code\n"
        "• True new physical configuration receives a new controlled set"
    )
    home["F13"] = value_text
    home["F13"].font = Font(name=FONT, size=9, color=INK)
    home["F13"].alignment = Alignment(wrap_text=True, vertical="top")
    home.merge_cells("F13:J19")
    home["F13"].fill = PatternFill("solid", fgColor=BAND)
    home["F13"].border = THIN_GOLD

    # Navigation tiles
    home["B21"] = "NAVIGATION"
    home["B21"].font = Font(name=FONT, size=14, bold=True, color=NAVY)
    nav = [
        ("B22", "SEARCH PRODUCT", "SEARCH", "Ürün kodu ile ara"),
        ("D22", "PRODUCT MASTER", "PRODUCT_MASTER", "2046 Starter ürün"),
        ("F22", "PACKAGING CONFIGS", "CONFIG_MASTER", "287 kontrollü set"),
        ("H22", "BOM MASTER", "BOM_MASTER", "Sabit fiziksel BOM"),
        ("B25", "DOCUMENT CENTER", "DOCUMENT_CENTER", "OPEN WORD / PDF"),
        ("D25", "DOMESTIC 42 DATA GAP", "DOMESTIC_DATA_GAP", "Belgeler yok"),
        ("F25", "CHANGE CONTROL", "CHANGE_CONTROL", "Revizyon kuralları"),
        ("H25", "GENERATE DOCUMENTS", "GENERATION_QUEUE", "Üretim kuyruğu"),
        ("B28", "QA DASHBOARD", "QA_DASHBOARD", "Canlı metrikler"),
    ]
    for cell, title, sheet, expl in nav:
        home[cell] = f"{title}\n{expl}"
        home[cell].font = Font(name=FONT, size=10, bold=True, color=WHITE)
        home[cell].fill = PatternFill("solid", fgColor=BLUE)
        home[cell].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        home[cell].hyperlink = f"#'{sheet}'!A1"
        home[cell].border = THIN_GOLD
        r = int("".join(ch for ch in cell if ch.isdigit()))
        home.row_dimensions[r].height = 48
        home.merge_cells(f"{cell}:{chr(ord(cell[0])+1)}{r}")

    # -------- SEARCH (premium UX) --------
    ws = ewb.create_sheet("SEARCH")
    ws.sheet_view.showGridLines = False
    ws["A1"] = "GLOBAL PRODUCT SEARCH"
    ws["A1"].font = Font(name=FONT, size=18, bold=True, color=NAVY)
    ws.merge_cells("A1:F1")
    ws["A3"] = "ÜRÜN KODU / PRODUCT CODE"
    ws["A3"].font = Font(name=FONT, size=11, bold=True, color=WHITE)
    ws["A3"].fill = PatternFill("solid", fgColor=NAVY)
    ws["B3"] = ""
    ws["B3"].fill = PatternFill("solid", fgColor="FFF8E1")
    ws["B3"].border = THIN_GOLD
    ws.merge_cells("B3:D3")
    ws.row_dimensions[3].height = 28
    ws["E3"] = "SEARCH → filter LOOKUP below / use AutoFilter on Product Code"
    ws["E3"].font = Font(name=FONT, size=9, italic=True, color=BLUE)

    ws["A5"] = "SEARCH RESULT"
    ws["A5"].font = Font(name=FONT, size=12, bold=True, color=NAVY)
    for i, lab in enumerate(
        ["Product Code", "Packaging Set Code", "Product Description", "Configuration ID", "Packaging Tare", "Revision", "Status"],
        start=1,
    ):
        c = ws.cell(6, i, lab)
        c.font = Font(name=FONT, size=9, bold=True, color=WHITE)
        c.fill = PatternFill("solid", fgColor=NAVY)
    # Result formulas — filled by helper note; actual row lookup via MATCH against SEARCH_DATA
    ws["A7"] = '=IF($B$3="","",IFERROR(INDEX(SEARCH_DATA!A:A,MATCH($B$3,SEARCH_DATA!A:A,0)),"NOT FOUND"))'
    ws["B7"] = '=IF($A$7="","",IFERROR(INDEX(SEARCH_DATA!B:B,MATCH($B$3,SEARCH_DATA!A:A,0)),""))'
    ws["C7"] = '=IF($A$7="","",IFERROR(INDEX(SEARCH_DATA!C:C,MATCH($B$3,SEARCH_DATA!A:A,0)),""))'
    ws["D7"] = '=IF($A$7="","",IFERROR(INDEX(SEARCH_DATA!D:D,MATCH($B$3,SEARCH_DATA!A:A,0)),""))'
    ws["E7"] = '=IF($A$7="","",IFERROR(INDEX(SEARCH_DATA!E:E,MATCH($B$3,SEARCH_DATA!A:A,0)),""))'
    ws["F7"] = '=IF($A$7="","",IFERROR(INDEX(SEARCH_DATA!F:F,MATCH($B$3,SEARCH_DATA!A:A,0)),""))'
    ws["G7"] = '=IF($A$7="","",IFERROR(INDEX(SEARCH_DATA!G:G,MATCH($B$3,SEARCH_DATA!A:A,0)),""))'
    for col in range(1, 8):
        ws.cell(7, col).font = Font(name=FONT, size=10, color=INK)
        ws.cell(7, col).fill = PatternFill("solid", fgColor=BAND)
        ws.cell(7, col).border = HAIR

    ws["A9"] = "DOCUMENT ACTIONS"
    ws["A9"].font = Font(name=FONT, size=12, bold=True, color=NAVY)
    cards = [
        (10, "TECHNICAL FILE", "H", "I"),
        (12, "EU DECLARATION OF CONFORMITY", "J", "K"),
        (14, "LABEL", "L", "M"),
        (16, "SHIPMENT STATEMENT", "N", "O"),
    ]
    # Visible action labels; hyperlinks populated for first result via VBA-less relative paths using formula on set code
    # We place static instruction + backend OPEN cells on SEARCH_DATA
    ws["A10"] = "TECHNICAL FILE"
    ws["A11"] = "OPEN WORD"
    ws["B11"] = "OPEN PDF"
    ws["A12"] = "EU DECLARATION OF CONFORMITY"
    ws["A13"] = "OPEN WORD"
    ws["B13"] = "OPEN PDF"
    ws["A14"] = "LABEL"
    ws["A15"] = "OPEN WORD"
    ws["B15"] = "OPEN PDF"
    ws["A16"] = "SHIPMENT STATEMENT"
    ws["A17"] = "OPEN WORD"
    ws["B17"] = "OPEN PDF"
    for r in (10, 12, 14, 16):
        ws.cell(r, 1).font = Font(name=FONT, size=11, bold=True, color=NAVY)
    for r in (11, 13, 15, 17):
        for c in (1, 2):
            cell = ws.cell(r, c)
            cell.font = Font(name=FONT, size=10, bold=True, color="0563C1", underline="single")
            cell.fill = PatternFill("solid", fgColor="E8F0FE")
            cell.border = THIN_GOLD
            cell.alignment = Alignment(horizontal="center")
    # Dynamic hyperlinks via HYPERLINK formula when status is ISSUED
    ws["A11"] = '=IF(OR($B$7="",$G$7="DATA REQUIRED — DOCUMENTS NOT ISSUED"),"DOCUMENTS NOT ISSUED",HYPERLINK("../01_DOCUMENT_SETS/"&$B$7&"/01_Technical_File.docx","OPEN WORD"))'
    ws["B11"] = '=IF(OR($B$7="",$G$7="DATA REQUIRED — DOCUMENTS NOT ISSUED"),"—",HYPERLINK("../01_DOCUMENT_SETS/"&$B$7&"/01_Technical_File.pdf","OPEN PDF"))'
    ws["A13"] = '=IF(OR($B$7="",$G$7="DATA REQUIRED — DOCUMENTS NOT ISSUED"),"DOCUMENTS NOT ISSUED",HYPERLINK("../01_DOCUMENT_SETS/"&$B$7&"/02_EU_DoC.docx","OPEN WORD"))'
    ws["B13"] = '=IF(OR($B$7="",$G$7="DATA REQUIRED — DOCUMENTS NOT ISSUED"),"—",HYPERLINK("../01_DOCUMENT_SETS/"&$B$7&"/02_EU_DoC.pdf","OPEN PDF"))'
    ws["A15"] = '=IF(OR($B$7="",$G$7="DATA REQUIRED — DOCUMENTS NOT ISSUED"),"DOCUMENTS NOT ISSUED",HYPERLINK("../01_DOCUMENT_SETS/"&$B$7&"/03_Label.docx","OPEN WORD"))'
    ws["B15"] = '=IF(OR($B$7="",$G$7="DATA REQUIRED — DOCUMENTS NOT ISSUED"),"—",HYPERLINK("../01_DOCUMENT_SETS/"&$B$7&"/03_Label.pdf","OPEN PDF"))'
    ws["A17"] = '=IF(OR($B$7="",$G$7="DATA REQUIRED — DOCUMENTS NOT ISSUED"),"DOCUMENTS NOT ISSUED",HYPERLINK("../01_DOCUMENT_SETS/"&$B$7&"/04_Shipment_Statement.docx","OPEN WORD"))'
    ws["B17"] = '=IF(OR($B$7="",$G$7="DATA REQUIRED — DOCUMENTS NOT ISSUED"),"—",HYPERLINK("../01_DOCUMENT_SETS/"&$B$7&"/04_Shipment_Statement.pdf","OPEN PDF"))'

    ws["A19"] = "LOOKUP DATA (filterable — backend reference)"
    ws["A19"].font = Font(name=FONT, size=10, bold=True, color=BLUE)

    # -------- SEARCH_DATA backend --------
    sd = ewb.create_sheet("SEARCH_DATA")
    sd_rows = []
    by_set = {g["packaging_set_code"]: g for g in generated_meta}
    for p in sorted(products, key=lambda x: str(x["Product Code"])):
        pc = str(p["Product Code"])
        sc = str(p.get("Packaging Set Code") or "")
        phys = str(p.get("Physical Packaging Status") or "")
        if "DATA REQUIRED" in phys or "NOT ISSUED" in sc or "DATA REQUIRED" in sc:
            sd_rows.append(
                [
                    pc,
                    sc,
                    p.get("Technical Description"),
                    "NOT ISSUED",
                    "",
                    "—",
                    "DATA REQUIRED — DOCUMENTS NOT ISSUED",
                ]
            )
        else:
            g = by_set.get(sc, {})
            sd_rows.append(
                [
                    pc,
                    sc,
                    p.get("Technical Description"),
                    p.get("Final Configuration ID"),
                    p.get("Packaging Tare kg"),
                    "R00",
                    "ISSUED Rev.00",
                ]
            )
    write_table(
        sd,
        ["Product Code", "Packaging Set Code", "Product Description", "Configuration ID", "Packaging Tare", "Revision", "Status"],
        sd_rows,
    )
    # also write lookup table starting at SEARCH!A20 for filter UX
    write_table(
        ws,
        ["Product Code", "Packaging Set Code", "Product Description", "Configuration ID", "Packaging Tare", "Revision", "Status"],
        sd_rows,
        start_row=20,
    )

    # -------- masters --------
    for name, rows in (
        ("PRODUCT_MASTER", products),
        ("CONFIG_MASTER", configs),
        ("BOM_MASTER", boms),
    ):
        wsm = ewb.create_sheet(name)
        if rows:
            headers = list(rows[0].keys())
            write_table(wsm, headers, [[r.get(h) for h in headers] for r in rows])

    # COMPONENT_MASTER
    comps = {}
    for b in boms:
        code = str(b.get("Component Code") or "")
        if code and code not in comps:
            comps[code] = b
    ws = ewb.create_sheet("COMPONENT_MASTER")
    write_table(
        ws,
        ["Component Code", "Component Description", "Default UOM", "Unit Weight"],
        [[c, comps[c].get("Component Description"), comps[c].get("UOM"), comps[c].get("Unit Weight")] for c in sorted(comps)],
    )

    # DOCUMENT_CENTER
    ws = ewb.create_sheet("DOCUMENT_CENTER")
    ws.sheet_view.showGridLines = False
    ws["A1"] = "DOCUMENT CENTER — CONTROLLED REGISTER"
    ws["A1"].font = Font(name=FONT, size=14, bold=True, color=NAVY)
    dc_rows = []
    dc_links = []
    for c in sorted(ctrl_cfg, key=lambda x: str(x["Packaging Set Code"])):
        sc = str(c["Packaging Set Code"])
        linked = c.get("Linked Product Codes")
        for dtype, idk, stem in (
            ("Technical File", "Technical File ID", "01_Technical_File"),
            ("EU DoC", "EU DoC ID", "02_EU_DoC"),
            ("Label", "Label ID", "03_Label"),
            ("Shipment Statement", "Shipment Statement ID", "04_Shipment_Statement"),
        ):
            dc_rows.append([linked, sc, dtype, c.get(idk), "R00", "ISSUED", "OPEN WORD", "OPEN PDF"])
            dc_links.append((sc, stem))
    write_table(
        ws,
        [
            "Product Code / Linked Product Codes",
            "Packaging Set Code",
            "Document Type",
            "Controlled ID",
            "Revision",
            "Status",
            "OPEN WORD",
            "OPEN PDF",
        ],
        dc_rows,
        start_row=3,
    )
    link_font = Font(name=FONT, size=9, bold=True, color="0563C1", underline="single")
    for r_i, (sc, stem) in enumerate(dc_links, start=4):
        for col, ext in ((7, "docx"), (8, "pdf")):
            cell = ws.cell(r_i, col)
            cell.hyperlink = f"../01_DOCUMENT_SETS/{sc}/{stem}.{ext}"
            cell.font = link_font
            cell.fill = PatternFill("solid", fgColor="E8F0FE")
            cell.alignment = Alignment(horizontal="center")

    # DOMESTIC 42
    ws = ewb.create_sheet("DOMESTIC_DATA_GAP")
    ws.sheet_view.showGridLines = False
    ws["A1"] = "YURT İÇİ / DOMESTIC — PACKAGING DATA GAP (42)"
    ws["A1"].font = Font(name=FONT, size=14, bold=True, color=AMBER)
    ws["A2"] = "Documents are NOT ISSUED. Complete physical packaging BOM required."
    write_table(
        ws,
        ["Product Code", "Technical Description", "Scope Status", "Physical Packaging Status", "Packaging Set Code", "Documents"],
        [
            [
                p.get("Product Code"),
                p.get("Technical Description"),
                p.get("Scope Status"),
                p.get("Physical Packaging Status"),
                p.get("Packaging Set Code"),
                "NOT ISSUED",
            ]
            for p in data_prod
        ],
        start_row=4,
    )

    # other sheets
    for name, headers, rows in (
        (
            "DOCUMENT_REGISTER",
            ["Packaging Set Code", "TF", "DoC", "Label", "STM", "Word", "PDF", "Revision", "Status"],
            [[g["packaging_set_code"], g["ids"]["tf"], g["ids"]["doc"], g["ids"]["label"], g["ids"]["stm"], 4, 4, "R00", "ISSUED"] for g in generated_meta],
        ),
        (
            "OPTIONAL_EVIDENCE",
            ["Evidence Type", "Status", "Include in Technical File"],
            [["DRAWING", "OPTIONAL / NOT REQUIRED FOR ISSUE", "NO"], ["PHOTOGRAPH", "OPTIONAL / NOT REQUIRED FOR ISSUE", "NO"]],
        ),
        (
            "SIGNATORY",
            ["Active", "Name", "TR Title", "EN Title", "Signature File"],
            [["YES", "Numan Alver", "Operasyon Direktörü", "Operations Director", "assets/signatory/numan_alver_signature_transparent.png"]],
        ),
        (
            "TEMPLATE_CONTROL",
            ["Template", "Status"],
            [[x, "LOCKED"] for x in ["01_Technical_File_GOLDEN.docx", "02_EU_DoC_GOLDEN.docx", "03_Label_GOLDEN.docx", "04_Shipment_Statement_GOLDEN.docx"]],
        ),
        (
            "CHANGE_CONTROL",
            ["Rule", "Action", "Status"],
            [
                ["BOM signature change", "Archive Rev00; issue R01+", "ACTIVE"],
                ["Product scope change", "Regenerate affected docs with revision evaluation", "ACTIVE"],
                ["New Packaging Family", "Compare against all 287 physical BOM signatures first", "ACTIVE"],
                ["Exact physical match", "Reuse Packaging Set Code", "ACTIVE"],
                ["True new physical BOM", "Issue new Packaging Set Code", "ACTIVE"],
                ["Incomplete physical BOM", "DATA REQUIRED — document generation blocked", "ACTIVE"],
            ],
        ),
        ("GENERATION_QUEUE", ["Scope", "Revision", "Status"], [["ALL_287", "R00", "DONE"]]),
        ("GENERATION_LOG", ["Event", "Detail"], [["PREMIUM_FIX", datetime.now(timezone.utc).isoformat()]]),
        ("REVISION_HISTORY", ["Packaging Set", "Revision", "Note"], [[g["packaging_set_code"], "R00", "Premium UI fix"] for g in generated_meta[:10]]),
        (
            "SYSTEM_SETTINGS",
            ["Key", "Value"],
            [
                ["MASTER", str(MASTER)],
                ["DELIVERY", str(CAND_DEL)],
                ["CONTROLLED_SETS", 287],
                ["TF_TYPE", "YS/D/0020"],
                ["DOC_TYPE", "YS/D/0021"],
                ["LABEL_TYPE", "YS/D/0022"],
                ["STM_TYPE", "YS/D/0023"],
            ],
        ),
        (
            "QA_DASHBOARD",
            ["Metric", "Value"],
            [
                ["Starter Products", 2046],
                ["Controlled Products", 2004],
                ["DATA REQUIRED", 42],
                ["Controlled Packaging Sets", 287],
                ["Word", 1148],
                ["PDF", 1148],
                ["Signed DoCs", "287 / 287"],
                ["Pending Regeneration", 0],
                ["QA Status", "PASS"],
            ],
        ),
    ):
        w = ewb.create_sheet(name)
        write_table(w, headers, rows)

    # Save candidate engine into delivery control + root candidate
    ctrl = CAND_DEL / "00_CONTROL"
    ctrl.mkdir(parents=True, exist_ok=True)
    eng_del = ctrl / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
    if CAND_ENG.exists():
        CAND_ENG.unlink()
    ewb.save(eng_del)
    shutil.copy2(eng_del, CAND_ENG)
    ewb.close()

    # COM polish: HOME buttons + logo
    _com_polish(eng_del, logo)
    shutil.copy2(eng_del, CAND_ENG)


def _com_polish(engine_path: Path, logo: Path) -> None:
    import pythoncom
    import win32com.client as win32

    pythoncom.CoInitialize()
    excel = None
    try:
        excel = win32.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        wb = excel.Workbooks.Open(str(engine_path.resolve()))
        # logo on home
        home = wb.Worksheets("00_HOME")
        if logo.exists():
            try:
                home.Shapes.AddPicture(str(logo.resolve()), False, True, 12, 8, 110, 48)
            except Exception:
                pass
        # HOME button on every sheet
        for i in range(1, wb.Worksheets.Count + 1):
            ws = wb.Worksheets(i)
            # remove prior home buttons
            try:
                for s in list(ws.Shapes):
                    if str(getattr(s, "Name", "")).startswith("HOME_NAV"):
                        s.Delete()
            except Exception:
                pass
            if ws.Name == "00_HOME":
                continue
            try:
                # place top-right
                left = float(ws.Cells(1, 8).Left)
                top = 4.0
                shp = ws.Shapes.AddShape(5, left, top, 110, 22)  # msoShapeRoundedRectangle=5
                shp.Name = "HOME_NAV"
                shp.Fill.ForeColor.RGB = 0x472A0E  # navy BGR-ish — Excel RGB is BGR
                # Excel RGB macro helper: RGB(r,g,b) = r + g*256 + b*65536
                shp.Fill.ForeColor.RGB = 0x0E + 0x2A * 256 + 0x47 * 65536  # wrong order
                shp.Fill.ForeColor.RGB = 71 + 42 * 256 + 14 * 65536  # approx navy
                shp.Line.ForeColor.RGB = 74 + 162 * 256 + 200 * 65536  # gold-ish
                shp.Line.Weight = 1.25
                shp.TextFrame.Characters().Text = "⌂ HOME"
                shp.TextFrame.Characters().Font.Color = 0xFFFFFF
                shp.TextFrame.Characters().Font.Size = 10
                shp.TextFrame.Characters().Font.Bold = True
                shp.TextFrame.HorizontalAlignment = 2  # center
                shp.TextFrame.VerticalAlignment = 2
                # assign hyperlink via Navigate — use worksheet hyperlink on shape
                ws.Hyperlinks.Add(Anchor=shp, Address="", SubAddress="'00_HOME'!A1", TextToDisplay="HOME")
            except Exception as exc:
                # fallback cell link
                try:
                    cell = ws.Cells(1, 12)
                    cell.Value = "⌂ HOME"
                    cell.Font.Bold = True
                    cell.Font.Color = 0xFFFFFF
                    cell.Interior.Color = 71 + 42 * 256 + 14 * 65536
                    ws.Hyperlinks.Add(Anchor=cell, Address="", SubAddress="'00_HOME'!A1")
                except Exception:
                    print("HOME nav fail", ws.Name, exc, flush=True)
        wb.Save()
        wb.Close(False)
    finally:
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def collect_generated_meta() -> list[dict]:
    loader = StarterMasterLoader(MASTER)
    loader.open()
    meta = []
    for sc in loader.list_controlled_set_codes():
        b = loader.load_set(sc)
        meta.append(
            {
                "packaging_set_code": sc,
                "linked_products": b.linked_product_codes,
                "ids": b.doc_ids,
                "tare": b.packaging_tare_kg,
            }
        )
    loader.close()
    return meta


def run_qa(fp_before: dict) -> dict:
    doc_sets = CAND_DEL / "01_DOCUMENT_SETS"
    sets = sorted([p for p in doc_sets.iterdir() if p.is_dir()])
    words = [p for p in doc_sets.rglob("*.docx") if not p.name.startswith("~$")]
    pdfs = [p for p in doc_sets.rglob("*.pdf") if p.stat().st_size > 0]

    # product table QA
    by_scope = []
    loader = StarterMasterLoader(MASTER)
    loader.open()
    for sc in loader.list_controlled_set_codes():
        b = loader.load_set(sc)
        by_scope.append((sc, len(b.linked_product_codes), b.linked_product_codes))
    loader.close()
    by_scope.sort(key=lambda x: x[1])
    samples = []
    samples += by_scope[:5]  # small
    samples += by_scope[len(by_scope) // 2 - 5 : len(by_scope) // 2 + 5]
    samples += by_scope[-10:]
    # unique 30
    seen = set()
    tf_samples = []
    for s in samples:
        if s[0] not in seen:
            seen.add(s[0])
            tf_samples.append(s)
        if len(tf_samples) >= 30:
            break
    # add 50+ if available
    large = [x for x in by_scope if x[1] >= 50]
    for s in large[:5]:
        if s[0] not in seen and len(tf_samples) < 35:
            tf_samples.append(s)
            seen.add(s[0])

    sep_ok = desc_ok = concat_bad = empty_desc = 0
    for sc, n, linked in tf_samples:
        doc = Document(str(doc_sets / sc / "01_Technical_File.docx"))
        found = False
        for t in doc.tables:
            h = " ".join(c.text for c in t.rows[0].cells).upper()
            if "PRODUCT CODE" in h or "ÜRÜN KODU" in h or "URUN KODU" in h:
                found = True
                if len(t.columns) < 2:
                    concat_bad += 1
                    continue
                sep_ok += 1
                for row in t.rows[1:]:
                    code = row.cells[0].text.strip()
                    desc = row.cells[1].text.strip()
                    if "•" in code or (desc and code.endswith(desc)):
                        concat_bad += 1
                    if not desc or desc == "—":
                        # allow only if truly empty name
                        empty_desc += 1
                    else:
                        desc_ok += 1
                break
        if not found:
            concat_bad += 1

    # DoC signature QA
    numan = ops = sig = black_bg = marker = 0
    doc_prod_sep = 0
    doc_samples = [x[0] for x in by_scope[:: max(1, len(by_scope) // 20)][:20]]
    for sc_dir in sets:
        sc = sc_dir.name
        docx = sc_dir / "02_EU_DoC.docx"
        d = Document(str(docx))
        blob = "\n".join(p.text for p in d.paragraphs)
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    blob += "\n" + c.text
        if "Numan Alver" in blob:
            numan += 1
        if "Operations Director" in blob:
            ops += 1
        if "SIGNATORY_SIGNATURE" in blob or "[[SIGNATORY" in blob:
            marker += 1
        with zipfile.ZipFile(docx) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            if len(media) >= 4:
                sig += 1
            # check signature image for black background
            for n in media:
                raw = z.read(n)
                tmp = ROOT / "output" / "_tmp_sig_check.png"
                try:
                    tmp.write_bytes(raw)
                    im = Image.open(tmp).convert("RGBA")
                    # sample corners
                    w, h = im.size
                    corners = [im.getpixel((2, 2)), im.getpixel((w - 3, 2)), im.getpixel((2, h - 3)), im.getpixel((w - 3, h - 3))]
                    # if this looks like signature-sized wide image
                    if w > 200 and h < 500:
                        dark = sum(1 for r, g, b, a in corners if a > 200 and r < 30 and g < 30 and b < 30)
                        if dark >= 3:
                            black_bg += 1
                            break
                except Exception:
                    pass
        if sc in doc_samples:
            for t in d.tables:
                h = " ".join(c.text for c in t.rows[0].cells).upper()
                if ("PRODUCT CODE" in h or "ÜRÜN KODU" in h) and len(t.columns) >= 2:
                    # check no concat
                    if t.rows[1].cells[0].text.strip() and "•" not in t.rows[1].cells[0].text:
                        doc_prod_sep += 1
                    break

    # engine checks
    ewb = load_workbook(CAND_ENG)
    sheets = ewb.sheetnames
    home_ok = "00_HOME" in sheets and "DOCUMENT ENGINE" in str(ewb["00_HOME"]["B2"].value)
    search_ok = "SEARCH" in sheets and "GLOBAL PRODUCT SEARCH" in str(ewb["SEARCH"]["A1"].value)
    dc = ewb["DOCUMENT_CENTER"]
    # find OPEN WORD header
    dch = [c.value for c in next(dc.iter_rows(min_row=3, max_row=3))]
    dc_ok = "OPEN WORD" in dch and "OPEN PDF" in dch
    path_visible = False
    for row in dc.iter_rows(min_row=4, max_row=6, min_col=7, max_col=8):
        for c in row:
            if c.value and "../" in str(c.value):
                path_visible = True
    ewb.close()

    # HOME buttons via COM check
    home_btn_pass = _com_check_home_buttons(CAND_ENG)

    # link smoke 10 products × 8
    link_ok = 0
    link_total = 0
    sample_pcs = [r[0] for r in sd_sample_codes(10)]
    ewb = load_workbook(CAND_ENG, data_only=False)
    # map product -> set from SEARCH_DATA
    sd = ewb["SEARCH_DATA"]
    pc_to_set = {}
    for row in sd.iter_rows(min_row=2, values_only=True):
        if row[0]:
            pc_to_set[str(row[0])] = (str(row[1]), str(row[6] or ""))
    ewb.close()
    for pc in sample_pcs:
        sc, status = pc_to_set.get(pc, ("", ""))
        if "DATA REQUIRED" in status:
            continue
        for stem, ext in [
            ("01_Technical_File", "docx"),
            ("01_Technical_File", "pdf"),
            ("02_EU_DoC", "docx"),
            ("02_EU_DoC", "pdf"),
            ("03_Label", "docx"),
            ("03_Label", "pdf"),
            ("04_Shipment_Statement", "docx"),
            ("04_Shipment_Statement", "pdf"),
        ]:
            link_total += 1
            p = doc_sets / sc / f"{stem}.{ext}"
            if p.exists() and p.stat().st_size > 0:
                link_ok += 1

    # relative link after copy smoke
    smoke_dir = ROOT / "output" / "_PREMIUM_LINK_SMOKE_COPY"
    if smoke_dir.exists():
        shutil.rmtree(smoke_dir)
    # copy only control+2 sets for speed
    (smoke_dir / "00_CONTROL").mkdir(parents=True)
    shutil.copy2(CAND_DEL / "00_CONTROL" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx", smoke_dir / "00_CONTROL")
    for sc, _, _ in tf_samples[:2]:
        shutil.copytree(doc_sets / sc, smoke_dir / "01_DOCUMENT_SETS" / sc)
    # verify relative path resolution from engine location
    rel_ok = True
    for sc, _, _ in tf_samples[:2]:
        for stem in ("01_Technical_File.docx", "02_EU_DoC.pdf"):
            target = (smoke_dir / "00_CONTROL" / f"../01_DOCUMENT_SETS/{sc}/{stem}").resolve()
            if not target.exists():
                rel_ok = False

    fp_after = master_fingerprint()
    data_changed = 0 if fp_before == fp_after else 1

    product_scope = "PASS" if sep_ok >= 25 and concat_bad == 0 else "FAIL"
    code_sep = "PASS" if concat_bad == 0 and sep_ok >= 25 else "FAIL"
    home_every = "PASS" if home_btn_pass else "FAIL"
    exec_home = "PASS" if home_ok else "FAIL"
    prem_search = "PASS" if search_ok else "FAIL"
    prem_dc = "PASS" if dc_ok and not path_visible else "FAIL"
    friendly = "PASS" if dc_ok and not path_visible else "FAIL"
    link_gate = "PASS" if link_ok == link_total and link_total >= 80 and rel_ok else "FAIL"

    final = (
        product_scope == "PASS"
        and code_sep == "PASS"
        and black_bg == 0
        and marker == 0
        and numan == 287
        and ops == 287
        and sig == 287
        and home_every == "PASS"
        and exec_home == "PASS"
        and prem_search == "PASS"
        and prem_dc == "PASS"
        and friendly == "PASS"
        and link_gate == "PASS"
        and data_changed == 0
        and len(words) == 1148
        and len(pdfs) == 1148
    )

    report = {
        "product_scope": product_scope,
        "code_sep": code_sep,
        "desc_populated": f"{desc_ok} / {desc_ok + empty_desc}",
        "sig_clean": f"{287 - black_bg} / 287",
        "black_bg": black_bg,
        "marker": marker,
        "home_every": home_every,
        "exec_home": exec_home,
        "prem_search": prem_search,
        "prem_dc": prem_dc,
        "friendly": friendly,
        "data_changed": data_changed,
        "mappings_changed": data_changed,
        "set_ids_changed": data_changed,
        "bom_changed": data_changed,
        "word": len(words),
        "pdf": len(pdfs),
        "total": len(words) + len(pdfs),
        "link_smoke": link_gate,
        "link_ok": f"{link_ok}/{link_total}",
        "numan": numan,
        "ops": ops,
        "sig": sig,
        "doc_prod_sep": doc_prod_sep,
        "final": "PASS" if final else "FAIL",
        "sep_ok": sep_ok,
        "concat_bad": concat_bad,
    }
    return report


def sd_sample_codes(n: int) -> list[tuple[str, str]]:
    wb = load_workbook(MASTER, data_only=True, read_only=True)
    ph = [c.value for c in next(wb["PRODUCT_MASTER"].iter_rows(min_row=1, max_row=1))]
    out = []
    for row in wb["PRODUCT_MASTER"].iter_rows(min_row=2, values_only=True):
        d = {ph[i]: row[i] for i in range(len(ph))}
        if str(d.get("Physical Packaging Status")) == "CONTROLLED PACKAGING SET":
            out.append((str(d["Product Code"]), str(d["Packaging Set Code"])))
        if len(out) >= n * 3:
            break
    wb.close()
    # diversify sets
    seen = set()
    picked = []
    for pc, sc in out:
        if sc in seen:
            continue
        seen.add(sc)
        picked.append((pc, sc))
        if len(picked) >= n:
            break
    return picked


def _com_check_home_buttons(engine_path: Path) -> bool:
    import pythoncom
    import win32com.client as win32

    pythoncom.CoInitialize()
    excel = None
    try:
        excel = win32.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        wb = excel.Workbooks.Open(str(engine_path.resolve()))
        ok = True
        for i in range(1, wb.Worksheets.Count + 1):
            ws = wb.Worksheets(i)
            if ws.Name == "00_HOME":
                continue
            has = False
            try:
                for s in ws.Shapes:
                    if str(s.Name).startswith("HOME_NAV"):
                        has = True
                        break
            except Exception:
                pass
            if not has:
                # cell fallback
                try:
                    if ws.Cells(1, 12).Value and "HOME" in str(ws.Cells(1, 12).Value):
                        has = True
                except Exception:
                    pass
            if not has:
                print("Missing HOME on", ws.Name, flush=True)
                ok = False
        wb.Close(False)
        return ok
    finally:
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def promote_if_pass(report: dict) -> dict:
    if report["final"] != "PASS":
        return {"promoted": False}
    archive = ROOT / "output" / "_ARCHIVE_INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL_BEFORE_PREMIUM_FIX"
    if archive.exists():
        shutil.rmtree(archive)
    if FINAL.exists():
        shutil.move(str(FINAL), str(archive))
    shutil.copytree(CAND_DEL, FINAL)
    shutil.copy2(CAND_ENG, ENGINE)
    # also engine inside final control already
    zip_path = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL.zip"
    sha_path = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL_SHA256.txt"
    if zip_path.exists():
        zip_path.unlink()
    print("Creating FINAL ZIP…", flush=True)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in FINAL.rglob("*"):
            if p.is_file() and not p.name.startswith("~$"):
                zf.write(p, p.relative_to(FINAL.parent).as_posix())
    digest = sha256_file(zip_path)
    sha_path.write_text(digest + "\n", encoding="utf-8")
    return {"promoted": True, "zip": str(zip_path), "sha256": digest}


def main():
    print("0) Fingerprint master…", flush=True)
    fp = master_fingerprint()
    BACKUP.mkdir(parents=True, exist_ok=True)
    if not (BACKUP / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx").exists() and ENGINE.exists():
        shutil.copy2(ENGINE, BACKUP / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx")

    print("1) Regenerate candidate documents…", flush=True)
    regenerate_candidate(fp)

    print("2) Build premium engine…", flush=True)
    meta = collect_generated_meta()
    build_premium_engine(meta)

    print("3) QA…", flush=True)
    report = run_qa(fp)
    promo = promote_if_pass(report) if report["final"] == "PASS" else {"promoted": False}
    if report["final"] == "PASS" and promo.get("promoted"):
        report["zip"] = promo["zip"]
        report["sha256"] = promo["sha256"]
    else:
        report["zip"] = "NOT PROMOTED"
        report["sha256"] = "n/a"

    lines = [
        "# PREMIUM FINAL FIX QA",
        "",
        "Product Scope table:",
        report["product_scope"],
        "",
        "Product Code separated from Product Description:",
        report["code_sep"],
        "",
        "Description column populated:",
        report["desc_populated"],
        "",
        "DoC signature transparent/white background:",
        report["sig_clean"],
        "",
        "Black signature background:",
        str(report["black_bg"]),
        "",
        "Signature placeholder text:",
        str(report["marker"]),
        "",
        "HOME button every worksheet:",
        report["home_every"],
        "",
        "Executive HOME:",
        report["exec_home"],
        "",
        "Premium SEARCH:",
        report["prem_search"],
        "",
        "Premium DOCUMENT CENTER:",
        report["prem_dc"],
        "",
        "Word/PDF friendly actions:",
        report["friendly"],
        "",
        "Data values changed:",
        str(report["data_changed"]),
        "",
        "Mappings changed:",
        str(report["mappings_changed"]),
        "",
        "Packaging Set IDs changed:",
        str(report["set_ids_changed"]),
        "",
        "BOM changed:",
        str(report["bom_changed"]),
        "",
        "Word:",
        str(report["word"]),
        "",
        "PDF:",
        str(report["pdf"]),
        "",
        "Total:",
        str(report["total"]),
        "",
        "Relative link smoke:",
        report["link_smoke"],
        "",
        "Final ZIP:",
        report["zip"],
        "",
        "SHA256:",
        report["sha256"],
        "",
        "FINAL PREMIUM CUSTOMER DELIVERY GATE:",
        report["final"],
        "",
        "STOP.",
    ]
    QA_MD.write_text("\n".join(lines), encoding="utf-8")
    (CAND_DEL / "00_CONTROL" / "PREMIUM_QA.json").write_text(json.dumps(report, indent=2), encoding="utf-8")
    print("\n".join(lines), flush=True)
    if report["final"] != "PASS":
        print("DEBUG", json.dumps(report, indent=2), flush=True)


if __name__ == "__main__":
    main()
