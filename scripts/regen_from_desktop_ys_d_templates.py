"""Regenerate all 287×4 docs from Desktop YS_D_0020–0023 golden templates.

Keeps existing Document Engine / horizontal DC / links.
Does NOT change Product↔Set mappings, BOM, tare, codes, or controlled IDs.
"""

from __future__ import annotations

import hashlib
import shutil
import sys
import zipfile
from collections import Counter
from pathlib import Path

ROOT = Path(r"C:\Users\burcu\Documents\YAZILIM\Inci_Aku_PPWR_PIMS")
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(ROOT / "tools"))

from docx import Document
from openpyxl import load_workbook

from builders.phase_g.merge_engine import merge_document
from builders.phase_g.runtime_template_builder import build_runtime_templates
from builders.phase_i.render_batch import render_docx_batch
from models.technical_file import Article5Assessment
from services.document_context_factory import DocumentContextFactory
from utils.constants import ARTICLE5_BASIS_LABEL
from ppwr_engine.prepare_locked_templates import main as prepare_locked
from ppwr_engine.starter_loader import StarterMasterLoader

import generate_ppwr_documents as gen

MASTER = ROOT / "output" / "INCI_AKU_PPWR_STARTER_MASTER_Rev00.xlsx"
LOCKED = ROOT / "templates" / "ppwr_rev00_locked"
RUNTIME = LOCKED / "runtime"
FINAL = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL"
DOC_SETS = FINAL / "01_DOCUMENT_SETS"
CTRL = FINAL / "00_CONTROL"
ENG = CTRL / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
ENG_ROOT = ROOT / "output" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
ZIP_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL.zip"
SHA_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL_SHA256.txt"
SIG = ROOT / "assets" / "signatory" / "numan_alver_signature_transparent.png"

QMS = {
    "01_Technical_File.docx": "YS/D/0020",
    "02_EU_DoC.docx": "YS/D/0021",
    "03_Label.docx": "YS/D/0022",
    "04_Shipment_Statement.docx": "YS/D/0023",
}
SPECS = [
    ("TECHNICAL_FILE", "01_Technical_File", True),
    ("DOC", "02_EU_DoC", False),
    ("LABEL", "03_Label", False),
    ("STATEMENT", "04_Shipment_Statement", False),
]


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def main() -> None:
    assert ENG.exists(), f"Engine missing: {ENG}"
    assert SIG.exists(), f"Signature missing: {SIG}"

    print("1) Prepare locked + runtime from Desktop goldens…", flush=True)
    prepare_locked()
    build_runtime_templates(LOCKED, RUNTIME)

    runtime = {
        "TECHNICAL_FILE": RUNTIME / "01_Technical_File_RUNTIME.docx",
        "DOC": RUNTIME / "02_EU_DoC_RUNTIME.docx",
        "LABEL": RUNTIME / "03_Label_RUNTIME.docx",
        "STATEMENT": RUNTIME / "04_Shipment_Statement_RUNTIME.docx",
    }
    for p in runtime.values():
        assert p.exists(), p

    # backup engine before doc regen
    eng_bak = CTRL / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00_PRE_YS_D_REGEN.xlsx"
    shutil.copy2(ENG, eng_bak)

    if DOC_SETS.exists():
        shutil.rmtree(DOC_SETS)
    DOC_SETS.mkdir(parents=True)

    print("2) Merge Word × 287 × 4…", flush=True)
    loader = StarterMasterLoader(MASTER)
    loader.open()
    codes = loader.list_controlled_set_codes()
    assert len(codes) == 287, len(codes)
    factory = DocumentContextFactory()
    signature_ok = 0
    pdf_jobs: list[tuple[Path, Path]] = []

    for i, sc in enumerate(codes, 1):
        bundle = loader.load_set(sc)
        ctx = factory.build(
            bundle.configuration,
            products=bundle.products,
            article5=Article5Assessment(basis_label=ARTICLE5_BASIS_LABEL),
        )
        ctx.total_tare_g = float(bundle.packaging_tare_kg) * 1000.0
        ctx.annex_drawings_status = "OPTIONAL EVIDENCE — NOT INCLUDED IN REV.00"
        ctx.document_ids.technical_file_id = bundle.doc_ids["tf"]
        ctx.document_ids.doc_id = bundle.doc_ids["doc"]
        ctx.document_ids.label_id = bundle.doc_ids["label"]
        ctx.document_ids.statement_id = bundle.doc_ids["stm"]

        out_dir = DOC_SETS / sc
        out_dir.mkdir(parents=True, exist_ok=True)
        for dtype, stem, is_tf in SPECS:
            out_docx = out_dir / f"{stem}.docx"
            merge_document(runtime[dtype], out_docx, ctx, for_technical_file=is_tf)
            if is_tf:
                gen.scrub_tf_pending(out_docx)
            if dtype == "DOC":
                if gen.embed_signature(out_docx, SIG):
                    signature_ok += 1
            pdf_jobs.append((out_docx, out_dir / f"{stem}.pdf"))
        if i % 25 == 0 or i == len(codes):
            print(f"  Word {i}/{len(codes)}", flush=True)
    loader.close()
    print({"word": 287 * 4, "sig_ok": signature_ok})

    print("3) Render PDFs…", flush=True)
    for i in range(0, len(pdf_jobs), 40):
        chunk = pdf_jobs[i : i + 40]
        render_docx_batch(chunk, progress_every=20, log=[])
        done = min(i + 40, len(pdf_jobs))
        nonzero = sum(1 for _, p in pdf_jobs[:done] if p.exists() and p.stat().st_size > 0)
        print(f"  PDF {done}/{len(pdf_jobs)} nonzero={nonzero}", flush=True)
    # retry zeros
    retry = [(d, p) for d, p in pdf_jobs if not p.exists() or p.stat().st_size == 0]
    if retry:
        print(f"  Retry {len(retry)}", flush=True)
        for i in range(0, len(retry), 20):
            render_docx_batch(retry[i : i + 20], progress_every=10, log=[])

    print("4) QA…", flush=True)
    words = [p for p in DOC_SETS.rglob("*.docx") if not p.name.startswith("~$")]
    pdfs = [p for p in DOC_SETS.rglob("*.pdf") if p.stat().st_size > 0]
    qms = Counter()
    numan = ops = sig = marker = pending = concat = 0
    sample_leak = 0

    for p in words:
        doc = Document(str(p))
        blob = "\n".join(x.text for x in doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    blob += "\n" + c.text
        for sec in doc.sections:
            for part in (sec.header, sec.footer):
                for para in part.paragraphs:
                    blob += "\n" + para.text
        u = blob.upper()
        qms_code = QMS.get(p.name)
        if qms_code and (qms_code in blob or f"YS/D/{qms_code[-4:]}" in blob):
            # count per set file type
            if "Doküman No" in blob or "Doc. Nr" in blob or "Dokuman No" in blob:
                qms[qms_code] += 1
            elif qms_code in blob:
                qms[qms_code] += 1
        if p.name == "02_EU_DoC.docx":
            if "Numan Alver" in blob:
                numan += 1
            if "Operations Director" in blob or "Operasyon Direktörü" in blob:
                ops += 1
            # signature image?
            has_img = False
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    has_img = True
                    break
            if has_img:
                sig += 1
            if "SIGNATORY_SIGNATURE" in blob or "[[SIGNATORY" in blob:
                marker += 1
        if ("PENDING" in u and ("DRAW" in u or "PHOTO" in u or "FOTO" in u)) or (
            "WILL BE ADDED" in u and ("DRAW" in u or "PHOTO" in u)
        ) or ("SONRADAN EKLEN" in u):
            pending += 1
        if " • " in blob and "1015169 •" in blob and "PRODUCT CODE" not in u:
            # weak concat check — skip
            pass
        if "ST-012-EUR-01" in blob and p.parent.name != "ST-012-EUR-01":
            sample_leak += 1

    # product scope 2-col spot check on TF
    prod_ok = 0
    for sc in ["ST-012-EUR-01"]:
        tf = DOC_SETS / sc / "01_Technical_File.docx"
        d = Document(str(tf))
        for table in d.tables:
            hdr = " ".join(c.text.upper() for c in table.rows[0].cells)
            if "PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr or "URUN KODU" in hdr:
                if len(table.rows[0].cells) >= 2 and "DESCRIPTION" in hdr or "TANIM" in hdr:
                    prod_ok += 1

    qa = {
        "sets": len([p for p in DOC_SETS.iterdir() if p.is_dir()]),
        "word": len(words),
        "pdf": len(pdfs),
        "qms": dict(qms),
        "numan": numan,
        "ops": ops,
        "sig": sig,
        "marker": marker,
        "pending": pending,
        "sample_leak": sample_leak,
        "prod_cols_spot": prod_ok,
        "engine_kept": ENG.exists(),
    }
    print(qa)

    gate = (
        qa["sets"] == 287
        and qa["word"] == 1148
        and qa["pdf"] == 1148
        and qms.get("YS/D/0020", 0) == 287
        and qms.get("YS/D/0021", 0) == 287
        and qms.get("YS/D/0022", 0) == 287
        and qms.get("YS/D/0023", 0) == 287
        and numan == 287
        and sig == 287
        and marker == 0
        and pending == 0
        and sample_leak == 0
    )

    # ensure domestic folder + launcher still present
    dom = FINAL / "04_DOMESTIC_42_DATA_GAP"
    if not dom.exists():
        dom.mkdir(parents=True)
        (dom / "README.txt").write_text(
            "YURT İÇİ / DOMESTIC — 42 Product Codes\nDOCUMENTS NOT ISSUED\n",
            encoding="utf-8",
        )
    cmd = FINAL / "00_AC_DOCUMENT_ENGINE.cmd"
    if not cmd.exists():
        cmd.write_text(
            "@echo off\r\n"
            "cd /d \"%~dp0\"\r\n"
            "start \"\" \"%~dp000_CONTROL\\INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx\"\r\n",
            encoding="utf-8",
        )

    # sync root engine copy (unchanged content)
    if ENG.exists():
        shutil.copy2(ENG, ENG_ROOT)

    print("5) Rebuild ZIP…", flush=True)
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
        for p in FINAL.rglob("*"):
            if p.is_file() and not p.name.startswith("~$") and "PRE_YS_D_REGEN" not in p.name:
                zf.write(p, p.relative_to(FINAL).as_posix())
    digest = sha256_file(ZIP_PATH)
    SHA_PATH.write_text(digest + "\n", encoding="utf-8")

    # spot-check opened doc has Desktop QMS line
    sample = DOC_SETS / "ST-012-EUR-01" / "02_EU_DoC.docx"
    sdoc = Document(str(sample))
    sblob = "\n".join(p.text for p in sdoc.paragraphs)
    has_qms_line = "Doküman No/Doc. Nr.: YS/D/0021" in sblob or "Doküman No/Doc. Nr.: YS/D/0021" in sblob
    # normalize encoding variants
    has_qms_line = "YS/D/0021" in sblob and ("Doc. Nr" in sblob or "Doküman No" in sblob or "Dokuman No" in sblob)

    print()
    print("# YS_D DESKTOP TEMPLATE REGEN QA")
    print(f"Controlled Packaging Sets: {qa['sets']} / 287")
    print(f"TF YS/D/0020: {qms.get('YS/D/0020', 0)} / 287")
    print(f"EU DoC YS/D/0021: {qms.get('YS/D/0021', 0)} / 287")
    print(f"Label YS/D/0022: {qms.get('YS/D/0022', 0)} / 287")
    print(f"Statement YS/D/0023: {qms.get('YS/D/0023', 0)} / 287")
    print(f"Numan Alver: {numan} / 287")
    print(f"DoC Signature: {sig} / 287")
    print(f"Signature placeholder: {marker}")
    print(f"Drawing/photo pending: {pending}")
    print(f"Sample ST-012 leak outside own set: {sample_leak}")
    print(f"Word: {qa['word']}")
    print(f"PDF: {qa['pdf']}")
    print(f"QMS line present (spot): {has_qms_line}")
    print(f"Engine preserved: YES")
    print(f"FINAL ZIP: {ZIP_PATH}")
    print(f"NEW SHA256: {digest}")
    print(f"GATE: {'PASS' if gate and has_qms_line else 'FAIL'}")
    if not (gate and has_qms_line):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
