"""Representative packaging component photo annex for Technical Files."""

from __future__ import annotations

import csv
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PHOTO_ROOT = ROOT / "assets" / "representative_component_photos"
MAPPING = PHOTO_ROOT / "PHOTO_MAPPING_INDEX.csv"
LOG = ROOT / "output" / "_photo_annex_mapping.log"

NAVY = RGBColor(0x0E, 0x2A, 0x47)
SCOPE_DIRS = {
    "STARTER": "01_STARTER_GENERAL",
    "INDUSTRIAL": "02_INDUSTRIAL",
    "CONTAINER": "03_CONTAINER_LOADING",
    "COMPONENT": "04_COMPONENT_SPARE",
    "COMPONENT_SPARE": "04_COMPONENT_SPARE",
    "COMPONENT_1ROW": "04_COMPONENT_SPARE",
    "COMPONENT_2ROW": "04_COMPONENT_SPARE",
    "COMMON": "01_STARTER_GENERAL",
}
SCOPE_ALIASES = {
    "COMPONENT": "COMPONENT",
    "COMPONENT_SPARE": "COMPONENT",
    "CONTAINER_LOADING": "CONTAINER",
}


@dataclass(frozen=True)
class PhotoHit:
    component_code: str
    name_tr: str
    name_en: str
    path: Path
    note: str = ""


_MAPPING_CACHE: list[dict] | None = None


def _load_mapping() -> list[dict]:
    global _MAPPING_CACHE
    if _MAPPING_CACHE is not None:
        return _MAPPING_CACHE
    if not MAPPING.exists():
        _MAPPING_CACHE = []
        return _MAPPING_CACHE
    rows: list[dict] = []
    with MAPPING.open("r", encoding="utf-8-sig", newline="") as f:
        # delimiter may be ; or ,
        sample = f.read(2048)
        f.seek(0)
        delim = ";" if sample.count(";") >= sample.count(",") else ","
        reader = csv.DictReader(f, delimiter=delim)
        for row in reader:
            rows.append({(k or "").strip(): (v or "").strip() for k, v in row.items()})
    _MAPPING_CACHE = rows
    return _MAPPING_CACHE


def _norm(s: str) -> str:
    return (
        (s or "")
        .upper()
        .replace("İ", "I")
        .replace("Ş", "S")
        .replace("Ğ", "G")
        .replace("Ü", "U")
        .replace("Ö", "O")
        .replace("Ç", "C")
    )


def clear_mapping_cache() -> None:
    global _MAPPING_CACHE
    _MAPPING_CACHE = None


def resolve_photos(
    *,
    scope: str,
    bom_lines: list[dict],
    prefer_prefix: str | None = None,
) -> list[PhotoHit]:
    """bom_lines: dicts with component_code, description (and optional name_tr/name_en)."""
    mapping = _load_mapping()
    scope_key = SCOPE_ALIASES.get(scope.upper(), scope.upper())
    preferred = SCOPE_DIRS.get(scope_key, SCOPE_DIRS.get(scope.upper(), "01_STARTER_GENERAL"))
    hits: list[PhotoHit] = []
    seen_files: set[str] = set()
    missing: list[str] = []
    pref = (prefer_prefix or "").upper()

    for line in bom_lines:
        code = str(line.get("component_code") or "").strip()
        desc = str(line.get("description") or line.get("component_description") or "").strip()
        # strip already-bilingual EN half for matching
        if " / " in desc:
            desc_match = desc.split(" / ", 1)[0].strip()
        else:
            desc_match = desc.split("\n", 1)[0].strip()
        name_tr = str(line.get("name_tr") or desc_match).strip()
        name_en = str(line.get("name_en") or "").strip()
        if not code and not desc_match:
            continue

        photo_rel = None
        note = ""
        nd = _norm(desc_match)
        nc = _norm(code)
        candidates: list[tuple[int, dict, str]] = []
        for row in mapping:
            row_scope = _norm(row.get("Scope") or row.get("SCOPE") or "")
            src = (
                row.get("Source Component Description")
                or row.get("Source Component")
                or row.get("Component Description")
                or row.get("Component")
                or ""
            )
            if not src:
                continue
            ns = _norm(src)
            matched = False
            match_q = 9
            if ns and ns == nd:
                matched = True
                match_q = 0
            elif ns and (ns in nd):
                matched = True
                match_q = 1
            elif nc and nc in ns:
                matched = True
                match_q = 2
            else:
                token = ns.split(",")[0].split()[0] if ns else ""
                # require token word-boundary style: token must appear as whole token in desc
                if token and len(token) >= 4:
                    desc_tokens = set(nd.replace(",", " ").replace("*", " ").split())
                    if token in desc_tokens:
                        matched = True
                        match_q = 3
            if not matched:
                continue
            photo_name = (row.get("Photo File") or row.get("Photo") or "").replace("\\", "/").split("/")[-1].upper()
            if pref and photo_name and not photo_name.startswith(pref):
                rank_boost = 5
            else:
                rank_boost = 0
            if row_scope in (_norm(scope_key), "COMMON", "ALL", ""):
                rank = 0 + rank_boost
            elif scope_key in row_scope or row_scope in scope_key:
                rank = 1 + rank_boost
            else:
                rank = 2 + rank_boost
            candidates.append((rank, match_q, -len(ns), row, ns))
        if candidates:
            candidates.sort(key=lambda x: (x[0], x[1], x[2]))
            best = candidates[0][3]
            photo_rel = best.get("Photo File") or best.get("Photo") or ""
            note = best.get("Note") or best.get("Mapping Type") or ""

        path = None
        if photo_rel:
            cand = PHOTO_ROOT / photo_rel.replace("/", "\\")
            if cand.exists():
                path = cand
        if path is None:
            tokens = [t for t in _norm(desc_match).replace(",", " ").split() if len(t) >= 4][:3]
            search_dirs = [
                PHOTO_ROOT / preferred,
                PHOTO_ROOT / "01_STARTER_GENERAL",
                PHOTO_ROOT / "02_INDUSTRIAL",
                PHOTO_ROOT / "03_CONTAINER_LOADING",
                PHOTO_ROOT / "04_COMPONENT_SPARE",
            ]
            for d in search_dirs:
                if not d.exists():
                    continue
                for img in sorted(d.glob("*.*")):
                    if img.suffix.lower() not in {".jpg", ".jpeg", ".png", ".webp"}:
                        continue
                    if pref and not img.name.upper().startswith(pref):
                        continue
                    iname = _norm(img.stem)
                    if tokens and sum(1 for t in tokens if t in iname) >= min(2, len(tokens)):
                        path = img
                        note = "FALLBACK_FILENAME"
                        break
                if path:
                    break

        if path is None:
            missing.append(f"{code}|{desc_match}")
            continue
        key = str(path.resolve()).lower()
        if key in seen_files:
            continue
        seen_files.add(key)
        hits.append(
            PhotoHit(
                component_code=code or "N/A",
                name_tr=name_tr or desc_match or path.stem,
                name_en=name_en or desc_match or path.stem,
                path=path,
                note=note,
            )
        )

    if missing:
        LOG.parent.mkdir(parents=True, exist_ok=True)
        with LOG.open("a", encoding="utf-8") as f:
            f.write(f"scope={scope} missing={len(missing)}\n")
            for m in missing[:50]:
                f.write(f"  {m}\n")
    return hits


ANNEX_MARKERS = (
    "Representative Packaging Component Photos",
    "Temsilî Ambalaj Bileşeni Görselleri",
    "Temsili Ambalaj Bileseni Gorselleri",
)


def _strip_photo_annex(doc: Document) -> bool:
    """Remove existing photo annex from end of document (heading + following content)."""
    body = doc.element.body
    children = list(body)
    cut_idx = None
    for i, child in enumerate(children):
        text = "".join(node.text or "" for node in child.iter() if hasattr(node, "text"))
        if any(m in text for m in ANNEX_MARKERS):
            cut_idx = i
            # also drop preceding page-break-only paragraph if present
            if i > 0:
                prev = children[i - 1]
                prev_text = "".join(node.text or "" for node in prev.iter() if hasattr(node, "text")).strip()
                if not prev_text:
                    cut_idx = i - 1
            break
    if cut_idx is None:
        return False
    for child in children[cut_idx:]:
        # keep sectPr at end
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)
    return True


def replace_photo_annex(docx_path: Path, photos: list[PhotoHit], *, title_extra: str = "") -> int:
    """Strip any prior annex and append fresh photos."""
    doc = Document(str(docx_path))
    _strip_photo_annex(doc)
    doc.save(str(docx_path))
    if not photos:
        return 0
    return append_photo_annex(docx_path, photos, title_extra=title_extra)


def append_photo_annex(docx_path: Path, photos: list[PhotoHit], *, title_extra: str = "") -> int:
    """Append informational photo annex page(s) at end of Technical File."""
    if not photos:
        return 0
    doc = Document(str(docx_path))
    # page break
    doc.add_page_break()
    h = doc.add_paragraph()
    r = h.add_run(
        "EK / ANNEX — Representative Packaging Component Photos\n"
        "Temsilî Ambalaj Bileşeni Görselleri"
    )
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    r.font.name = "Tahoma"
    note = doc.add_paragraph()
    nr = note.add_run(
        "Bu ek bilgilendirme amaçlıdır; zorunlu ana teknik içerik değildir. "
        "Görseller temsilîdir (Representative Image). "
        f"{title_extra}".strip()
    )
    nr.font.size = Pt(9)
    nr.font.name = "Tahoma"
    nr.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    # 2-column visual grid via table
    cols = 2
    rows_n = (len(photos) + cols - 1) // cols
    table = doc.add_table(rows=rows_n, cols=cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for idx, photo in enumerate(photos):
        r_i, c_i = divmod(idx, cols)
        cell = table.cell(r_i, c_i)
        # clear
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(str(photo.path), width=Cm(7.2))
        except Exception:
            p.add_run(f"[image missing: {photo.path.name}]")
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t = cap.add_run(
            f"{photo.component_code}\n{photo.name_tr}"
        )
        t.font.size = Pt(8)
        t.font.name = "Tahoma"
        t.font.color.rgb = RGBColor(0x1C, 0x24, 0x30)
        cap2 = cell.add_paragraph()
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t2 = cap2.add_run(f"{photo.name_en}\nTemsilî Görsel / Representative Image")
        t2.italic = True
        t2.font.size = Pt(8)
        t2.font.name = "Tahoma"
        t2.font.color.rgb = RGBColor(0x1C, 0x24, 0x30)

    doc.save(str(docx_path))
    return len(photos)
