"""Build one PPWR 4-doc pack from YS/D runtime golden templates (İnci Akü format).

- TR normal / EN italic in product + BOM tables
- Date 11.08.2026
- Photo annex
- No Word COM
"""

from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

import sys

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from bilingual_format import (  # noqa: E402
    set_bilingual_cell,
    translate_component,
    translate_product,
)
from photo_annex import replace_photo_annex, resolve_photos  # noqa: E402

RUNTIME = ROOT / "templates" / "ppwr_rev00_locked" / "runtime"
SIG = ROOT / "assets" / "signatory" / "numan_alver_signature_transparent.png"
DATE = "11.08.2026"

FILES = {
    "01_Technical_File.docx": RUNTIME / "01_Technical_File_RUNTIME.docx",
    "02_EU_DoC.docx": RUNTIME / "02_EU_DoC_RUNTIME.docx",
    "03_Label.docx": RUNTIME / "03_Label_RUNTIME.docx",
    "04_Shipment_Statement.docx": RUNTIME / "04_Shipment_Statement_RUNTIME.docx",
}


def _replace_tokens_in_docx(path: Path, mapping: dict[str, str]) -> None:
    """Replace {{TOKEN}} in document.xml + headers/footers keeping run formatting."""
    import zipfile
    from io import BytesIO

    buf = BytesIO()
    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    text = data.decode("utf-8")
                    for k, v in mapping.items():
                        text = text.replace("{{" + k + "}}", _xml_escape(v))
                    data = text.encode("utf-8")
                zout.writestr(item, data)
    path.write_bytes(buf.getvalue())


def _xml_escape(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _find_bom_table(doc: Document):
    for table in doc.tables:
        hdr = " ".join(c.text for c in table.rows[0].cells).upper()
        if "BİLEŞEN" in hdr or "BILESEN" in hdr or "COMPONENT" in hdr and "MATERIAL" in hdr:
            if "{{BOM" in "".join(c.text for row in table.rows for c in row.cells) or len(table.rows) >= 2:
                return table
        if "BOM_CODE" in "".join(c.text for row in table.rows for c in row.cells):
            return table
    # fallback: table with Kod + Miktar headers
    for table in doc.tables:
        hdr = " ".join(c.text.upper() for c in table.rows[0].cells)
        if "KOD" in hdr and "MIKTAR" in hdr and "BIRIM" in hdr.replace("İ", "I"):
            return table
    return None


def _find_product_table(doc: Document):
    for table in doc.tables:
        hdr = " ".join(c.text.upper() for c in table.rows[0].cells)
        if ("ÜRÜN KODU" in hdr or "URUN KODU" in hdr or "PRODUCT CODE" in hdr) and (
            "TANIM" in hdr or "DESCRIPTION" in hdr
        ):
            return table
    return None


def _fill_bom(table, bom_lines: list[dict], total_tare_kg: float | None) -> None:
    tbl = table._tbl
    rows = list(tbl.findall(qn("w:tr")))
    if len(rows) < 2:
        return
    template_tr = rows[1]
    # keep header + total if last looks like TOTAL
    last_txt = "".join(table.rows[-1].cells[0].text).upper()
    has_total = "TOPLAM" in last_txt or "TOTAL" in last_txt
    total_tr = rows[-1] if has_total else None
    # remove data rows
    for tr in rows[1:(-1 if has_total else None) or None]:
        if tr is total_tr:
            continue
        tbl.remove(tr)
    # if no total, remove all after header
    if not has_total:
        for tr in list(tbl.findall(qn("w:tr")))[1:]:
            tbl.remove(tr)
        total_tr = None

    # re-get total
    if has_total:
        total_tr = tbl.findall(qn("w:tr"))[-1]

    for line in bom_lines:
        new_tr = deepcopy(template_tr)
        if total_tr is not None:
            total_tr.addprevious(new_tr)
        else:
            tbl.append(new_tr)

    # refill via API
    data_rows = table.rows[1:-1] if has_total else table.rows[1:]
    for row, line in zip(data_rows, bom_lines):
        code = str(line.get("component_code") or "")
        desc_tr = str(line.get("description") or line.get("name_tr") or "")
        desc_en = str(line.get("name_en") or translate_component(desc_tr))
        qty = line.get("qty")
        uom = str(line.get("uom") or "ADT").upper()
        uom_tr = {"ADT": "adet", "PCS": "adet", "M": "m", "KG": "kg", "GR": "g"}.get(uom, uom.lower())
        uom_en = {"ADT": "pcs", "PCS": "pcs", "M": "m", "KG": "kg", "GR": "g"}.get(uom, uom.lower())
        unit_wt = line.get("unit_weight")
        line_wt = line.get("line_weight")
        if len(row.cells) >= 1:
            row.cells[0].text = code
        if len(row.cells) >= 2:
            set_bilingual_cell(row.cells[1], desc_tr, desc_en)
        if len(row.cells) >= 3:
            row.cells[2].text = "" if qty is None else f"{qty}".rstrip("0").rstrip(".") if isinstance(qty, float) else str(qty)
        if len(row.cells) >= 4:
            set_bilingual_cell(row.cells[3], uom_tr, uom_en, size=8)
        if len(row.cells) >= 5:
            row.cells[4].text = "—" if unit_wt in (None, "") else str(unit_wt)
        if len(row.cells) >= 6:
            row.cells[5].text = "—" if line_wt in (None, "") else str(line_wt)

    if has_total and total_tare_kg is not None and len(table.rows[-1].cells) >= 1:
        # put tare on last numeric cell
        table.rows[-1].cells[-1].text = f"{total_tare_kg:.4f} kg"


def _fill_product_table(table, product_code: str, desc_tr: str, desc_en: str) -> None:
    if len(table.rows) < 2:
        return
    row = table.rows[1]
    row.cells[0].text = product_code
    if len(row.cells) > 1:
        set_bilingual_cell(row.cells[1], desc_tr, desc_en)


def _stamp_dates(doc: Document) -> None:
    pat = re.compile(r"\b\d{2}\.\d{2}\.\d{4}\b")
    for p in doc.paragraphs:
        if pat.search(p.text or "") and ("Yayın" in (p.text or "") or "Rel. Date" in (p.text or "") or "Rev" in (p.text or "")):
            for r in p.runs:
                if r.text and pat.search(r.text):
                    r.text = pat.sub(DATE, r.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.text and pat.search(r.text):
                            r.text = pat.sub(DATE, r.text)


def _ensure_doc_signature(doc: Document) -> None:
    blob = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                blob += "\n" + c.text
    # fill signature table if present
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            left = row.cells[0].text.upper().replace("İ", "I").replace("Ö", "O")
            if ("SIGNATURE" in left or "IMZA" in left) and ("NAME" in left or "AD /" in left or "FUNCTION" in left):
                cell = row.cells[-1]
                for p in list(cell.paragraphs):
                    el = p._p
                    parent = el.getparent()
                    if parent is not None:
                        parent.remove(el)
                p1 = cell.add_paragraph()
                r = p1.add_run("Ad / Name: Numan Alver")
                r.bold = True
                r.font.name = "Tahoma"
                r.font.size = Pt(9)
                p2 = cell.add_paragraph()
                r = p2.add_run("Görev / Function: Operasyon Direktörü / Operations Director")
                r.bold = True
                r.font.name = "Tahoma"
                r.font.size = Pt(9)
                p2b = cell.add_paragraph()
                r = p2b.add_run("Function: Operations Director")
                r.italic = True
                r.font.name = "Tahoma"
                r.font.size = Pt(9)
                p3 = cell.add_paragraph()
                p3.add_run("İmza / Signature:").font.name = "Tahoma"
                if SIG.exists():
                    cell.add_paragraph().add_run().add_picture(str(SIG), width=Cm(4.2))
                return
    if "Numan Alver" not in blob:
        doc.add_paragraph("Ad / Name: Numan Alver")
        doc.add_paragraph("Görev / Function: Operasyon Direktörü / Operations Director")
        if SIG.exists():
            doc.add_paragraph().add_run().add_picture(str(SIG), width=Cm(4.2))


def build_premium_pack(
    dest: Path,
    *,
    key: str,
    description_tr: str,
    description_en: str | None = None,
    set_code: str,
    bom_lines: list[dict],
    scope: str,
    total_tare_kg: float | None = None,
    config_id: str | None = None,
    photo_prefix: str | None = None,
) -> dict:
    if dest.exists():
        shutil.rmtree(dest)
    dest.mkdir(parents=True)

    desc_en = description_en or translate_product(description_tr)
    cfg = set_code or key
    cid = config_id or f"IA-{key}"
    ids = {
        "CONFIG_SET_CODE": cfg,
        "CONFIG_ID": cid,
        "SOURCE_CONFIG_ID": cid,
        "PRODUCT_CODE": key,
        "PRODUCT_DESCRIPTION": description_tr,  # bilingual applied in table after
        "VARIANT_BASIS_TR": description_tr,
        "DOC_ID": f"IA-PPWR-DOC-{key}-R00",
        "LABEL_ID": f"IA-PPWR-LBL-{key}-R00",
        "STM_ID": f"IA-PPWR-STM-{key}-R00",
        "REVISION_PAIR": f"Rev.00 / {DATE}",
        "TOTAL_TARE_KG": f"{total_tare_kg:.4f} kg" if total_tare_kg is not None else "—",
        # BOM placeholders cleared — filled structurally
        "BOM_CODE": "",
        "BOM_DESC": "",
        "BOM_QTY": "",
        "BOM_UOM": "",
        "BOM_UNIT_WT": "",
        "BOM_LINE_WT": "",
    }

    photos_n = 0
    for fname, src in FILES.items():
        assert src.exists(), src
        out = dest / fname
        shutil.copy2(src, out)
        _replace_tokens_in_docx(out, ids)
        doc = Document(str(out))
        _stamp_dates(doc)

        if fname.startswith("01_"):
            pt = _find_product_table(doc)
            if pt is not None:
                _fill_product_table(pt, key, description_tr, desc_en)
            bt = _find_bom_table(doc)
            if bt is not None:
                _fill_bom(bt, bom_lines, total_tare_kg)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "{{BOM_" in cell.text:
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    if "{{BOM_" in (r.text or ""):
                                        r.text = ""
            doc.save(str(out))
            photos = resolve_photos(scope=scope, bom_lines=bom_lines, prefer_prefix=photo_prefix)
            photos_n = replace_photo_annex(out, photos, title_extra=f"{key}")
        elif fname.startswith("02_"):
            pt = _find_product_table(doc)
            if pt is not None:
                _fill_product_table(pt, key, description_tr, desc_en)
            _ensure_doc_signature(doc)
            doc.save(str(out))
        else:
            pt = _find_product_table(doc)
            if pt is not None:
                _fill_product_table(pt, key, description_tr, desc_en)
            doc.save(str(out))

    return {"key": key, "photos": photos_n, "bom": len(bom_lines)}
