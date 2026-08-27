"""Style DoC 'Kontrollü ürün kapsamı' as a clean navy/gold data table on all 287 DoCs."""

from __future__ import annotations

import tempfile
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\burcu\Documents\YAZILIM\Inci_Aku_PPWR_PIMS")
DOC = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL" / "01_DOCUMENT_SETS"

NAVY = "0E2A47"
GOLD = "C8A24A"
BAND = "F3F6F9"
WHITE = "FFFFFF"


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove old shd
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    # remove old borders
    for child in list(tblPr):
        if child.tag == qn("w:tblBorders"):
            tblPr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), GOLD)
        borders.append(el)
    tblPr.append(borders)


def _style_cell(cell, text: str, *, header: bool, bold: bool = False) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    # keep bilingual lines
    lines = text.split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i == 0:
            if p.runs:
                r = p.runs[0]
                r.text = line
            else:
                r = p.add_run(line)
        else:
            p.add_run("\n" + line) if False else None
            # better: new paragraph for EN line
            break
    # rewrite cleanly
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    while len(cell.paragraphs) > 1:
        el = cell.paragraphs[-1]._p
        el.getparent().remove(el)
    p = cell.paragraphs[0]
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line) if i == 0 or True else p.add_run(line)
        if i > 0 and len(p.runs) > 1:
            # last run is the new one after break — simplify
            pass
    # simpler approach: one run with \n
    for r in p.runs:
        r.text = ""
    r = p.add_run(text)
    r.bold = True if header else bold
    r.font.size = Pt(8) if header else Pt(9)
    r.font.name = "Tahoma"
    if header:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        r.font.color.rgb = RGBColor(0x1C, 0x24, 0x30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def _style_cell_simple(cell, text: str, *, header: bool, bold: bool = False) -> None:
    for p in cell.paragraphs:
        for r in list(p.runs):
            r.text = ""
    while len(cell.paragraphs) > 1:
        el = cell.paragraphs[-1]._p
        el.getparent().remove(el)
    p = cell.paragraphs[0]
    parts = (text or "").split("\n")
    first = True
    for part in parts:
        if not first:
            p.add_run().add_break()
        r = p.add_run(part)
        r.bold = True if header else bold
        r.font.size = Pt(8) if header else Pt(9)
        r.font.name = "Tahoma"
        if header:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        else:
            r.font.color.rgb = RGBColor(0x1C, 0x24, 0x30)
        first = False
    # clear empty leading runs left from initial clear
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    _shade(cell, NAVY if header else BAND)


def style_doc(path: Path) -> bool:
    doc = Document(str(path))
    target = None
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = " ".join(c.text.upper() for c in table.rows[0].cells)
        if ("ÜRÜN KODU" in hdr or "URUN KODU" in hdr or "PRODUCT CODE" in hdr) and (
            "TANIM" in hdr or "DESCRIPTION" in hdr
        ):
            # Prefer the one near "Kontrollü ürün" — if multiple, take last matching (DoC insert)
            target = table
    if target is None:
        return False

    _borders(target)
    # set widths ~ 4cm / 12cm
    for row in target.rows:
        if len(row.cells) >= 2:
            row.cells[0].width = Cm(4.2)
            row.cells[1].width = Cm(12.0)

    for ri, row in enumerate(target.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text
            if ri == 0:
                if ci == 0:
                    text = "ÜRÜN KODU\nPRODUCT CODE"
                elif ci == 1:
                    text = "ÜRÜN TANIMI\nPRODUCT DESCRIPTION"
                _style_cell_simple(cell, text, header=True)
            else:
                _style_cell_simple(cell, text, header=False, bold=(ci == 0))
                if ri % 2 == 0:
                    _shade(cell, "FFFFFF")
                else:
                    _shade(cell, BAND)

    # ensure heading paragraph exists and is styled
    for p in doc.paragraphs:
        if "Kontrollü ürün kapsamı" in (p.text or "") or "Controlled product scope" in (p.text or ""):
            for r in p.runs:
                r.bold = True
                r.font.name = "Tahoma"
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x0E, 0x2A, 0x47)
            break

    doc.save(str(path))
    return True


def main() -> None:
    docs = sorted(DOC.rglob("02_EU_DoC.docx"))
    print(f"Styling {len(docs)} DoCs…", flush=True)
    n = 0
    for i, p in enumerate(docs, 1):
        if style_doc(p):
            n += 1
        if i % 50 == 0:
            print(f"  {i}/{len(docs)}", flush=True)
    print(f"styled {n}/{len(docs)}", flush=True)

    # sample verify
    s = DOC / "ST-012-EUR-01" / "02_EU_DoC.docx"
    d = Document(str(s))
    for t in d.tables:
        hdr = " ".join(c.text for c in t.rows[0].cells)
        if "ÜRÜN KODU" in hdr or "PRODUCT CODE" in hdr:
            print("sample header:", [c.text for c in t.rows[0].cells])
            print("sample row:", [c.text for c in t.rows[1].cells] if len(t.rows) > 1 else None)
            break


if __name__ == "__main__":
    main()
