"""Bilingual DOCX helpers: Turkish normal + English italic."""

from __future__ import annotations

import re
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT = "Tahoma"
INK = RGBColor(0x1C, 0x24, 0x30)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Component / packaging phrase translations (longest keys first when matching)
_COMPONENT_MAP = {
    "KARTON SEPERATÖR": "Cardboard Separator",
    "KARTON SEPARATÖR": "Cardboard Separator",
    "KARTON SEPERATOR": "Cardboard Separator",
    "KARTON KÖŞEBENT": "Cardboard Edge Protector",
    "KARTON KOSEBENT": "Cardboard Edge Protector",
    "PLASTİK SARI KÖŞE": "Yellow Plastic Corner Protector",
    "PLASTIK SARI KOSE": "Yellow Plastic Corner Protector",
    "HAVA YASTIĞI": "Dunnage Air Bag",
    "HAVA YASTIGI": "Dunnage Air Bag",
    "BAĞLAMA HALATI": "Lashing Rope",
    "BAGLAMA HALATI": "Lashing Rope",
    "PALET SABİTLEME TAKOZU": "Pallet Securing Block",
    "PALET SABITLEME TAKOZU": "Pallet Securing Block",
    "ÜÇGEN TAKOZ": "Triangular Wedge",
    "UCGEN TAKOZ": "Triangular Wedge",
    "HALAT TOKASI": "Rope Buckle",
    "YANICI ETİKET": "Flammable Label",
    "YANICI ETIKET": "Flammable Label",
    "ETİKET CORROSIVE": "Corrosive Label",
    "ETIKET CORROSIVE": "Corrosive Label",
    "TÜM ETİKETLER": "All Labels",
    "TUM ETIKETLER": "All Labels",
    "END HÜCRE SEVK TAPASI": "Cell Shipping Vent Plug",
    "END HÜCRE SEVK TAPASI (SUR KALIP)": "Cell Shipping Vent Plug (SUR Mould)",
    "END HUCRE SEVK TAPASI": "Cell Shipping Vent Plug",
    "M10 KUTUPBAŞI TAPA SARI": "M10 Terminal Cap Yellow",
    "M10 KUTUPBASI TAPA SARI": "M10 Terminal Cap Yellow",
    "END OTO DOLUM SAF SU BİDONU 20L": "Auto-fill Demineralized Water Jerry Can 20L",
    "END OTO DOLUM SAF SU BIDONU 20L": "Auto-fill Demineralized Water Jerry Can 20L",
    "END OTO DOLUM SAF SU BİDONU 25L": "Auto-fill Demineralized Water Jerry Can 25L",
    "END OTO DOLUM SAF SU BIDONU 25L": "Auto-fill Demineralized Water Jerry Can 25L",
    "STRECH PED": "Stretch Pad",
    "STRECH": "Stretch Film",
    "STREÇ": "Stretch Film",
    "STREC": "Stretch Film",
    "SIGNODE ÇEMBER": "Signode Strapping",
    "SIGNODE CEMBER": "Signode Strapping",
    "PALET EURO": "Euro Pallet",
    "PALET STANDART": "Standard Pallet",
    "PALET -1": "Pallet",
    "PALET": "Pallet",
    "SHRINK": "Shrink Film",
    "STRAFOR": "EPS Foam",
    "KARTON KOLI": "Carton Box",
    "KARTON KUTU": "Carton Box",
    "POŞET": "Plastic Bag",
    "POSET": "Plastic Bag",
    "BANT": "Tape",
    "ETİKET": "Label",
    "ETIKET": "Label",
    "OSB": "OSB Board",
    "KÖPÜK": "Foam Protection",
    "KOPUK": "Foam Protection",
    "ADET": "pcs",
    "TOPLAM": "TOTAL",
}

_PRODUCT_MAP = {
    "AKÜ": "BATTERY",
    "AKU": "BATTERY",
}


def _norm(s: str) -> str:
    return (
        (s or "")
        .upper()
        .replace("İ", "I")
        .replace("Ş", "S")
        .replace("Ğ", "G")
        .replace("Ü", "U")
        .replace("Ö", "O")
        .replace("Ç", "C")
    )


def translate_component(tr: str) -> str:
    raw = (tr or "").strip()
    if not raw:
        return ""
    # already bilingual?
    if "\n" in raw:
        parts = raw.split("\n", 1)
        return parts[1].strip() or parts[0].strip()
    n = _norm(raw)
    # exact / startswith family match (longest first)
    keys = sorted(_COMPONENT_MAP.keys(), key=lambda k: len(k), reverse=True)
    for k in keys:
        if _norm(k) in n or n.startswith(_norm(k)):
            # keep technical suffixes (sizes) from original
            en = _COMPONENT_MAP[k]
            # append dimension-like tail if present
            m = re.search(r"(\d+(?:[.,]\d+)?\s*[xX×\*]\s*\d+(?:[.,]\d+)?(?:\s*[xX×\*]\s*\d+(?:[.,]\d+)?)*)", raw)
            if m and m.group(1) not in en:
                return f"{en} {m.group(1)}"
            return en
    # OSB special
    if "OSB" in n:
        return re.sub(r"(?i)osb", "OSB Board", raw)
    # fallback: keep technical string as EN (acceptable for codes/specs)
    return raw


def translate_product(tr: str) -> str:
    raw = (tr or "").strip()
    if not raw:
        return ""
    if "\n" in raw:
        return raw.split("\n", 1)[-1].strip()
    en = raw
    for a, b in _PRODUCT_MAP.items():
        en = re.sub(rf"(?i)\b{re.escape(a)}\b", b, en)
    return en


def _font(run, *, size=9, bold=False, italic=False, color=INK) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    except Exception:
        pass


def clear_cell(cell) -> None:
    # keep one empty paragraph
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    while len(cell.paragraphs) > 1:
        el = cell.paragraphs[-1]._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def set_bilingual_cell(
    cell,
    tr: str,
    en: str | None = None,
    *,
    header: bool = False,
    bold: bool = False,
    size: int = 9,
) -> None:
    """Write TR (normal) + EN (italic) into a table cell."""
    tr = (tr or "").strip()
    if en is None:
        en = translate_component(tr) if tr else ""
    en = (en or "").strip()
    if en == tr:
        # still show italic EN line for consistency when identical technical text
        pass
    clear_cell(cell)
    p = cell.paragraphs[0]
    color = WHITE if header else INK
    r1 = p.add_run(tr)
    _font(r1, size=size, bold=bold or header, italic=False, color=color)
    if en:
        p.add_run().add_break()
        r2 = p.add_run(en)
        _font(r2, size=size, bold=False, italic=True, color=color)


def set_bilingual_from_text(cell, text: str, *, header: bool = False, size: int = 9) -> None:
    text = text or ""
    if "\n" in text:
        tr, en = text.split("\n", 1)
        set_bilingual_cell(cell, tr.strip(), en.strip(), header=header, size=size)
    else:
        # try translate
        set_bilingual_cell(cell, text.strip(), translate_component(text), header=header, size=size)
