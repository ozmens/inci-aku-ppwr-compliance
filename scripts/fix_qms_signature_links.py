"""Stamp QMS type numbers YS/D/0020-0023 on all Word docs, ensure Numan Alver
signature on DoCs, fix Excel links (SEARCH + Document Center), rebuild PDFs/ZIP.

Does NOT change BOM/mappings/set codes/product data.
"""

from __future__ import annotations

import hashlib
import os
import shutil
import sys
import zipfile
from pathlib import Path

ROOT = Path(r"C:\Users\burcu\Documents\YAZILIM\Inci_Aku_PPWR_PIMS")
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(ROOT / "tools"))

from docx import Document
from docx.shared import Cm
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from builders.phase_i.render_batch import render_docx_batch

FINAL = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL"
DOC_SETS = FINAL / "01_DOCUMENT_SETS"
ENG_CTRL = FINAL / "00_CONTROL" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
ENG_ROOT = ROOT / "output" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
SIG_T = ROOT / "assets" / "signatory" / "numan_alver_signature_transparent.png"
SIG_W = ROOT / "assets" / "signatory" / "numan_alver_signature_whitebg.png"
MARKER = "[[SIGNATORY_SIGNATURE_IMAGE]]"
ZIP_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL.zip"
SHA_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL_SHA256.txt"
BACKUP_ENG = ROOT / "output" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00_PRE_QMS_LINK_FIX.xlsx"

QMS = {
    "01_Technical_File.docx": ("YS/D/0020", "Technical File", "Teknik Dosya"),
    "02_EU_DoC.docx": ("YS/D/0021", "EU DoC", "AB Uygunluk Beyanı"),
    "03_Label.docx": ("YS/D/0022", "Label", "Etiket"),
    "04_Shipment_Statement.docx": ("YS/D/0023", "Shipment Statement", "Sevkiyat Beyanı"),
}

NAVY, BLUE, WHITE, GOLD = "0E2A47", "1F4E79", "FFFFFF", "C8A24A"
FONT = "Tahoma"
LINK_FONT = Font(name=FONT, size=9, bold=True, color="0563C1", underline="single")
GOLD_B = Border(
    left=Side(style="thin", color=GOLD),
    right=Side(style="thin", color=GOLD),
    top=Side(style="thin", color=GOLD),
    bottom=Side(style="thin", color=GOLD),
)


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def win_rel(sc: str, fname: str) -> str:
    return f"..\\01_DOCUMENT_SETS\\{sc}\\{fname}"


def stamp_qms_and_signature() -> dict:
    sig = SIG_T if SIG_T.exists() else SIG_W
    assert sig.exists(), "signature asset missing"
    stats = {"word": 0, "doc_signed": 0, "marker_cleared": 0, "titles": 0}

    for set_dir in sorted(p for p in DOC_SETS.iterdir() if p.is_dir()):
        for fname, (qms, en, tr) in QMS.items():
            path = set_dir / fname
            if not path.exists():
                continue
            doc = Document(str(path))
            # Core title → appears as YS_D_0020 - Technical File in Word temp names
            title = f"{qms} - {en}"
            doc.core_properties.title = title
            doc.core_properties.subject = f"{qms} | {tr} / {en}"
            doc.core_properties.keywords = f"{qms};PPWR;Inci Aku;{en}"
            stats["titles"] += 1

            # Visible stamp in footer (first section) if not already present
            for section in doc.sections:
                footer = section.footer
                blob = "\n".join(p.text for p in footer.paragraphs)
                if qms not in blob:
                    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                    # append type no line
                    if p.runs:
                        # add new paragraph to avoid wrecking existing footer
                        p2 = footer.add_paragraph()
                        run = p2.add_run(f"Doküman Tip No / Document Type No: {qms}")
                    else:
                        run = p.add_run(f"Doküman Tip No / Document Type No: {qms}")
                    run.font.name = "Tahoma"
                    run.font.size = __import__("docx").shared.Pt(8)
                    run.bold = True

            if fname == "02_EU_DoC.docx":
                blob = "\n".join(p.text for p in doc.paragraphs)
                for t in doc.tables:
                    for row in t.rows:
                        for c in row.cells:
                            blob += "\n" + c.text
                for p in doc.paragraphs:
                    if MARKER in p.text or "SIGNATORY_SIGNATURE" in p.text:
                        for r in p.runs:
                            r.text = ""
                        stats["marker_cleared"] += 1
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if MARKER in p.text or "SIGNATORY_SIGNATURE" in p.text:
                                    for r in p.runs:
                                        r.text = ""
                                    stats["marker_cleared"] += 1

                with zipfile.ZipFile(path) as z:
                    media_before = len([n for n in z.namelist() if n.startswith("word/media/")])

                has_numan = "Numan Alver" in blob
                has_ops = "Operations Director" in blob
                # Only embed signature image if clearly missing (media typically logo+icons+sig >=4)
                if media_before < 4:
                    placed = False
                    for t in doc.tables:
                        for row in t.rows:
                            for cell in row.cells:
                                if "Numan Alver" in cell.text or "İmza" in cell.text or "Signature" in cell.text:
                                    p_img = cell.add_paragraph()
                                    p_img.add_run().add_picture(str(sig), width=Cm(4.2))
                                    placed = True
                                    break
                            if placed:
                                break
                        if placed:
                            break
                    if not placed:
                        doc.add_paragraph().add_run().add_picture(str(sig), width=Cm(4.2))
                    stats["doc_signed"] += 1
                else:
                    stats["doc_signed"] += 1

                if not has_numan:
                    doc.add_paragraph("Ad / Name: Numan Alver")
                if not has_ops:
                    doc.add_paragraph("Görev / Function: Operasyon Direktörü / Operations Director")

            doc.save(str(path))
            stats["word"] += 1
        if stats["word"] % 200 == 0:
            print(f"stamped {stats['word']}…", flush=True)
    return stats


def render_all_pdfs() -> int:
    jobs = []
    for docx in sorted(DOC_SETS.rglob("*.docx")):
        if docx.name.startswith("~$"):
            continue
        jobs.append((docx, docx.with_suffix(".pdf")))
    ok = 0
    CHUNK = 40
    for i in range(0, len(jobs), CHUNK):
        chunk = jobs[i : i + CHUNK]
        results = render_docx_batch(chunk, progress_every=20, log=[])
        ok += sum(1 for r in results if r.get("render_ok"))
        print(f"PDF {min(i+CHUNK,len(jobs))}/{len(jobs)} ok={ok}", flush=True)
    # retry
    retry = [(d, p) for d, p in jobs if not (p.exists() and p.stat().st_size > 0)]
    if retry:
        for i in range(0, len(retry), 20):
            render_docx_batch(retry[i : i + 20], progress_every=10, log=[])
    return sum(1 for _d, p in jobs if p.exists() and p.stat().st_size > 0)


def fix_excel_links() -> None:
    """Make SEARCH + Document Center links actually clickable from 00_CONTROL.

    SEARCH: replace IF(HYPERLINK()) with real Hyperlink objects updated via
    formulas that Excel treats as clickable — use HYPERLINK as OUTERMOST and
    CELL(\"filename\") so path resolves from workbook location.
    Also rewrite Document Center hyperlinks via COM with backslash relatives.
    """
    shutil.copy2(ENG_CTRL, BACKUP_ENG)
    wb = load_workbook(ENG_CTRL)

    # HOME QMS legend
    if "00_HOME" in wb.sheetnames:
        home = wb["00_HOME"]
        home["B27"] = (
            "QMS Document Type Nos:  "
            "Technical File YS/D/0020  •  EU DoC YS/D/0021  •  "
            "Label YS/D/0022  •  Shipment Statement YS/D/0023"
        )
        home["B27"].font = Font(name=FONT, size=9, bold=True, color=NAVY)
        home.merge_cells("B27:I27")
        home["B28"] = (
            "IMPORTANT: Open this Document Engine ONLY from the 00_CONTROL folder "
            "inside the customer delivery package so document links resolve correctly."
        )
        home["B28"].font = Font(name=FONT, size=9, bold=True, color="A12622")
        home.merge_cells("B28:I28")

    # SEARCH — outermost HYPERLINK + CELL filename base
    ws = wb["SEARCH"]
    # workbook folder helper (requires calc)
    # Z2 = folder path of this workbook
    ws["Z2"] = '=IFERROR(LEFT(CELL("filename",$B$4),FIND("[",CELL("filename",$B$4))-1),"")'
    ws.column_dimensions["Z"].hidden = True

    issued = (
        'AND($B$8<>"",$A$8<>"NOT FOUND",'
        'ISNUMBER(SEARCH("ISSUED",$G$8)),'
        'NOT(ISNUMBER(SEARCH("NOT ISSUED",$G$8))),'
        'NOT(ISNUMBER(SEARCH("YURT",$G$8))))'
    )
    domestic = 'OR(ISNUMBER(SEARCH("YURT",$G$8)),ISNUMBER(SEARCH("NOT ISSUED",$G$8)))'

    def action_formula(stem: str, is_word: bool) -> str:
        label = "OPEN WORD" if is_word else "OPEN PDF"
        ext = "docx" if is_word else "pdf"
        # Outer HYPERLINK — clickable; address empty when not issued
        # Path: workbook_dir & ..\01_DOCUMENT_SETS\SET\file
        path_expr = f'$Z$2&"..\\01_DOCUMENT_SETS\\"&$B$8&"\\{stem}.{ext}"'
        return (
            f'=IF(OR($B$4="",$A$8="",$A$8="NOT FOUND"),"",'
            f'IF({domestic},"DOCUMENTS NOT ISSUED",'
            f'IF({issued},HYPERLINK({path_expr},"{label}"),"")))'
        )

    # rows: TF 13, DoC 15, Label 17, STM 19
    mapping = [
        (13, "01_Technical_File"),
        (15, "02_EU_DoC"),
        (17, "03_Label"),
        (19, "04_Shipment_Statement"),
    ]
    for r, stem in mapping:
        ws.cell(r, 1).value = action_formula(stem, True)
        ws.cell(r, 2).value = action_formula(stem, False)
        for c in (1, 2):
            cell = ws.cell(r, c)
            cell.font = Font(name=FONT, size=10, bold=True, color="0563C1")
            cell.fill = PatternFill("solid", fgColor="E8F0FE")
            cell.border = GOLD_B
            cell.alignment = Alignment(horizontal="center")

    # Document Center — ensure friendly labels (COM will set real hyperlinks)
    if "DOCUMENT_CENTER" in wb.sheetnames:
        dc = wb["DOCUMENT_CENTER"]
        # reinforce group titles with QMS nos
        dc["E3"] = "TECHNICAL FILE (YS/D/0020)"
        dc["H3"] = "EU DECLARATION OF CONFORMITY (YS/D/0021)"
        dc["K3"] = "LABEL (YS/D/0022)"
        dc["N3"] = "SHIPMENT STATEMENT (YS/D/0023)"
        for col in (5, 8, 11, 14):
            dc.cell(3, col).font = Font(name=FONT, size=9, bold=True, color=WHITE)
            dc.cell(3, col).fill = PatternFill("solid", fgColor=NAVY if col in (5, 11) else BLUE)

    # Save to temp then replace (avoids lock on open Excel)
    tmp = ENG_CTRL.with_suffix(".xlsx.tmp")
    wb.save(tmp)
    wb.close()
    try:
        shutil.move(str(tmp), str(ENG_CTRL))
    except PermissionError:
        alt = ENG_CTRL.parent / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00_QMS_LINK_FIXED.xlsx"
        shutil.move(str(tmp), str(alt))
        print(f"CTRL engine locked — saved as {alt.name}. Close Excel and replace manually.", flush=True)
        # still COM-polish the alt file and copy to root
        eng_for_com = alt
    else:
        eng_for_com = ENG_CTRL

    # COM: rewrite DC hyperlinks with Windows relative paths; HOME buttons
    import pythoncom
    import win32com.client as win32

    pythoncom.CoInitialize()
    excel = None
    try:
        excel = win32.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        wb = excel.Workbooks.Open(str(ENG_CTRL.resolve()))
        dc = wb.Worksheets("DOCUMENT_CENTER")
        r = 5
        while dc.Cells(r, 2).Value:
            sc = str(dc.Cells(r, 2).Value).strip()
            mapping = {
                6: win_rel(sc, "01_Technical_File.docx"),
                7: win_rel(sc, "01_Technical_File.pdf"),
                9: win_rel(sc, "02_EU_DoC.docx"),
                10: win_rel(sc, "02_EU_DoC.pdf"),
                12: win_rel(sc, "03_Label.docx"),
                13: win_rel(sc, "03_Label.pdf"),
                15: win_rel(sc, "04_Shipment_Statement.docx"),
                16: win_rel(sc, "04_Shipment_Statement.pdf"),
            }
            for col, addr in mapping.items():
                cell = dc.Cells(r, col)
                try:
                    if cell.Hyperlinks.Count >= 1:
                        cell.Hyperlinks.Delete()
                except Exception:
                    pass
                label = "OPEN WORD" if addr.endswith(".docx") else "OPEN PDF"
                cell.Value = label
                dc.Hyperlinks.Add(Anchor=cell, Address=addr, TextToDisplay=label)
                cell.Font.Color = 0xC16305
                cell.Font.Underline = 2
            r += 1
            if r > 400:
                break

        # HOME buttons far right
        navy = 0x0E + 256 * 0x2A + 65536 * 0x47
        gold = 0xC8 + 256 * 0xA2 + 65536 * 0x4A
        for i in range(1, wb.Worksheets.Count + 1):
            ws = wb.Worksheets(i)
            try:
                for s in list(ws.Shapes):
                    if str(getattr(s, "Name", "")).startswith("HOME_NAV"):
                        s.Delete()
            except Exception:
                pass
            if ws.Name == "00_HOME":
                continue
            try:
                left = float(ws.Cells(1, 18).Left) if ws.Name == "DOCUMENT_CENTER" else float(ws.Cells(1, 12).Left)
                shp = ws.Shapes.AddShape(5, left, 2, 90, 18)
                shp.Name = "HOME_NAV"
                shp.Fill.ForeColor.RGB = navy
                shp.Line.ForeColor.RGB = gold
                shp.TextFrame.Characters().Text = "⌂ HOME"
                shp.TextFrame.Characters().Font.Color = 0xFFFFFF
                shp.TextFrame.Characters().Font.Size = 9
                shp.TextFrame.Characters().Font.Bold = True
                shp.TextFrame.HorizontalAlignment = 2
                ws.Hyperlinks.Add(Anchor=shp, Address="", SubAddress="'00_HOME'!A1")
            except Exception:
                pass

        wb.Save()
        wb.Close(False)
    finally:
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    shutil.copy2(ENG_CTRL, ENG_ROOT)


def verify() -> dict:
    # sample docs
    sample = DOC_SETS / "ST-012-EUR-01"
    out = {}
    for fname, (qms, en, tr) in QMS.items():
        d = Document(str(sample / fname))
        out[fname] = {
            "title": d.core_properties.title,
            "qms_in_title": qms in (d.core_properties.title or ""),
        }
        if fname.startswith("02"):
            blob = "\n".join(p.text for p in d.paragraphs)
            for t in d.tables:
                for row in t.rows:
                    for c in row.cells:
                        blob += "\n" + c.text
            out[fname]["numan"] = "Numan Alver" in blob
            with zipfile.ZipFile(sample / fname) as z:
                out[fname]["media"] = len([n for n in z.namelist() if n.startswith("word/media/")])
            out[fname]["marker"] = MARKER in blob or "SIGNATORY_SIGNATURE_IMAGE" in blob

    # link follow from CTRL
    import pythoncom
    import win32com.client as win32

    pythoncom.CoInitialize()
    excel = None
    try:
        excel = win32.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        wb = excel.Workbooks.Open(str(ENG_CTRL.resolve()))
        dc = wb.Worksheets("DOCUMENT_CENTER")
        addr = dc.Cells(5, 6).Hyperlinks(1).Address
        try:
            wb.FollowHyperlink(Address=addr)
            out["dc_follow"] = True
        except Exception as e:
            out["dc_follow"] = False
            out["dc_follow_err"] = str(e)
        # SEARCH
        s = wb.Worksheets("SEARCH")
        s.Range("B4").NumberFormat = "@"
        s.Range("B4").Value = "1015169"
        excel.CalculateFull()
        out["search_a13"] = str(s.Range("A13").Value or "")
        out["search_b8"] = str(s.Range("B8").Value or "")
        wb.Close(False)
    finally:
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()
    return out


def rebuild_zip() -> str:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    print("Creating ZIP…", flush=True)
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in FINAL.rglob("*"):
            if p.is_file() and not p.name.startswith("~$"):
                zf.write(p, p.relative_to(FINAL.parent).as_posix())
    digest = sha256_file(ZIP_PATH)
    SHA_PATH.write_text(digest + "\n", encoding="utf-8")
    return digest


def main():
    print("1) Stamp QMS titles + Numan Alver signature…", flush=True)
    st = stamp_qms_and_signature()
    print(st, flush=True)

    print("2) Re-render all PDFs…", flush=True)
    pdf_ok = render_all_pdfs()
    print({"pdf_ok": pdf_ok}, flush=True)

    print("3) Fix Excel links + QMS legend…", flush=True)
    fix_excel_links()

    print("4) Verify…", flush=True)
    v = verify()
    print(v, flush=True)

    print("5) Rebuild ZIP…", flush=True)
    digest = rebuild_zip()

    words = len([p for p in DOC_SETS.rglob("*.docx") if not p.name.startswith("~$")])
    pdfs = len([p for p in DOC_SETS.rglob("*.pdf") if p.stat().st_size > 0])
    print(
        {
            "qms": "YS/D/0020..0023 stamped",
            "word": words,
            "pdf": pdfs,
            "zip_sha": digest,
            "verify": v,
        },
        flush=True,
    )


if __name__ == "__main__":
    main()
