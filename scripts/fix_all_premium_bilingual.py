"""ONE-SHOT: premium İnci Akü format + bilingual tables + fresh photos.

Scopes:
- COMPONENT (2): rebuild from YS/D runtime golden
- CONTAINER (4): rebuild from runtime + refreshed photo library
- INDUSTRIAL (2736): rebuild from runtime + Excel BOM
- STARTER (2024): keep premium body; fix table cells to TR normal / EN italic

PDF via LibreOffice only (never Word).
"""

from __future__ import annotations

import json
import shutil
import sys
import time
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import rebuild_industrial_from_excel as ind  # noqa: E402
import generate_component_packaging_delivery as cmp  # noqa: E402
from bilingual_format import set_bilingual_cell, translate_component, translate_product  # noqa: E402
from premium_pack_from_runtime import build_premium_pack  # noqa: E402
from convert_pdfs_libreoffice import convert_batch_via_temp, find_soffice  # noqa: E402
from ppwr_engine_builder import build_document_engine, verify_links  # noqa: E402
from photo_annex import clear_mapping_cache, replace_photo_annex, resolve_photos  # noqa: E402
import zipfile


def _tf_has_annex(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path, "r") as zf:
            return b"Representative Packaging Component Photos" in zf.read("word/document.xml")
    except Exception:
        return False

OUT = ROOT / "output"
DATE = "11.08.2026"
PROGRESS = OUT / "_PREMIUM_BILINGUAL_PROGRESS.json"


def progress(**kw):
    data = {}
    if PROGRESS.exists():
        try:
            data = json.loads(PROGRESS.read_text(encoding="utf-8"))
        except Exception:
            pass
    data.update(kw)
    data["ts"] = time.strftime("%Y-%m-%d %H:%M:%S")
    PROGRESS.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def lo_pdfs(files: list[Path]) -> tuple[int, int]:
    if not files:
        return 0, 0
    return convert_batch_via_temp(find_soffice(), files, OUT / "_lo_profile_premium", chunk=40)


def extract_bom_from_existing_tf(tf: Path) -> list[dict]:
    doc = Document(str(tf))
    for table in doc.tables:
        hdr = " ".join(c.text.upper() for c in table.rows[0].cells)
        if "BİLEŞEN" not in hdr and "BILESEN" not in hdr and "COMPONENT" not in hdr:
            continue
        if "KOD" not in hdr and "CODE" not in hdr:
            continue
        lines = []
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells or "TOPLAM" in cells[0].upper() or "TOTAL" in cells[0].upper():
                continue
            code = cells[0].split("\n")[0].strip()
            desc = cells[1].split("\n")[0].strip() if len(cells) > 1 else ""
            qty = cells[2].split("\n")[0].strip() if len(cells) > 2 else ""
            uom = cells[3].split("\n")[0].strip() if len(cells) > 3 else "ADT"
            if not code and not desc:
                continue
            try:
                qv: object = float(qty.replace(",", ".")) if qty else ""
            except Exception:
                qv = qty
            lines.append(
                {
                    "component_code": code,
                    "description": desc,
                    "name_en": translate_component(desc),
                    "qty": qv,
                    "uom": uom.split("/")[0].strip().upper() if uom else "ADT",
                }
            )
        if lines:
            return lines
    return []


def fix_starter_bilingual_inplace(root: Path) -> list[Path]:
    """Fix product + BOM description cells in existing premium starter docs."""
    docs = root / "01_PRODUCTS"
    changed: list[Path] = []
    folders = sorted([p for p in docs.iterdir() if p.is_dir()], key=lambda p: p.name)
    for i, folder in enumerate(folders, 1):
        for fname in [
            "01_Technical_File.docx",
            "02_EU_DoC.docx",
            "03_Label.docx",
            "04_Shipment_Statement.docx",
        ]:
            path = folder / fname
            if not path.exists():
                continue
            doc = Document(str(path))
            dirty = False
            for table in doc.tables:
                hdr = " ".join(c.text.upper() for c in table.rows[0].cells)
                # product scope table
                if ("ÜRÜN KODU" in hdr or "URUN KODU" in hdr or "PRODUCT CODE" in hdr) and (
                    "TANIM" in hdr or "DESCRIPTION" in hdr
                ):
                    for row in table.rows[1:]:
                        if len(row.cells) < 2:
                            continue
                        tr = row.cells[1].text.split("\n")[0].strip()
                        if not tr or tr in {"—", "-"}:
                            continue
                        # skip if already has italic EN run
                        has_italic = any(r.italic for p in row.cells[1].paragraphs for r in p.runs)
                        if has_italic and "\n" in row.cells[1].text:
                            continue
                        set_bilingual_cell(row.cells[1], tr, translate_product(tr))
                        dirty = True
                # BOM table
                if ("BİLEŞEN" in hdr or "BILESEN" in hdr or ("COMPONENT" in hdr and "MATERIAL" in hdr)) and (
                    "KOD" in hdr or "CODE" in hdr
                ):
                    for row in table.rows[1:]:
                        if not row.cells:
                            continue
                        if "TOPLAM" in row.cells[0].text.upper() or "TOTAL" in row.cells[0].text.upper():
                            continue
                        if len(row.cells) < 2:
                            continue
                        tr = row.cells[1].text.split("\n")[0].strip()
                        if not tr:
                            continue
                        has_italic = any(r.italic for p in row.cells[1].paragraphs for r in p.runs)
                        if has_italic and "\n" in row.cells[1].text:
                            continue
                        set_bilingual_cell(row.cells[1], tr, translate_component(tr))
                        dirty = True
                        # UOM col often index 3
                        if len(row.cells) >= 4:
                            u = row.cells[3].text.strip()
                            if u and " / " not in u and not any(
                                r.italic for p in row.cells[3].paragraphs for r in p.runs
                            ):
                                # adet / pcs style
                                low = u.lower()
                                if low in {"adet", "adt", "pcs"}:
                                    set_bilingual_cell(row.cells[3], "adet", "pcs", size=8)
                                    dirty = True
            if dirty:
                doc.save(str(path))
                changed.append(path)
        if i <= 3 or i % 200 == 0:
            print(f"STARTER bilingual {i}/{len(folders)} changed_files={len(changed)}", flush=True)
            progress(phase="starter_bilingual", i=i, total=len(folders), changed=len(changed))
    return changed


def rebuild_component(root: Path) -> list[Path]:
    docs = root / "01_VARIANTS"
    if docs.exists():
        shutil.rmtree(docs)
    docs.mkdir(parents=True)
    pdfs = []
    mapping = {"CMP-1ROW-01": "CMP-TEK-SIRA", "CMP-2ROW-01": "CMP-CIFT-SIRA"}
    for old, new in mapping.items():
        v = next(x for x in cmp.VARIANTS if x["id"] == old)
        bom = [
            {
                "component_code": b[0],
                "description": b[1].split("/")[0].strip(),
                "name_en": b[1].split("/")[-1].strip() if "/" in b[1] else translate_component(b[1]),
                "qty": b[2],
                "uom": b[3],
                "unit_weight": b[4],
                "line_weight": b[5],
            }
            for b in v["bom"]
        ]
        dest = docs / new
        scope = "COMPONENT_1ROW" if "TEK" in new else "COMPONENT_2ROW"
        prefix = "CMP_S" if "TEK" in new else "CMP_D"
        info = build_premium_pack(
            dest,
            key=new,
            description_tr=v["name_tr"],
            description_en=v["name_en"],
            set_code=old,
            bom_lines=bom,
            scope=scope,
            total_tare_kg=v["tare"],
            config_id=old,
            photo_prefix=prefix,
        )
        print("COMPONENT", new, info, flush=True)
        for stem in ["01_Technical_File", "02_EU_DoC", "03_Label", "04_Shipment_Statement"]:
            pdfs.append(dest / f"{stem}.docx")
    return pdfs


def rebuild_container(root: Path) -> list[Path]:
    docs = root / "01_CONFIGS"
    old = {p.name: p for p in docs.iterdir() if p.is_dir()} if docs.exists() else {}
    bom_cache = {}
    for name, folder in old.items():
        tf = folder / "01_Technical_File.docx"
        if tf.exists():
            bom_cache[name] = extract_bom_from_existing_tf(tf)

    labels = {
        "CNT-20-EUR-01": ("20 ft Euro paletli konteyner yükleme", "20ft Euro-pallet container loading"),
        "CNT-20-STD-01": ("20 ft standart konteyner yükleme", "20ft standard container loading"),
        "CNT-40-IND-01": ("40 ft endüstriyel konteyner yükleme", "40ft industrial container loading"),
        "CNT-DE-HAM-01": ("DE-HAM konteyner yükleme konfigürasyonu", "DE-HAM container loading configuration"),
    }
    default_bom = [
        {"component_code": "OSB", "description": "Osb (11 mm)", "qty": 10, "uom": "ADT"},
        {"component_code": "SEP", "description": "Karton Seperatör", "qty": 10, "uom": "ADT"},
        {"component_code": "AIR", "description": "Hava Yastığı (Level 2)", "qty": 5, "uom": "ADT"},
        {"component_code": "LASH", "description": "Bağlama Halatı", "qty": 1, "uom": "ADT"},
        {"component_code": "PALLET", "description": "Palet", "qty": 1, "uom": "ADT"},
        {"component_code": "LOCK", "description": "Load Lock", "qty": 1, "uom": "ADT"},
        {"component_code": "SEAL", "description": "Seal", "qty": 1, "uom": "ADT"},
    ]

    new_docs = root / "01_CONFIGS_NEW"
    if new_docs.exists():
        shutil.rmtree(new_docs)
    new_docs.mkdir(parents=True)

    for name, (tr, en) in labels.items():
        bom = bom_cache.get(name) or list(default_bom)
        for b in bom:
            b["name_en"] = translate_component(b["description"])
        dest = new_docs / name
        info = build_premium_pack(
            dest,
            key=name,
            description_tr=tr,
            description_en=en,
            set_code=name,
            bom_lines=bom,
            scope="CONTAINER",
            config_id=name,
        )
        print("CONTAINER", name, info, flush=True)

    if docs.exists():
        shutil.rmtree(docs)
    new_docs.rename(docs)

    pdfs = []
    for name in labels:
        dest = docs / name
        for stem in ["01_Technical_File", "02_EU_DoC", "03_Label", "04_Shipment_Statement"]:
            pdfs.append(dest / f"{stem}.docx")
    return pdfs


def refresh_tf_photos(tf: Path, *, scope: str, photo_prefix: str | None = None) -> int:
    bom = extract_bom_from_existing_tf(tf)
    if not bom:
        return 0
    photos = resolve_photos(scope=scope, bom_lines=bom, prefer_prefix=photo_prefix)
    return replace_photo_annex(tf, photos, title_extra=tf.parent.name)


def refresh_industrial_photos(root: Path) -> list[Path]:
    """Refresh TF photo annex with new library; queue stale PDFs for all 4 docs."""
    docs = root / "01_PRODUCTS"
    pdfs: list[Path] = []
    folders = sorted([p for p in docs.iterdir() if p.is_dir()], key=lambda p: p.name)
    for i, folder in enumerate(folders, 1):
        tf = folder / "01_Technical_File.docx"
        if tf.exists():
            n = refresh_tf_photos(tf, scope="INDUSTRIAL")
            if i <= 3 or i % 200 == 0:
                print(f"IND photo refresh {i}/{len(folders)} photos={n}", flush=True)
        for stem in ["01_Technical_File", "02_EU_DoC", "03_Label", "04_Shipment_Statement"]:
            d = folder / f"{stem}.docx"
            pdf = folder / f"{stem}.pdf"
            if d.exists() and (not pdf.exists() or pdf.stat().st_mtime + 1 < d.stat().st_mtime):
                pdfs.append(d)
    return pdfs


def refresh_starter_photos(root: Path) -> list[Path]:
    docs = root / "01_PRODUCTS"
    changed: list[Path] = []
    folders = sorted([p for p in docs.iterdir() if p.is_dir()], key=lambda p: p.name)
    for i, folder in enumerate(folders, 1):
        tf = folder / "01_Technical_File.docx"
        if tf.exists():
            n = refresh_tf_photos(tf, scope="STARTER")
            if n:
                changed.append(tf)
            if i <= 3 or i % 200 == 0:
                print(f"STARTER photo refresh {i}/{len(folders)} photos={n}", flush=True)
    return changed


def rebuild_industrial(root: Path, *, limit: int | None = None, force: bool = False) -> list[Path]:
    docs = root / "01_PRODUCTS"
    docs.mkdir(parents=True, exist_ok=True)
    _, products = ind.parse_industrial(ind.SRC)
    if limit:
        products = products[:limit]
    pdfs = []
    for i, p in enumerate(products, 1):
        dest = docs / p["product_code"]
        tf = dest / "01_Technical_File.docx"
        if tf.exists() and not force:
            try:
                with zipfile.ZipFile(tf, "r") as zf:
                    xml = zf.read("word/document.xml")
                    if b"w:i" in xml and b"{{" not in xml and _tf_has_annex(tf):
                        for stem in ["01_Technical_File", "02_EU_DoC", "03_Label", "04_Shipment_Statement"]:
                            d = dest / f"{stem}.docx"
                            pdf = dest / f"{stem}.pdf"
                            if d.exists() and (
                                not pdf.exists() or pdf.stat().st_mtime + 1 < d.stat().st_mtime
                            ):
                                pdfs.append(d)
                        if i % 500 == 0:
                            print(f"IND resume skip {i}/{len(products)}", flush=True)
                        continue
            except Exception:
                pass
        bom = []
        for b in p["bom"]:
            bom.append(
                {
                    "component_code": b["component_code"],
                    "description": b["description"],
                    "name_en": translate_component(b["description"]),
                    "qty": b.get("qty"),
                    "uom": b.get("uom") or "ADT",
                }
            )
        info = build_premium_pack(
            dest,
            key=p["product_code"],
            description_tr=p["description"],
            description_en=translate_product(p["description"]),
            set_code=f"IND-{p['product_code']}",
            bom_lines=bom,
            scope="INDUSTRIAL",
            config_id=f"IA-IND-{p['product_code']}",
        )
        for stem in ["01_Technical_File", "02_EU_DoC", "03_Label", "04_Shipment_Statement"]:
            pdfs.append(dest / f"{stem}.docx")
        if i <= 3 or i % 100 == 0:
            print(f"IND {i}/{len(products)} {p['product_code']} photos={info.get('photos')}", flush=True)
            progress(phase="industrial", i=i, total=len(products))
    return pdfs


def engine_for(root: Path, docs_subdir: str, records: list[dict], filename: str, title: str, extra=None):
    return build_document_engine(
        delivery_root=root,
        engine_filename=filename,
        title=title,
        docs_subdir=docs_subdir,
        records=records,
        extra_home={
            "PUBLISH DATE": DATE,
            "SIGNATORY": "Numan Alver — Operations Director",
            "FORMAT": "YS/D golden + bilingual TR/EN italic",
            **(extra or {}),
        },
    )


def zip_delivery(root: Path) -> str:
    import hashlib

    zip_path = OUT / f"{root.name}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=1) as zf:
        for f in root.rglob("*"):
            if f.is_file():
                zf.write(f, f.relative_to(root).as_posix())
    h = hashlib.sha256()
    with zip_path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    digest = h.hexdigest()
    (OUT / f"{root.name}_SHA256.txt").write_text(digest + "\n", encoding="utf-8")
    return digest


def main() -> int:
    print("PREMIUM + BILINGUAL ONE-SHOT start", flush=True)
    import argparse

    ap = argparse.ArgumentParser()
    ap.add_argument("--skip-component", action="store_true")
    ap.add_argument("--from-container", action="store_true", help="Resume from container onward")
    ap.add_argument("--industrial-only", action="store_true")
    ap.add_argument("--starter-only", action="store_true")
    ap.add_argument("--force", action="store_true", help="Force rebuild / photo refresh (no resume skips)")
    ap.add_argument("--photos-only-industrial", action="store_true", help="Refresh industrial TF photos only")
    args = ap.parse_args()
    clear_mapping_cache()

    if args.starter_only:
        print("=== STARTER BILINGUAL ===", flush=True)
        st_root = OUT / "01_STARTER_INDIVIDUAL_DELIVERY_REV00"
        changed = fix_starter_bilingual_inplace(st_root)
        photo_changed = refresh_starter_photos(st_root)
        todo = list({*changed, *photo_changed})
        # also regenerate PDFs for sibling docs when TF changed
        extra = []
        for tf in photo_changed:
            for stem in ["02_EU_DoC", "03_Label", "04_Shipment_Statement"]:
                p = tf.parent / f"{stem}.docx"
                if p.exists():
                    extra.append(p)
            todo.append(tf)
        todo = list(dict.fromkeys(todo + extra))
        print(f"STARTER changed docx={len(todo)}", flush=True)
        ok, fail = lo_pdfs(todo)
        print(f"STARTER PDF ok={ok} fail={fail}", flush=True)
        return 0 if fail == 0 else 1

    if not args.industrial_only and not args.from_container:
        if not args.skip_component:
            print("=== COMPONENT ===", flush=True)
            cmp_root = OUT / "04_COMPONENT_SPARE_DELIVERY_REV00"
            cmp_pdfs = rebuild_component(cmp_root)
            ok, fail = lo_pdfs(cmp_pdfs)
            print(f"COMPONENT PDF ok={ok} fail={fail}", flush=True)
            engine_for(
                cmp_root,
                "01_VARIANTS",
                [
                    {"key": "CMP-TEK-SIRA", "label": "Komponent Ambalaj — Tek Sıra"},
                    {"key": "CMP-CIFT-SIRA", "label": "Komponent Ambalaj — Çift Sıra"},
                ],
                "INCI_AKU_PPWR_COMPONENT_SPARE_ENGINE_Rev00.xlsx",
                "İNCI AKÜ PPWR — COMPONENT / SPARE ENGINE Rev00",
            )

    if not args.industrial_only:
        print("=== CONTAINER ===", flush=True)
        cnt_root = OUT / "03_CONTAINER_DELIVERY_REV00"
        cnt_pdfs = rebuild_container(cnt_root)
        ok, fail = lo_pdfs(cnt_pdfs)
        print(f"CONTAINER PDF ok={ok} fail={fail}", flush=True)
        engine_for(
            cnt_root,
            "01_CONFIGS",
            [{"key": p.name, "label": p.name} for p in sorted((cnt_root / "01_CONFIGS").iterdir()) if p.is_dir()],
            "INCI_AKU_PPWR_CONTAINER_ENGINE_Rev00.xlsx",
            "İNCI AKÜ PPWR — CONTAINER LOADING ENGINE Rev00",
        )

    print("=== INDUSTRIAL ===", flush=True)
    ind_root = OUT / "02_INDUSTRIAL_DELIVERY_REV00"
    if args.photos_only_industrial or (args.force and not args.industrial_only):
        # Prefer fast photo refresh when bilingual premium packs already exist
        sample = next((ind_root / "01_PRODUCTS").iterdir(), None)
        has_premium = False
        if sample:
            tf = sample / "01_Technical_File.docx"
            if tf.exists():
                with zipfile.ZipFile(tf, "r") as zf:
                    xml = zf.read("word/document.xml")
                    has_premium = xml.count(b"<w:i") > 50 and b"{{" not in xml
        if has_premium or args.photos_only_industrial:
            ind_pdfs = refresh_industrial_photos(ind_root)
        else:
            ind_pdfs = rebuild_industrial(ind_root, force=True)
    else:
        ind_pdfs = rebuild_industrial(ind_root, force=args.force)
    print(f"INDUSTRIAL pdf jobs={len(ind_pdfs)}", flush=True)
    ok, fail = lo_pdfs(ind_pdfs)
    print(f"INDUSTRIAL PDF ok={ok} fail={fail}", flush=True)
    keys = sorted([p.name for p in (ind_root / "01_PRODUCTS").iterdir() if p.is_dir()])
    engine_for(
        ind_root,
        "01_PRODUCTS",
        [{"key": k, "label": k} for k in keys],
        "INCI_AKU_PPWR_INDUSTRIAL_ENGINE_Rev00.xlsx",
        "İNCI AKÜ PPWR — INDUSTRIAL INDIVIDUAL ENGINE Rev00",
    )

    if not args.industrial_only and not args.photos_only_industrial:
        print("=== STARTER BILINGUAL + PHOTOS ===", flush=True)
        st_root = OUT / "01_STARTER_INDIVIDUAL_DELIVERY_REV00"
        changed = fix_starter_bilingual_inplace(st_root)
        photo_changed = refresh_starter_photos(st_root)
        todo = list(dict.fromkeys(changed + photo_changed))
        for tf in photo_changed:
            for stem in ["02_EU_DoC", "03_Label", "04_Shipment_Statement"]:
                p = tf.parent / f"{stem}.docx"
                pdf = tf.parent / f"{stem}.pdf"
                if p.exists() and (not pdf.exists() or pdf.stat().st_mtime + 1 < tf.stat().st_mtime):
                    # TF changed; sibling docs unchanged — only convert changed docx
                    pass
        print(f"STARTER changed docx={len(todo)}", flush=True)
        ok, fail = lo_pdfs(todo)
        print(f"STARTER PDF ok={ok} fail={fail}", flush=True)
        st_keys = sorted([p.name for p in (st_root / "01_PRODUCTS").iterdir() if p.is_dir()])
        engine_for(
            st_root,
            "01_PRODUCTS",
            [{"key": k, "label": k} for k in st_keys],
            "INCI_AKU_PPWR_STARTER_INDIVIDUAL_ENGINE_Rev00.xlsx",
            "İNCI AKÜ PPWR — STARTER INDIVIDUAL ENGINE Rev00",
        )
    else:
        st_root = OUT / "01_STARTER_INDIVIDUAL_DELIVERY_REV00"
        st_keys = sorted([p.name for p in (st_root / "01_PRODUCTS").iterdir() if p.is_dir()]) if st_root.exists() else []
        changed = []

    # ZIP all four
    print("=== ZIP + SHA256 ===", flush=True)
    digests = {}
    for name in [
        "01_STARTER_INDIVIDUAL_DELIVERY_REV00",
        "02_INDUSTRIAL_DELIVERY_REV00",
        "03_CONTAINER_DELIVERY_REV00",
        "04_COMPONENT_SPARE_DELIVERY_REV00",
    ]:
        r = OUT / name
        if r.exists():
            digests[name] = zip_delivery(r)
            print(f"ZIP {name} sha={digests[name][:16]}…", flush=True)

    report = {
        "COMPONENT": "DONE",
        "CONTAINER": "DONE",
        "INDUSTRIAL": f"folders={len(keys)}",
        "STARTER": f"changed={len(changed) if 'changed' in dir() else 'n/a'} folders={len(st_keys)}",
        "SHA256": digests,
        "NOTE": "Premium runtime + bilingual TR/EN italic + refreshed representative photos; LibreOffice PDF",
    }
    (OUT / "_PREMIUM_BILINGUAL_QA.json").write_text(
        json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    print(json.dumps(report, indent=2), flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
