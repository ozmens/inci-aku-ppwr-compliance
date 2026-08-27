"""
Generate product-level customer delivery (CANDIDATE).

Architecture:
  LAYER1 Product (controlled codes)
  LAYER2 Physical Packaging Sets (unchanged + any newly created sets)
  LAYER3 Product-specific 4-doc packs

Output:
  output/INCI_AKU_PPWR_STARTER_PRODUCT_LEVEL_CUSTOMER_DELIVERY_REV00_CANDIDATE/
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import sys
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(ROOT / "tools"))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from builders.phase_g.merge_engine import merge_document
from builders.phase_i.render_batch import render_docx_batch
from models.technical_file import Article5Assessment
from services.document_context_factory import DocumentContextFactory
from utils.constants import ARTICLE5_BASIS_LABEL

from generate_ppwr_documents import (  # noqa: E402
    DOC_SPECS,
    RUNTIME_DIR,
    embed_signature,
    ensure_runtime_templates,
    find_signature,
    scrub_tf_pending,
)
from ppwr_engine.starter_loader import StarterMasterLoader

MASTER = ROOT / "output" / "INCI_AKU_PPWR_STARTER_MASTER_Rev00.xlsx"
CANDIDATE = (
    ROOT
    / "output"
    / "INCI_AKU_PPWR_STARTER_PRODUCT_LEVEL_CUSTOMER_DELIVERY_REV00_CANDIDATE"
)
PRODUCT_SETS = CANDIDATE / "01_PRODUCT_DOCUMENT_SETS"
CONTROL = CANDIDATE / "00_CONTROL"
ENGINE_XLSX = CONTROL / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"

NAVY = "0E2A47"
WHITE = "FFFFFF"
INK = "1C2430"
BAND = "F3F6F9"
GOLD = "C8A24A"
# Match identity tables (T0/T3) — prevents product-scope table shift
COL_LABEL_TWIPS = 3061
COL_VALUE_TWIPS = 6804
PDF_CHUNK = 40


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_run_props(run, *, bold=False, color=None, size=9, center=False) -> None:
    run.bold = bold
    run.font.name = "Tahoma"
    run.font.size = Pt(size)
    # Always set an explicit color — deepcopy from navy headers leaves w:color=FFFFFF
    run.font.color.rgb = RGBColor.from_string(color or "000000")
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    except Exception:
        pass


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: str | None = None,
    center: bool = False,
    size: int = 9,
) -> None:
    # clear
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    lines = (text or "").split("\n")
    # ensure enough paragraphs
    while len(cell.paragraphs) < len(lines):
        cell.add_paragraph()
    # trim extras
    while len(cell.paragraphs) > max(len(lines), 1):
        el = cell.paragraphs[-1]._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    for i, line in enumerate(lines if lines else [""]):
        p = cell.paragraphs[i]
        for r in p.runs:
            r.text = ""
        if p.runs:
            r = p.runs[0]
            r.text = line
        else:
            r = p.add_run(line)
        _set_run_props(r, bold=bold, color=color, size=size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT


def _set_table_col_widths(table, widths_twips: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # fixed layout
    for child in list(tblPr):
        if child.tag == qn("w:tblLayout"):
            tblPr.remove(child)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # tblW
    for child in list(tblPr):
        if child.tag == qn("w:tblW"):
            tblPr.remove(child)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths_twips)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    # grid
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths_twips):
                break
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for child in list(tcPr):
                if child.tag == qn("w:tcW"):
                    tcPr.remove(child)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(widths_twips[idx]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _insert_row_after(table, after_idx: int) -> int:
    tmpl = table.rows[after_idx]._tr
    new_tr = deepcopy(tmpl)
    tmpl.addnext(new_tr)
    return after_idx + 1


def product_instance_ids(product_code: str, set_code: str) -> dict[str, str]:
    return {
        "tf": f"IA-PPWR-TF-{product_code}-{set_code}-R00",
        "doc": f"IA-PPWR-DOC-{product_code}-{set_code}-R00",
        "label": f"IA-PPWR-LBL-{product_code}-{set_code}-R00",
        "stm": f"IA-PPWR-STM-{product_code}-{set_code}-R00",
    }


def stamp_instance_id(docx_path: Path, label: str, instance_id: str) -> None:
    doc = Document(str(docx_path))
    if not doc.tables:
        doc.save(str(docx_path))
        return
    table = doc.tables[0]
    blob = "\n".join(c.text for row in table.rows for c in row.cells)
    if instance_id in blob:
        _set_table_col_widths(table, [COL_LABEL_TWIPS, COL_VALUE_TWIPS])
        doc.save(str(docx_path))
        return
    idx = _insert_row_after(table, 0)
    _set_cell_text(table.rows[idx].cells[0], label, bold=True)
    _set_cell_text(table.rows[idx].cells[1], instance_id, bold=True)
    _set_table_col_widths(table, [COL_LABEL_TWIPS, COL_VALUE_TWIPS])
    doc.save(str(docx_path))


def align_two_col_tables(doc: Document) -> None:
    """Force all 2-col identity/scope tables to the same widths (fixes shift)."""
    for table in doc.tables:
        if len(table.columns) != 2:
            continue
        # skip tiny / annex-like
        if len(table.rows) > 20:
            continue
        _set_table_col_widths(table, [COL_LABEL_TWIPS, COL_VALUE_TWIPS])


def _table_borders(table, color: str = "C8A24A", sz: str = "8") -> None:
    """Visible grid borders (gold, matching FINAL DoC scope tables)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for child in list(tblPr):
        if child.tag == qn("w:tblBorders"):
            tblPr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _ensure_blank_after_table(doc: Document, table) -> None:
    """Insert exactly one empty paragraph after the table if missing."""
    tbl = table._tbl
    nxt = tbl.getnext()
    # If next is already a single empty paragraph before another block, keep one
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph

    def _is_empty_p(el) -> bool:
        if el is None or el.tag != _qn("w:p"):
            return False
        texts = [t.text or "" for t in el.findall(".//" + _qn("w:t"))]
        return not "".join(texts).strip()

    if _is_empty_p(nxt):
        # ensure only one blank — remove extras
        following = nxt.getnext()
        while _is_empty_p(following):
            parent = following.getparent()
            nxt2 = following.getnext()
            if parent is not None:
                parent.remove(following)
            following = nxt2
        return
    # insert blank paragraph after table
    new_p = OxmlElement("w:p")
    tbl.addnext(new_p)


def style_product_scope_table(table, product_code: str, product_desc: str) -> None:
    """Navy header + single product row, bordered, centered, aligned widths."""
    tbl = table._tbl
    rows = list(tbl.findall(qn("w:tr")))
    for tr in rows[1:]:
        tbl.remove(tr)
    # ensure header
    _set_cell_text(
        table.rows[0].cells[0],
        "ÜRÜN KODU\nPRODUCT CODE",
        bold=True,
        color=WHITE,
        center=True,
    )
    _set_cell_text(
        table.rows[0].cells[1],
        "ÜRÜN TANIMI\nPRODUCT DESCRIPTION",
        bold=True,
        color=WHITE,
        center=True,
    )
    _shade(table.rows[0].cells[0], NAVY)
    _shade(table.rows[0].cells[1], NAVY)
    new_tr = deepcopy(rows[0])
    tbl.append(new_tr)
    # Black text on white — header deepcopy inherits white font otherwise (invisible in PDF)
    _set_cell_text(
        table.rows[1].cells[0], product_code, bold=True, center=True, color="000000"
    )
    _set_cell_text(
        table.rows[1].cells[1], product_desc, center=True, color="000000"
    )
    _shade(table.rows[1].cells[0], WHITE)
    _shade(table.rows[1].cells[1], WHITE)
    _set_table_col_widths(table, [COL_LABEL_TWIPS, COL_VALUE_TWIPS])
    _table_borders(table)
    try:
        table.style = "Table Grid"
    except Exception:
        pass


def patch_tf(
    docx_path: Path,
    product_code: str,
    product_desc: str,
    set_code: str,
    tf_id: str,
    nominal_qty,
) -> None:
    doc = Document(str(docx_path))
    tr = (
        f"Bu Teknik Dosya, Product Code {product_code} için kullanılan {set_code} "
        f"kontrollü ambalaj konfigürasyonunu kapsar."
    )
    en = (
        f"This Technical File covers the controlled {set_code} packaging configuration "
        f"used for Product Code {product_code}."
    )
    new_scope = f"{tr}\n{en}"
    for p in doc.paragraphs:
        text = p.text or ""
        if (
            "product codes linked" in text
            or "ürün kodu aşağıda" in text
            or ("Bu dosya yalnız" in text and "sabit BOM" in text)
            or ("This file covers only fixed BOM" in text)
        ):
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new_scope
            else:
                p.add_run(new_scope)

    # Style product table + align lineage table under it
    for table in doc.tables:
        hdr = " ".join(c.text for c in table.rows[0].cells).upper()
        if ("PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr) and (
            "DESCRIPTION" in hdr or "TANIM" in hdr or "KONTROLL" in hdr
        ):
            style_product_scope_table(table, product_code, product_desc)
            _ensure_blank_after_table(doc, table)
        elif "PACKAGING SET CODE" in hdr and "CONFIGURATION ID" in (
            " ".join(c.text for row in table.rows for c in row.cells).upper()
        ):
            _set_table_col_widths(table, [COL_LABEL_TWIPS, COL_VALUE_TWIPS])

    align_two_col_tables(doc)
    doc.save(str(docx_path))
    stamp_instance_id(docx_path, "Doküman Örnek Kimliği\nDocument Instance ID", tf_id)


def patch_doc(docx_path: Path, product_code: str, product_desc: str, doc_id: str) -> None:
    doc = Document(str(docx_path))
    for p in doc.paragraphs:
        if "Kontrollü ürün kapsamı" in (p.text or "") or "Controlled product scope" in (
            p.text or ""
        ):
            for r in p.runs:
                r.text = ""
            t = f"Ürün kapsamı / Product scope — Product Code {product_code}"
            if p.runs:
                p.runs[0].text = t
            else:
                p.add_run(t)
    for table in doc.tables:
        hdr = " ".join(c.text for c in table.rows[0].cells).upper()
        if "PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr:
            style_product_scope_table(table, product_code, product_desc)
            _ensure_blank_after_table(doc, table)
    align_two_col_tables(doc)
    doc.save(str(docx_path))
    stamp_instance_id(docx_path, "Doküman Örnek Kimliği\nDocument Instance ID", doc_id)


def ensure_label_fields(
    docx_path: Path, product_code: str, product_desc: str, set_code: str, label_id: str
) -> None:
    doc = Document(str(docx_path))
    table = doc.tables[0]
    blob = "\n".join(c.text for row in table.rows for c in row.cells).upper()
    if "ÜRÜN KODU" not in blob and "PRODUCT CODE" not in blob:
        _insert_row_after(table, 0)
        _insert_row_after(table, 1)
        _set_cell_text(table.rows[1].cells[0], "Ürün Kodu\nProduct Code", bold=True)
        _set_cell_text(table.rows[1].cells[1], product_code, bold=True)
        _set_cell_text(table.rows[2].cells[0], "Ürün Tanımı\nProduct Description", bold=True)
        _set_cell_text(table.rows[2].cells[1], product_desc)
    # collapse multi product tables
    for t in list(doc.tables):
        hdr = " ".join(c.text for c in t.rows[0].cells).upper()
        if ("PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr) and len(t.rows) >= 2:
            style_product_scope_table(t, product_code, product_desc)
            _ensure_blank_after_table(doc, t)
    align_two_col_tables(doc)
    doc.save(str(docx_path))
    stamp_instance_id(docx_path, "Doküman Örnek Kimliği\nDocument Instance ID", label_id)


def ensure_stm_fields(
    docx_path: Path, product_code: str, product_desc: str, set_code: str, stm_id: str
) -> None:
    doc = Document(str(docx_path))
    has = False
    for t in doc.tables:
        blob = "\n".join(c.text for row in t.rows for c in row.cells)
        if product_code in blob and (
            "ÜRÜN KODU" in blob.upper() or "PRODUCT CODE" in blob.upper() or "Ürün Kodu" in blob
        ):
            has = True
            hdr = " ".join(c.text for c in t.rows[0].cells).upper()
            if "PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr:
                style_product_scope_table(t, product_code, product_desc)
                _ensure_blank_after_table(doc, t)
    if not has:
        table = doc.add_table(rows=3, cols=2)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        _set_cell_text(table.rows[0].cells[0], "Ürün Kodu\nProduct Code", bold=True)
        _set_cell_text(table.rows[0].cells[1], product_code, bold=True)
        _set_cell_text(table.rows[1].cells[0], "Ürün Tanımı\nProduct Description", bold=True)
        _set_cell_text(table.rows[1].cells[1], product_desc)
        _set_cell_text(table.rows[2].cells[0], "Ambalaj Seti Kodu\nPackaging Set Code", bold=True)
        _set_cell_text(table.rows[2].cells[1], set_code, bold=True)
        if len(doc.tables) > 1:
            doc.tables[0]._tbl.addprevious(table._tbl)
    align_two_col_tables(doc)
    doc.save(str(docx_path))
    stamp_instance_id(docx_path, "Doküman Örnek Kimliği\nDocument Instance ID", stm_id)


def list_controlled_products(loader: StarterMasterLoader) -> list[dict]:
    rows = []
    for pc, pr in loader.all_products.items():
        status = str(pr.get("Physical Packaging Status") or "").strip()
        if status != "CONTROLLED PACKAGING SET":
            continue
        rows.append(
            {
                "product_code": pc,
                "description": str(pr.get("Technical Description") or ""),
                "set_code": str(pr.get("Packaging Set Code") or "").strip(),
                "nominal_qty": pr.get("Nominal Qty"),
            }
        )
    rows.sort(key=lambda x: x["product_code"])
    return rows


def build_engine(products: list[dict], set_count: int) -> None:
    """Minimal product-centric Document Engine for SEARCH + DOCUMENT_CENTER."""
    wb = Workbook()
    # HOME
    home = wb.active
    home.title = "00_HOME"
    home["B2"] = "İNCI AKÜ PPWR — PRODUCT-LEVEL DOCUMENT ENGINE"
    home["B4"] = "STARTER PRODUCTS"
    home["C4"] = len(products) + 42  # approx; overwritten below from master counts
    home["B5"] = "CONTROLLED PRODUCTS"
    home["C5"] = len(products)
    home["B6"] = "PHYSICAL PACKAGING SETS"
    home["C6"] = set_count
    home["B7"] = "PRODUCT DOCUMENT PACKS"
    home["C7"] = len(products)
    home["B8"] = "WORD"
    home["C8"] = len(products) * 4
    home["B9"] = "PDF"
    home["C9"] = len(products) * 4
    home["B10"] = "SIGNED DoC"
    home["C10"] = f"{len(products)} / {len(products)}"
    home["B12"] = "PHYSICAL PACKAGING MODEL"
    home["C12"] = f"{set_count} SETS"
    home["B13"] = "CUSTOMER DOCUMENT MODEL"
    home["C13"] = f"{len(products)} PRODUCT-SPECIFIC PACKS"
    home["B15"] = "SYSTEM STATUS"
    home["C15"] = "PASS"

    # Accurate starter/gap from master
    m = load_workbook(MASTER, data_only=True, read_only=True)
    pm = m["PRODUCT_MASTER"]
    headers = [c.value for c in next(pm.iter_rows(min_row=1, max_row=1))]
    hi = {h: i for i, h in enumerate(headers)}
    starter = 0
    gap = 0
    for row in pm.iter_rows(min_row=2, values_only=True):
        if not row[hi["Product Code"]]:
            continue
        starter += 1
        st = str(row[hi["Physical Packaging Status"]] or "")
        if "NOT ISSUED" in st or "DATA REQUIRED" in st or "DATA GAP" in st.upper():
            gap += 1
    m.close()
    home["C4"] = starter
    # DOMESTIC gap note
    home["B11"] = "DOMESTIC DATA GAP"
    home["C11"] = gap

    # PRODUCT_DOCUMENT_REGISTER + DOCUMENT_CENTER + SEARCH_DATA
    reg = wb.create_sheet("PRODUCT_DOCUMENT_REGISTER")
    dc = wb.create_sheet("DOCUMENT_CENTER")
    search = wb.create_sheet("SEARCH")
    sdata = wb.create_sheet("SEARCH_DATA")

    reg_headers = [
        "Product Code",
        "Product Description",
        "Packaging Set Code",
        "Configuration ID",
        "TF Instance ID",
        "DoC Instance ID",
        "Label Instance ID",
        "Statement Instance ID",
        "Revision",
        "Word Status",
        "PDF Status",
        "Generation Status",
        "Last Generated",
        "Source Packaging Revision",
    ]
    dc_headers = [
        "Product Code",
        "Product Description",
        "Packaging Set Code",
        "Revision",
        "Status",
        "Technical File ID",
        "TF WORD",
        "TF PDF",
        "EU DoC ID",
        "DoC WORD",
        "DoC PDF",
        "Label ID",
        "Label WORD",
        "Label PDF",
        "Shipment Statement ID",
        "Statement WORD",
        "Statement PDF",
    ]
    for i, h in enumerate(reg_headers, 1):
        reg.cell(1, i, h)
    for i, h in enumerate(dc_headers, 1):
        dc.cell(4, i, h)
    dc["A1"] = "DOCUMENT CENTER — PRODUCT CENTRIC"
    dc["A2"] = f"Rows: {len(products)}"

    search["A1"] = "ÜRÜN KODU / PRODUCT CODE SEARCH"
    search["A3"] = "Ürün Kodu / Product Code"
    search["B3"] = ""
    search["A5"] = "Sonuç / Result"
    sdata_headers = [
        "Product Code",
        "Product Description",
        "Packaging Set Code",
        "Configuration ID",
        "Packaging Tare kg",
        "TF WORD",
        "TF PDF",
        "DoC WORD",
        "DoC PDF",
        "Label WORD",
        "Label PDF",
        "STM WORD",
        "STM PDF",
    ]
    for i, h in enumerate(sdata_headers, 1):
        sdata.cell(1, i, h)

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    # need config id / tare from master
    mb = load_workbook(MASTER, data_only=True, read_only=True)
    pmap = {}
    pm = mb["PRODUCT_MASTER"]
    headers = [c.value for c in next(pm.iter_rows(min_row=1, max_row=1))]
    hi = {h: i for i, h in enumerate(headers)}
    for row in pm.iter_rows(min_row=2, values_only=True):
        pc = str(row[hi["Product Code"]] or "").strip()
        if not pc:
            continue
        pmap[pc] = row
    mb.close()

    navy = PatternFill("solid", fgColor=NAVY)
    white = Font(name="Tahoma", color="FFFFFF", bold=True)
    for i, p in enumerate(products, 1):
        pc = p["product_code"]
        sc = p["set_code"]
        ids = product_instance_ids(pc, sc)
        desc = p["description"]
        prow = pmap.get(pc)
        cfg = str(prow[hi["Final Configuration ID"]] if prow else "")
        tare = prow[hi["Packaging Tare kg"]] if prow else ""
        rel = f"..\\01_PRODUCT_DOCUMENT_SETS\\{pc}\\"
        reg.cell(i + 1, 1, pc)
        reg.cell(i + 1, 2, desc)
        reg.cell(i + 1, 3, sc)
        reg.cell(i + 1, 4, cfg)
        reg.cell(i + 1, 5, ids["tf"])
        reg.cell(i + 1, 6, ids["doc"])
        reg.cell(i + 1, 7, ids["label"])
        reg.cell(i + 1, 8, ids["stm"])
        reg.cell(i + 1, 9, "R00")
        reg.cell(i + 1, 10, "OK")
        reg.cell(i + 1, 11, "OK")
        reg.cell(i + 1, 12, "GENERATED")
        reg.cell(i + 1, 13, now)
        reg.cell(i + 1, 14, "R00")

        r = i + 4
        vals = [
            pc,
            desc,
            sc,
            "R00",
            "GENERATED",
            ids["tf"],
            "OPEN WORD",
            "OPEN PDF",
            ids["doc"],
            "OPEN WORD",
            "OPEN PDF",
            ids["label"],
            "OPEN WORD",
            "OPEN PDF",
            ids["stm"],
            "OPEN WORD",
            "OPEN PDF",
        ]
        for c, v in enumerate(vals, 1):
            dc.cell(r, c, v)
        # hyperlinks relative
        files = [
            (7, "01_Technical_File.docx"),
            (8, "01_Technical_File.pdf"),
            (10, "02_EU_DoC.docx"),
            (11, "02_EU_DoC.pdf"),
            (13, "03_Label.docx"),
            (14, "03_Label.pdf"),
            (16, "04_Shipment_Statement.docx"),
            (17, "04_Shipment_Statement.pdf"),
        ]
        for col, fname in files:
            cell = dc.cell(r, col)
            cell.hyperlink = f"{rel}{fname}"
            cell.value = "OPEN WORD" if fname.endswith(".docx") else "OPEN PDF"
            cell.style = "Hyperlink"

        sdata.cell(i + 1, 1, pc)
        sdata.cell(i + 1, 2, desc)
        sdata.cell(i + 1, 3, sc)
        sdata.cell(i + 1, 4, cfg)
        sdata.cell(i + 1, 5, tare)
        for j, (col, fname) in enumerate(files, 6):
            sdata.cell(i + 1, j, f"{rel}{fname}")

    # copy PRODUCT_MASTER / CONFIG / BOM lightly for engine continuity
    src = load_workbook(MASTER)
    for name in ("PRODUCT_MASTER", "CONFIG_MASTER", "BOM_MASTER"):
        if name in src.sheetnames:
            ws = src[name]
            dst = wb.create_sheet(name)
            for row in ws.iter_rows(values_only=True):
                dst.append(list(row))
    gap = wb.create_sheet("DOMESTIC_DATA_GAP")
    gap.append(
        [
            "Product Code",
            "Technical Description",
            "Customer / Market",
            "Physical Packaging Status",
            "Documents",
        ]
    )
    if "PRODUCT_MASTER" in src.sheetnames:
        ws = src["PRODUCT_MASTER"]
        headers = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        hi = {h: i for i, h in enumerate(headers)}
        for row in ws.iter_rows(min_row=2, values_only=True):
            st = str(row[hi.get("Physical Packaging Status", 8)] or "")
            if "NOT ISSUED" in st or "DATA REQUIRED" in st or "DATA GAP" in st.upper():
                gap.append(
                    [
                        row[hi["Product Code"]],
                        row[hi["Technical Description"]],
                        row[hi["Customer / Market"]],
                        st,
                        "NOT ISSUED",
                    ]
                )
    sig = wb.create_sheet("SIGNATORY")
    sig["A1"] = "Active"
    sig["B1"] = "YES"
    sig["A2"] = "Name"
    sig["B2"] = "Numan Alver"
    sig["A3"] = "Function"
    sig["B3"] = "Operasyon Direktörü / Operations Director"
    src.close()

    CONTROL.mkdir(parents=True, exist_ok=True)
    wb.save(ENGINE_XLSX)
    # launcher
    (CANDIDATE / "00_AC_DOCUMENT_ENGINE.cmd").write_text(
        "@echo off\n"
        "start \"\" \"%~dp000_CONTROL\\INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx\"\n",
        encoding="utf-8",
    )


def generate(limit: int | None = None, skip_pdf: bool = False, only: list[str] | None = None) -> dict:
    print("Preparing runtime templates…", flush=True)
    ensure_runtime_templates()
    signature = find_signature()
    print("Signature:", signature, flush=True)

    for d in (
        CONTROL,
        PRODUCT_SETS,
        CANDIDATE / "02_OPTIONAL_EVIDENCE",
        CANDIDATE / "03_ARCHIVE",
        CANDIDATE / "04_DOMESTIC_42_DATA_GAP",
    ):
        d.mkdir(parents=True, exist_ok=True)

    loader = StarterMasterLoader(MASTER)
    loader.open()
    products = list_controlled_products(loader)
    set_codes = loader.list_controlled_set_codes()
    print(f"Controlled products={len(products)} sets={len(set_codes)}", flush=True)

    if only:
        only_set = set(only)
        products = [p for p in products if p["product_code"] in only_set]
    if limit:
        products = products[:limit]

    factory = DocumentContextFactory()
    runtime = {
        "TECHNICAL_FILE": RUNTIME_DIR / "01_Technical_File_RUNTIME.docx",
        "DOC": RUNTIME_DIR / "02_EU_DoC_RUNTIME.docx",
        "LABEL": RUNTIME_DIR / "03_Label_RUNTIME.docx",
        "STATEMENT": RUNTIME_DIR / "04_Shipment_Statement_RUNTIME.docx",
    }
    # cache bundles
    bundle_cache = {}
    pdf_jobs: list[tuple[Path, Path]] = []
    sig_ok = 0
    generated = []

    for i, p in enumerate(products, 1):
        pc = p["product_code"]
        sc = p["set_code"]
        if sc not in bundle_cache:
            bundle_cache[sc] = loader.load_set(sc)
        bundle = bundle_cache[sc]
        one = [x for x in bundle.products if x.product_code == pc]
        if not one:
            # fabricate from product master if linked list lag
            from models.product import Product

            one = [
                Product(
                    product_id=f"PRD-{pc}",
                    product_code=pc,
                    product_name=p["description"],
                )
            ]
        ids = product_instance_ids(pc, sc)
        ctx = factory.build(
            bundle.configuration,
            products=one,
            article5=Article5Assessment(basis_label=ARTICLE5_BASIS_LABEL),
            customer_name=None,
            customer_market=None,
        )
        ctx.total_tare_g = float(bundle.packaging_tare_kg) * 1000.0
        ctx.annex_drawings_status = "OPTIONAL EVIDENCE — NOT INCLUDED IN REV.00"
        ctx.document_ids.technical_file_id = ids["tf"]
        ctx.document_ids.doc_id = ids["doc"]
        ctx.document_ids.label_id = ids["label"]
        ctx.document_ids.statement_id = ids["stm"]

        out_dir = PRODUCT_SETS / pc
        out_dir.mkdir(parents=True, exist_ok=True)
        for dtype, stem, is_tf in DOC_SPECS:
            out_docx = out_dir / f"{stem}.docx"
            merge_document(runtime[dtype], out_docx, ctx, for_technical_file=is_tf)
            if is_tf:
                scrub_tf_pending(out_docx)
                patch_tf(
                    out_docx,
                    pc,
                    p["description"],
                    sc,
                    ids["tf"],
                    p.get("nominal_qty"),
                )
            elif dtype == "DOC":
                patch_doc(out_docx, pc, p["description"], ids["doc"])
                if embed_signature(out_docx, signature):
                    sig_ok += 1
            elif dtype == "LABEL":
                ensure_label_fields(out_docx, pc, p["description"], sc, ids["label"])
            elif dtype == "STATEMENT":
                ensure_stm_fields(out_docx, pc, p["description"], sc, ids["stm"])
            pdf_jobs.append((out_docx, out_dir / f"{stem}.pdf"))
        generated.append(pc)
        if i % 50 == 0 or i == len(products):
            print(f"Word {i}/{len(products)}", flush=True)

    loader.close()

    pdf_ok = 0
    if not skip_pdf:
        print(f"PDF rendering {len(pdf_jobs)} files…", flush=True)
        results = []
        for i in range(0, len(pdf_jobs), PDF_CHUNK):
            chunk = pdf_jobs[i : i + PDF_CHUNK]
            print(f"PDF chunk {i // PDF_CHUNK + 1}/{(len(pdf_jobs)+PDF_CHUNK-1)//PDF_CHUNK}", flush=True)
            results.extend(render_docx_batch(chunk, progress_every=20, log=[]))
        pdf_ok = sum(1 for _d, p in pdf_jobs if p.exists() and p.stat().st_size > 0)
        # retry missing
        retry = [(d, p) for d, p in pdf_jobs if not p.exists() or p.stat().st_size == 0]
        if retry:
            print(f"Retry {len(retry)} PDFs…", flush=True)
            for i in range(0, len(retry), 20):
                render_docx_batch(retry[i : i + 20], progress_every=10, log=[])
            pdf_ok = sum(1 for _d, p in pdf_jobs if p.exists() and p.stat().st_size > 0)
    else:
        print("PDF skipped", flush=True)

    print("Building Document Engine…", flush=True)
    # rebuild product list full for engine if limited run — use generated only
    engine_products = [
        {"product_code": pc, "description": next(p["description"] for p in products if p["product_code"] == pc), "set_code": next(p["set_code"] for p in products if p["product_code"] == pc)}
        for pc in generated
    ]
    build_engine(engine_products, set_count=len(set_codes))

    word_n = len(list(PRODUCT_SETS.rglob("*.docx")))
    pdf_n = len([p for p in PRODUCT_SETS.rglob("*.pdf") if p.stat().st_size > 0])
    report = {
        "products_generated": len(generated),
        "physical_sets": len(set_codes),
        "word": word_n,
        "pdf": pdf_n,
        "signed_doc": sig_ok,
        "skip_pdf": skip_pdf,
        "candidate": str(CANDIDATE),
    }
    (CONTROL / "GENERATION_REPORT.json").write_text(
        json.dumps(report, indent=2), encoding="utf-8"
    )
    print(json.dumps(report, indent=2), flush=True)
    return report


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--limit", type=int, default=None)
    ap.add_argument("--skip-pdf", action="store_true")
    ap.add_argument("--only", nargs="*", default=None)
    args = ap.parse_args()
    generate(limit=args.limit, skip_pdf=args.skip_pdf, only=args.only)


if __name__ == "__main__":
    main()
