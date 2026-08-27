"""Fix product-level DOCX visibility + date issues (no Word/Excel UI).

1) Product-scope data rows inherited white font from navy headers → invisible in PDF
2) DoC place/date underscore → Manisa / Türkiye • 11.08.2026
3) Label packing-date underscore → 11.08.2026
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PRODUCT_SETS = (
    ROOT
    / "output"
    / "INCI_AKU_PPWR_STARTER_PRODUCT_LEVEL_CUSTOMER_DELIVERY_REV00_CANDIDATE"
    / "01_PRODUCT_DOCUMENT_SETS"
)
DATE = "11.08.2026"
PLACE = f"Manisa / Türkiye  •  {DATE}"
BLACK = "000000"
WHITE = "FFFFFF"
NAVY = "0E2A47"


def _set_run_props(run, *, bold=False, color=BLACK, size=9) -> None:
    run.bold = bold
    run.font.name = "Tahoma"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    except Exception:
        pass


def _set_cell_text(cell, text: str, *, bold=False, color=BLACK, center=True, size=9) -> None:
    lines = (text or "").split("\n")
    while len(cell.paragraphs) < len(lines):
        cell.add_paragraph()
    while len(cell.paragraphs) > max(len(lines), 1):
        el = cell.paragraphs[-1]._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    for i, line in enumerate(lines if lines else [""]):
        p = cell.paragraphs[i]
        for r in p.runs:
            r.text = ""
        if p.runs:
            r = p.runs[0]
            r.text = line
        else:
            r = p.add_run(line)
        _set_run_props(r, bold=bold, color=color, size=size)
        # clear leftover empty runs that still carry white color
        for r in p.runs[1:]:
            r.text = ""
            _set_run_props(r, bold=False, color=color, size=size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT


def _force_black_runs_in_cell(cell) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            if (r.text or "").strip():
                _set_run_props(r, bold=bool(r.bold), color=BLACK, size=9)
            else:
                # wipe invisible leftover runs
                r.text = ""


def fix_scope_tables(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.columns) < 2:
            continue
        hdr = " ".join(c.text for c in table.rows[0].cells).upper()
        if not (("PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr) and ("DESCRIPTION" in hdr or "TANIM" in hdr)):
            continue
        # header stays white-on-navy; data rows must be black
        for row in table.rows[1:]:
            for cell in row.cells:
                _force_black_runs_in_cell(cell)
                # if somehow empty but we expect content, leave as-is
            n += 1
    return n


def fix_doc_place_date(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            left = row.cells[0].text.upper()
            if "VERİLİŞ" in left or "PLACE AND DATE" in left:
                right = row.cells[1].text
                if "Manisa" in right or "______" in right or "Türkiye" in right:
                    _set_cell_text(row.cells[1], PLACE, bold=False, color=BLACK, center=False)
                    n += 1
    return n


def fix_label_packing_date(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        hdr = " ".join(c.text for c in table.rows[0].cells).upper()
        if "PACKING DATE" not in hdr and "PAKETLEME TARİHİ" not in hdr and "PAKETLEME TARIHI" not in hdr:
            continue
        # find packing date column
        col = None
        for i, c in enumerate(table.rows[0].cells):
            t = c.text.upper().replace("İ", "I")
            if "PACKING DATE" in t or "PAKETLEME TARIHI" in t or "PAKETLEME TARİHİ" in c.text.upper():
                col = i
                break
            # Turkish header without ASCII fold
            if "Paketleme tarihi" in c.text:
                col = i
                break
        if col is None:
            continue
        for row in table.rows[1:]:
            if col < len(row.cells):
                _set_cell_text(row.cells[col], DATE, bold=True, color=BLACK, center=True)
                n += 1
    return n


def fix_file(path: Path) -> dict:
    doc = Document(str(path))
    info = {"scope": 0, "place": 0, "pack": 0}
    info["scope"] = fix_scope_tables(doc)
    name = path.name.lower()
    if "doc" in name or "02_" in name:
        info["place"] = fix_doc_place_date(doc)
    if "label" in name or "03_" in name:
        info["pack"] = fix_label_packing_date(doc)
        info["scope"] = fix_scope_tables(doc)  # again ok
    # also fix place date if present in any doc
    if info["place"] == 0:
        info["place"] = fix_doc_place_date(doc)
    doc.save(str(path))
    return info


def main() -> int:
    only = set(sys.argv[1:]) if len(sys.argv) > 1 else None
    folders = sorted(p for p in PRODUCT_SETS.iterdir() if p.is_dir())
    if only:
        folders = [p for p in folders if p.name in only]
    total = {"files": 0, "scope": 0, "place": 0, "pack": 0}
    for i, folder in enumerate(folders, 1):
        for stem in (
            "01_Technical_File.docx",
            "02_EU_DoC.docx",
            "03_Label.docx",
            "04_Shipment_Statement.docx",
        ):
            path = folder / stem
            if not path.exists():
                continue
            info = fix_file(path)
            total["files"] += 1
            total["scope"] += info["scope"]
            total["place"] += info["place"]
            total["pack"] += info["pack"]
        if i % 200 == 0:
            print(f"… {i}/{len(folders)}", flush=True)
    print("DONE", total, flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
