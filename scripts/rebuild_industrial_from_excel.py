"""Rebuild Industrial delivery from Desktop Endustriyel Excel.

Reality: Excel has 2736 industrial PRODUCT rows (not only 3 voltage families).
This script:
1) Builds a proper industrial master from the Excel
2) Generates PILOT packs only (1 simple BOM + 1 full BOM) with premium component-style docs
3) Photo annex from refreshed library
4) LibreOffice PDF only (no Word)

Full 2736 run is a separate explicit step after pilot PASS.
"""

from __future__ import annotations

import json
import re
import shutil
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from photo_annex import append_photo_annex, resolve_photos  # noqa: E402
from convert_pdfs_libreoffice import convert_batch_via_temp, find_soffice  # noqa: E402

SRC = Path(r"C:\Users\burcu\Desktop") / "İNCİ AKÜ PPWR" / "Endustriyel_ambalaj_FINAL_ADET_VE_AGIRLIKLAR.xlsx"
OUT = ROOT / "output"
CAND = OUT / "02_INDUSTRIAL_DELIVERY_REV00_REBUILD"
MASTER = OUT / "INCI_AKU_PPWR_INDUSTRIAL_MASTER_FROM_EXCEL_Rev00.xlsx"
DATE = "11.08.2026"
SIG = ROOT / "assets" / "signatory" / "numan_alver_signature_transparent.png"
NAVY = RGBColor(0x0E, 0x2A, 0x47)
FONT = "Tahoma"


def parse_industrial(src: Path) -> tuple[dict[str, list[dict]], list[dict]]:
    """Return (base_by_voltage, products)."""
    wb = load_workbook(src, data_only=True)
    ws = wb["Sayfa1"]

    base: dict[str, list[dict]] = {"24 V": [], "48 V": [], "80 V": []}
    colmap = {"24 V": 1, "48 V": 5, "80 V": 9}
    for volt, start in colmap.items():
        for r in range(3, 9):  # rows 3-8 materials
            code = ws.cell(r, start).value
            desc = ws.cell(r, start + 1).value
            w = ws.cell(r, start + 2).value
            u = ws.cell(r, start + 3).value
            if not desc:
                continue
            base[volt].append(
                {
                    "component_code": str(code or "").strip(),
                    "description": str(desc).strip(),
                    "qty": 1,
                    "uom": str(u or "").strip(),
                    "unit_weight_g": w,
                    "source": "VOLTAGE_BASE",
                }
            )

    products: list[dict] = []
    for r in range(14, ws.max_row + 1):
        code = ws.cell(r, 1).value
        desc = ws.cell(r, 2).value
        volt = str(ws.cell(r, 3).value or "").strip()
        if code is None or desc is None:
            continue
        try:
            int(str(code))
        except Exception:
            continue
        bom: list[dict] = []
        for c in range(4, 56, 4):
            mcode = ws.cell(r, c).value
            mdesc = ws.cell(r, c + 1).value
            qty = ws.cell(r, c + 2).value
            uom = ws.cell(r, c + 3).value
            if mcode is None and mdesc is None:
                continue
            if str(mdesc or "").strip() in {"", "Tanım"}:
                continue
            bom.append(
                {
                    "component_code": str(mcode or "").strip(),
                    "description": str(mdesc or "").strip(),
                    "qty": qty,
                    "uom": str(uom or "").strip(),
                    "source": "PRODUCT_BOM",
                }
            )
        # merge voltage base (without duplicating same description family)
        vkey = volt if volt in base else None
        if vkey:
            have = {_norm(x["description"]) for x in bom}
            for b in base[vkey]:
                if _norm(b["description"]) not in have:
                    bom.append(dict(b))
        products.append(
            {
                "product_code": str(code).strip(),
                "description": str(desc).strip(),
                "voltage": volt,
                "bom": bom,
                "bom_slots": sum(1 for x in bom if x.get("source") == "PRODUCT_BOM"),
            }
        )
    wb.close()
    return base, products


def _norm(s: str) -> str:
    return (
        (s or "")
        .upper()
        .replace("İ", "I")
        .replace("Ş", "S")
        .replace("Ğ", "G")
        .replace("Ü", "U")
        .replace("Ö", "O")
        .replace("Ç", "C")
    )


def write_master(products: list[dict], path: Path) -> None:
    wb = Workbook()
    home = wb.active
    home.title = "00_HOME"
    home["B2"] = "İNCI AKÜ PPWR — INDUSTRIAL MASTER (FROM EXCEL) Rev00"
    home["B4"] = "PRODUCTS"
    home["C4"] = len(products)
    home["B5"] = "PUBLISH DATE"
    home["C5"] = DATE
    home["B6"] = "SOURCE"
    home["C6"] = SRC.name

    pm = wb.create_sheet("PRODUCT_MASTER")
    pm.append(
        [
            "Product Code",
            "Technical Description",
            "Voltage",
            "BOM Slot Count",
            "Packaging Set Code",
            "Scope Status",
        ]
    )
    bm = wb.create_sheet("BOM_MASTER")
    bm.append(
        [
            "Product Code",
            "Component Code",
            "Component Description",
            "Quantity",
            "UOM",
            "Source",
        ]
    )
    for p in products:
        set_code = f"IND-{p['product_code']}"
        pm.append(
            [
                p["product_code"],
                p["description"],
                p["voltage"],
                p["bom_slots"],
                set_code,
                "EXPORT-READY INDUSTRIAL",
            ]
        )
        for b in p["bom"]:
            bm.append(
                [
                    p["product_code"],
                    b["component_code"],
                    b["description"],
                    b["qty"],
                    b["uom"],
                    b["source"],
                ]
            )
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)


def _set_run(p, text: str, *, bold=False, size=10, color=None) -> None:
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _header(doc: Document, qms: str, title_tr: str, title_en: str, doc_id: str) -> None:
    p = doc.add_paragraph()
    _set_run(p, "İNCİ AKÜ SANAYİ VE TİCARET A.Ş.", bold=True, size=12, color=NAVY)
    p = doc.add_paragraph()
    _set_run(
        p,
        "Manisa OSB 2. Kısım Keçiliköy OSB Mh., Gaziler Cad. No:6, 45030 Yunusemre-Manisa / Türkiye",
        size=8,
    )
    p = doc.add_paragraph()
    _set_run(p, "info@inciaku.com  ·  www.inciaku.com  ·  +90 236 233 25 10", size=8)
    p = doc.add_paragraph()
    _set_run(
        p,
        f"Doküman No/Doc. Nr.: {qms}    Yayın Trh./Rel. Date: {DATE}   Rev.No/Rev.Nr.: 00   Rev.Trh./Rev.Date: - - -",
        bold=True,
        size=9,
    )
    p = doc.add_paragraph()
    _set_run(p, f"{title_tr}\n{title_en}", bold=True, size=13, color=NAVY)
    p = doc.add_paragraph()
    _set_run(p, f"Doküman Kimliği / Document ID: {doc_id}", bold=True, size=9)


def _kv(doc: Document, rows: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v


def build_tf(path: Path, product: dict) -> None:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)
    pid = product["product_code"]
    _header(
        doc,
        "YS/D/0020",
        "AMBALAJ KONFİGÜRASYONU TEKNİK DOSYASI",
        "PACKAGING CONFIGURATION TECHNICAL FILE",
        f"IA-PPWR-TF-IND-{pid}-R00",
    )
    p = doc.add_paragraph()
    _set_run(p, "01  ÜRÜN KAPSAMI / 01  PRODUCT SCOPE", bold=True, size=11, color=NAVY)
    _kv(
        doc,
        [
            ("Product Code", pid),
            ("Product Description", product["description"]),
            ("Voltage", product["voltage"]),
            ("Revision", f"Rev00 / {DATE}"),
            ("Pallet rule", "EURO PALLET (locked)"),
        ],
    )
    p = doc.add_paragraph()
    _set_run(p, "02  FİZİKSEL BOM / 02  PHYSICAL BOM", bold=True, size=11, color=NAVY)
    t = doc.add_table(rows=1 + len(product["bom"]), cols=4)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    hdr = ["Component Code", "Component Description", "Qty", "UOM"]
    for i, h in enumerate(hdr):
        t.rows[0].cells[i].text = h
    for i, b in enumerate(product["bom"], 1):
        t.rows[i].cells[0].text = str(b["component_code"])
        t.rows[i].cells[1].text = str(b["description"])
        t.rows[i].cells[2].text = str(b["qty"] if b["qty"] is not None else "")
        t.rows[i].cells[3].text = str(b["uom"])
    p = doc.add_paragraph()
    _set_run(
        p,
        "03  SONUÇ / 03  CONCLUSION\n"
        f"Bu Teknik Dosya, endüstriyel ürün {pid} için kontrollü ambalaj konfigürasyonunu tanımlar.\n"
        f"This Technical File defines the controlled packaging configuration for industrial product {pid}.",
        size=9,
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_doc(path: Path, product: dict) -> None:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)
    pid = product["product_code"]
    _header(
        doc,
        "YS/D/0021",
        "AB UYGUNLUK BEYANI",
        "EU DECLARATION OF CONFORMITY",
        f"IA-PPWR-DOC-IND-{pid}-R00",
    )
    _kv(
        doc,
        [
            ("Product Code", pid),
            ("Product Description", product["description"]),
            ("Object of declaration", f"Industrial packaging configuration for {pid}"),
            ("Revision / issue date", f"Rev.00 / {DATE}"),
        ],
    )
    p = doc.add_paragraph()
    _set_run(p, "03  İMZA / 03  SIGNATURE", bold=True, size=11, color=NAVY)
    t = doc.add_table(rows=2, cols=2)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    t.rows[0].cells[0].text = "Veriliş yeri ve tarihi\nPlace and date of issue"
    t.rows[0].cells[1].text = f"Manisa / Türkiye  •  {DATE}"
    t.rows[1].cells[0].text = "Ad / görev / imza\nName / function / signature"
    cell = t.rows[1].cells[1]
    cell.text = ""
    p1 = cell.paragraphs[0]
    _set_run(p1, "Ad / Name: Numan Alver", bold=True, size=9)
    p2 = cell.add_paragraph()
    _set_run(p2, "Görev / Function: Operasyon Direktörü / Operations Director", bold=True, size=9)
    p3 = cell.add_paragraph()
    _set_run(p3, "İmza / Signature:", size=9)
    p4 = cell.add_paragraph()
    if SIG.exists():
        p4.add_run().add_picture(str(SIG), width=Cm(4.2))
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_label(path: Path, product: dict) -> None:
    doc = Document()
    pid = product["product_code"]
    _header(doc, "YS/D/0022", "AMBALAJ ETİKETİ", "PACKAGING LABEL", f"IA-PPWR-LBL-IND-{pid}-R00")
    _kv(
        doc,
        [
            ("Product Code", pid),
            ("Product Description", product["description"]),
            ("Packing date", DATE),
            ("Revision", "Rev00"),
        ],
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def build_stm(path: Path, product: dict) -> None:
    doc = Document()
    pid = product["product_code"]
    _header(
        doc,
        "YS/D/0023",
        "SEVKİYAT BEYANI",
        "SHIPMENT STATEMENT",
        f"IA-PPWR-STM-IND-{pid}-R00",
    )
    _kv(
        doc,
        [
            ("Product Code", pid),
            ("Product Description", product["description"]),
            ("Statement date", DATE),
            ("Revision", "Rev00"),
        ],
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def generate_pack(product: dict, dest: Path) -> dict:
    if dest.exists():
        shutil.rmtree(dest)
    dest.mkdir(parents=True)
    build_tf(dest / "01_Technical_File.docx", product)
    build_doc(dest / "02_EU_DoC.docx", product)
    build_label(dest / "03_Label.docx", product)
    build_stm(dest / "04_Shipment_Statement.docx", product)
    photos = resolve_photos(scope="INDUSTRIAL", bom_lines=product["bom"])
    n = append_photo_annex(
        dest / "01_Technical_File.docx",
        photos,
        title_extra=f"Product {product['product_code']}",
    )
    return {"photos": n, "bom_lines": len(product["bom"])}


def main() -> int:
    print("Parse industrial Excel…", flush=True)
    base, products = parse_industrial(SRC)
    print(f"products={len(products)} base_keys={list(base)}", flush=True)
    write_master(products, MASTER)
    print("MASTER", MASTER, flush=True)

    # pick pilots
    simple = next(p for p in products if p["bom_slots"] == 3)
    full = next(p for p in products if p["bom_slots"] >= 10)
    pilots = [simple, full]
    print(
        "PILOTS",
        [(p["product_code"], p["bom_slots"], p["voltage"]) for p in pilots],
        flush=True,
    )

    if CAND.exists():
        shutil.rmtree(CAND)
    control = CAND / "00_CONTROL"
    docs = CAND / "01_PRODUCTS"
    control.mkdir(parents=True)
    docs.mkdir(parents=True)

    shutil.copy2(SRC, control / SRC.name)
    shutil.copy2(MASTER, control / MASTER.name)

    for p in pilots:
        dest = docs / p["product_code"]
        info = generate_pack(p, dest)
        print(f"PACK {p['product_code']} photos={info['photos']} bom={info['bom_lines']}", flush=True)

    # LO PDFs
    files = []
    for p in pilots:
        folder = docs / p["product_code"]
        for stem in ["01_Technical_File", "02_EU_DoC", "03_Label", "04_Shipment_Statement"]:
            files.append(folder / f"{stem}.docx")
    ok, fail = convert_batch_via_temp(find_soffice(), files, OUT / "_lo_profile_pdf", chunk=8)
    print(f"PDF ok={ok} fail={fail}", flush=True)

    # mini engine
    wb = Workbook()
    home = wb.active
    home.title = "00_HOME"
    home["B2"] = "İNCI AKÜ PPWR — INDUSTRIAL ENGINE (PILOT REBUILD)"
    home["B4"] = "TOTAL PRODUCTS IN MASTER"
    home["C4"] = len(products)
    home["B5"] = "PILOT ISSUED"
    home["C5"] = len(pilots)
    home["B6"] = "DATE"
    home["C6"] = DATE
    dc = wb.create_sheet("DOCUMENT_CENTER")
    headers = ["Key", "Description", "TF WORD", "TF PDF", "DoC WORD", "DoC PDF", "Label WORD", "Label PDF", "STM WORD", "STM PDF"]
    for i, h in enumerate(headers, 1):
        cell = dc.cell(4, i, h)
        cell.fill = PatternFill("solid", fgColor="0E2A47")
        cell.font = Font(name="Tahoma", color="FFFFFF", bold=True)
    for i, p in enumerate(pilots, 1):
        rel = f"..\\01_PRODUCTS\\{p['product_code']}\\"
        r = i + 4
        dc.cell(r, 1, p["product_code"])
        dc.cell(r, 2, p["description"])
        for col, fname in [
            (3, "01_Technical_File.docx"),
            (4, "01_Technical_File.pdf"),
            (5, "02_EU_DoC.docx"),
            (6, "02_EU_DoC.pdf"),
            (7, "03_Label.docx"),
            (8, "03_Label.pdf"),
            (9, "04_Shipment_Statement.docx"),
            (10, "04_Shipment_Statement.pdf"),
        ]:
            cell = dc.cell(r, col)
            cell.value = "OPEN WORD" if fname.endswith(".docx") else "OPEN PDF"
            cell.hyperlink = rel + fname
            cell.font = Font(name="Tahoma", color="0563C1", underline="single")
    eng = control / "INCI_AKU_PPWR_INDUSTRIAL_ENGINE_Rev00.xlsx"
    wb.save(eng)
    (CAND / "00_AC_DOCUMENT_ENGINE.cmd").write_text(
        "@echo off\r\nstart \"\" \"%~dp000_CONTROL\\INCI_AKU_PPWR_INDUSTRIAL_ENGINE_Rev00.xlsx\"\r\n",
        encoding="utf-8",
    )

    report = {
        "master_products": len(products),
        "pilots": [p["product_code"] for p in pilots],
        "pdf_ok": ok,
        "pdf_fail": fail,
        "candidate": str(CAND),
        "master": str(MASTER),
        "NOTE": "Full 2736 production NOT run yet — pilot first",
    }
    (control / "QA_INDUSTRIAL_PILOT.json").write_text(
        json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    print(json.dumps(report, indent=2), flush=True)
    print("INDUSTRIAL PILOT READY — review before full 2736", flush=True)
    return 0 if fail == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
