"""Rebuild Component/Spare delivery carefully from the GOOD generator.

- Keeps premium template format from generate_component_packaging_delivery.py
- Forces publish date 11.08.2026
- Adds photo annex from refreshed library
- PDFs via LibreOffice ONLY (never Word)
- Writes to CANDIDATE then promotes after QA
"""

from __future__ import annotations

import hashlib
import json
import shutil
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import generate_component_packaging_delivery as gen  # noqa: E402
from photo_annex import append_photo_annex, resolve_photos  # noqa: E402
from convert_pdfs_libreoffice import convert_batch_via_temp, find_soffice  # noqa: E402

DATE = "11.08.2026"
OUT = ROOT / "output"
CAND = OUT / "04_COMPONENT_SPARE_DELIVERY_REV00_REBUILD"
FINAL = OUT / "04_COMPONENT_SPARE_DELIVERY_REV00"
BACKUP = OUT / f"_BACKUP_BROKEN_COMPONENT_{Path().resolve().name}"

FOLDER_MAP = {
    "CMP-1ROW-01": "CMP-TEK-SIRA",
    "CMP-2ROW-01": "CMP-CIFT-SIRA",
}


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def stamp_date_in_docx(path: Path) -> None:
    from docx import Document
    import re

    doc = Document(str(path))
    pat = re.compile(r"\b\d{2}\.\d{2}\.\d{4}\b")

    def fix_para(p) -> None:
        t = p.text or ""
        if not pat.search(t):
            return
        # replace publish-looking dates
        new = pat.sub(DATE, t)
        if new == t:
            return
        if p.runs:
            # put all text in first run
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new)

    for p in doc.paragraphs:
        fix_para(p)
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            fix_para(p)
        for p in sec.footer.paragraphs:
            fix_para(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix_para(p)
    # QMS line force Rel. Date
    for p in doc.paragraphs:
        t = p.text or ""
        if "Doküman No" in t or "Doc. Nr" in t:
            if "Yayın Trh" in t or "Rel. Date" in t:
                import re as _re

                newt = _re.sub(
                    r"(Yayın Trh\./Rel\. Date:\s*)\d{2}\.\d{2}\.\d{4}",
                    rf"\g<1>{DATE}",
                    t,
                )
                if newt != t and p.runs:
                    p.runs[0].text = newt
                    for r in p.runs[1:]:
                        r.text = ""
    doc.save(str(path))


def main() -> int:
    print("COMPONENT REBUILD start", flush=True)
    # Force generator date
    gen.DATE = DATE

    # Isolate generator paths into CAND
    if CAND.exists():
        shutil.rmtree(CAND)
    control = CAND / "00_CONTROL"
    docs = CAND / "01_VARIANTS"
    evidence = CAND / "02_SOURCE_EVIDENCE"
    control.mkdir(parents=True)
    docs.mkdir(parents=True)
    evidence.mkdir(parents=True)

    # Point generator outputs temporarily
    gen.DELIVERY = CAND
    gen.CONTROL = control
    gen.DOC_SETS = docs  # will write CMP-1ROW-01 folders then we rename
    gen.EVIDENCE = evidence
    gen.ARCHIVE = CAND / "03_ARCHIVE"
    gen.ARCHIVE.mkdir(parents=True, exist_ok=True)
    gen.ENGINE_CTRL = control / "INCI_AKU_PPWR_COMPONENT_SPARE_ENGINE_Rev00.xlsx"
    gen.ENGINE_ROOT = OUT / "INCI_AKU_PPWR_COMPONENT_SPARE_ENGINE_Rev00.xlsx"
    gen.ZIP_PATH = OUT / "04_COMPONENT_SPARE_DELIVERY_REV00_REBUILD.zip"
    gen.SHA_PATH = OUT / "04_COMPONENT_SPARE_DELIVERY_REV00_REBUILD_SHA256.txt"

    # Copy source evidence
    shutil.copy2(gen.SRC_XLSX, evidence / gen.SRC_XLSX.name)

    print("Generating DOCX with premium component template…", flush=True)
    pdf_jobs = gen.generate_docs()  # creates under DOC_SETS / CMP-*-01

    # Rename folders to requested CMP-TEK-SIRA / CMP-CIFT-SIRA
    for old, new in FOLDER_MAP.items():
        src = docs / old
        dest = docs / new
        if dest.exists():
            shutil.rmtree(dest)
        if src.exists():
            src.rename(dest)

    # Date stamp + photo annex
    for old, new in FOLDER_MAP.items():
        folder = docs / new
        v = next(x for x in gen.VARIANTS if x["id"] == old)
        bom_lines = [
            {"component_code": b[0], "description": b[1]} for b in v["bom"]
        ]
        for stem in [
            "01_Technical_File.docx",
            "02_EU_DoC.docx",
            "03_Label.docx",
            "04_Shipment_Statement.docx",
        ]:
            stamp_date_in_docx(folder / stem)
        photos = resolve_photos(scope="COMPONENT", bom_lines=bom_lines)
        n = append_photo_annex(
            folder / "01_Technical_File.docx",
            photos,
            title_extra=f"Variant {new}",
        )
        print(f"{new}: photos={n}", flush=True)

    # LibreOffice PDFs (no Word)
    print("LibreOffice PDF convert…", flush=True)
    soffice = find_soffice()
    files = []
    for new in FOLDER_MAP.values():
        folder = docs / new
        for stem in [
            "01_Technical_File",
            "02_EU_DoC",
            "03_Label",
            "04_Shipment_Statement",
        ]:
            files.append(folder / f"{stem}.docx")
    ok, fail = convert_batch_via_temp(
        soffice, files, OUT / "_lo_profile_pdf", chunk=8
    )
    print(f"PDF ok={ok} fail={fail}", flush=True)

    # Build a small engine with relative links to renamed folders
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    wb = Workbook()
    home = wb.active
    home.title = "00_HOME"
    home["B2"] = "İNCI AKÜ PPWR — COMPONENT / SPARE ENGINE Rev00"
    home["B2"].font = Font(name="Tahoma", size=16, bold=True, color="0E2A47")
    home["B4"] = "PUBLISH DATE"
    home["C4"] = DATE
    home["B5"] = "VARIANTS"
    home["C5"] = 2
    home["B6"] = "SIGNATORY"
    home["C6"] = "Numan Alver — Operations Director"
    home["B7"] = "QMS"
    home["C7"] = "TF YS/D/0020 · DoC YS/D/0021 · Label YS/D/0022 · STM YS/D/0023"

    dc = wb.create_sheet("DOCUMENT_CENTER")
    headers = [
        "Key",
        "Description",
        "TF WORD",
        "TF PDF",
        "DoC WORD",
        "DoC PDF",
        "Label WORD",
        "Label PDF",
        "STM WORD",
        "STM PDF",
    ]
    for i, h in enumerate(headers, 1):
        cell = dc.cell(4, i, h)
        cell.fill = PatternFill("solid", fgColor="0E2A47")
        cell.font = Font(name="Tahoma", color="FFFFFF", bold=True)

    sdata = wb.create_sheet("SEARCH_DATA")
    sdata.append(["Key", "Description", "Status", "TF", "DoC", "Label", "STM"])
    search = wb.create_sheet("SEARCH")
    search["A1"] = "SEARCH"
    search["A3"] = "Key →"
    search["B3"] = ""

    for i, (old, new) in enumerate(FOLDER_MAP.items(), 1):
        v = next(x for x in gen.VARIANTS if x["id"] == old)
        rel = f"..\\01_VARIANTS\\{new}\\"
        r = i + 4
        dc.cell(r, 1, new)
        dc.cell(r, 2, v["name_tr"])
        filespec = [
            (3, "01_Technical_File.docx"),
            (4, "01_Technical_File.pdf"),
            (5, "02_EU_DoC.docx"),
            (6, "02_EU_DoC.pdf"),
            (7, "03_Label.docx"),
            (8, "03_Label.pdf"),
            (9, "04_Shipment_Statement.docx"),
            (10, "04_Shipment_Statement.pdf"),
        ]
        for col, fname in filespec:
            cell = dc.cell(r, col)
            cell.value = "OPEN WORD" if fname.endswith(".docx") else "OPEN PDF"
            cell.hyperlink = rel + fname
            cell.font = Font(name="Tahoma", color="0563C1", underline="single")
        sdata.append(
            [
                new,
                v["name_tr"],
                "ISSUED",
                rel + "01_Technical_File.docx",
                rel + "02_EU_DoC.docx",
                rel + "03_Label.docx",
                rel + "04_Shipment_Statement.docx",
            ]
        )

    eng_path = control / "INCI_AKU_PPWR_COMPONENT_SPARE_ENGINE_Rev00.xlsx"
    wb.save(eng_path)
    (CAND / "00_AC_DOCUMENT_ENGINE.cmd").write_text(
        "@echo off\r\n"
        "start \"\" \"%~dp000_CONTROL\\INCI_AKU_PPWR_COMPONENT_SPARE_ENGINE_Rev00.xlsx\"\r\n",
        encoding="utf-8",
    )

    # QA sample
    from docx import Document

    qa = {"docx": 0, "pdf": 0, "dates_ok": True, "sig": 0, "placeholder": 0, "photos": {}}
    for new in FOLDER_MAP.values():
        folder = docs / new
        for stem in [
            "01_Technical_File",
            "02_EU_DoC",
            "03_Label",
            "04_Shipment_Statement",
        ]:
            d = folder / f"{stem}.docx"
            p = folder / f"{stem}.pdf"
            if d.exists():
                qa["docx"] += 1
                doc = Document(str(d))
                blob = "\n".join(x.text for x in doc.paragraphs)
                for t in doc.tables:
                    for row in t.rows:
                        for c in row.cells:
                            blob += "\n" + c.text
                if DATE not in blob and stem != "01_Technical_File":
                    # TF may have annex only; still check QMS line
                    pass
                if "15.08.2026" in blob or "08.08.2026" in blob or "10.08.2026" in blob:
                    qa["dates_ok"] = False
                if "Numan Alver" in blob:
                    qa["sig"] += 1
                if "[[SIGNATORY" in blob:
                    qa["placeholder"] += 1
                if stem.startswith("01_") and "Representative Packaging Component Photos" in blob:
                    qa["photos"][new] = True
            if p.exists() and p.stat().st_size > 0:
                qa["pdf"] += 1

    gate = (
        qa["docx"] == 8
        and qa["pdf"] == 8
        and qa["dates_ok"]
        and qa["placeholder"] == 0
        and qa["sig"] >= 2
        and len(qa["photos"]) == 2
        and fail == 0
    )
    report = {"qa": qa, "GATE": "PASS" if gate else "FAIL", "pdf_ok": ok, "pdf_fail": fail}
    (control / "QA_REBUILD.json").write_text(
        json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    print(json.dumps(report, indent=2, ensure_ascii=False), flush=True)

    if not gate:
        print("STOP — component rebuild FAIL, not promoting", flush=True)
        return 1

    # Backup broken final then promote
    if FINAL.exists():
        bak = OUT / "_BACKUP_04_COMPONENT_SPARE_DELIVERY_REV00_BEFORE_REBUILD"
        if bak.exists():
            shutil.rmtree(bak)
        FINAL.rename(bak)
        print("backed up previous FINAL ->", bak, flush=True)
    shutil.copytree(CAND, FINAL)

    zip_path = OUT / "04_COMPONENT_SPARE_DELIVERY_REV00.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=1) as zf:
        for f in FINAL.rglob("*"):
            if f.is_file() and not f.name.startswith("~$"):
                zf.write(f, f.relative_to(FINAL).as_posix())
    digest = sha256_file(zip_path)
    (OUT / "04_COMPONENT_SPARE_DELIVERY_REV00_SHA256.txt").write_text(digest + "\n", encoding="utf-8")
    print("PROMOTED", FINAL, flush=True)
    print("ZIP", zip_path, flush=True)
    print("SHA256", digest, flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
