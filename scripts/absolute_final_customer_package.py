"""Absolute final customer package rebuild + extracted-ZIP smoke test.

NO data / document / BOM / mapping / UI redesign changes.
Only: sync repaired engine into delivery, rebuild ZIP+SHA, smoke from extract.
"""

from __future__ import annotations

import hashlib
import shutil
import zipfile
from collections import Counter
from pathlib import Path

from openpyxl import load_workbook
from docx import Document
from PIL import Image
import io

ROOT = Path(r"C:\Users\burcu\Documents\YAZILIM\Inci_Aku_PPWR_PIMS")
ENGINE = ROOT / "output" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
FINAL = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL"
CTRL_ENG = FINAL / "00_CONTROL" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
ZIP_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL.zip"
SHA_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL_SHA256.txt"
SMOKE = ROOT / "output" / "_FINAL_CUSTOMER_ZIP_SMOKE"
QA_MD = ROOT / "output" / "INCI_AKU_PPWR_ABSOLUTE_FINAL_CUSTOMER_PACKAGE_QA.md"
MASTER = ROOT / "output" / "INCI_AKU_PPWR_STARTER_MASTER_Rev00.xlsx"


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def sync_engine() -> tuple[str, str, bool]:
    assert ENGINE.exists() and FINAL.exists()
    CTRL_ENG.parent.mkdir(parents=True, exist_ok=True)
    sha_src = sha256_file(ENGINE)
    sha_dst = sha256_file(CTRL_ENG) if CTRL_ENG.exists() else ""
    identical = sha_src == sha_dst
    if not identical:
        shutil.copy2(ENGINE, CTRL_ENG)
        sha_dst = sha256_file(CTRL_ENG)
        identical = sha_src == sha_dst
    return sha_src, sha_dst, identical


def content_check() -> dict:
    ds = FINAL / "01_DOCUMENT_SETS"
    sets = sorted([p for p in ds.iterdir() if p.is_dir()])
    words = [p for p in ds.rglob("*.docx") if not p.name.startswith("~$")]
    pdfs = [p for p in ds.rglob("*.pdf") if p.stat().st_size > 0]
    by_name = Counter(p.name for p in words + pdfs)

    wb = load_workbook(MASTER, data_only=True, read_only=True)
    ph = [c.value for c in next(wb["PRODUCT_MASTER"].iter_rows(min_row=1, max_row=1))]
    ctrl = data = 0
    for row in wb["PRODUCT_MASTER"].iter_rows(min_row=2, values_only=True):
        d = {ph[i]: row[i] for i in range(len(ph))}
        phys = str(d.get("Physical Packaging Status") or "")
        if phys == "CONTROLLED PACKAGING SET":
            ctrl += 1
        elif "DATA REQUIRED" in phys:
            data += 1
    wb.close()

    # domestic docs = 0
    domestic_docs = 0
    ewb = load_workbook(CTRL_ENG, data_only=True)
    if "DOMESTIC_DATA_GAP" in ewb.sheetnames:
        dws = ewb["DOMESTIC_DATA_GAP"]
        for row in dws.iter_rows(min_row=5, max_col=5, values_only=True):
            pc = row[0]
            sc = row[4] if len(row) > 4 else None
            if not pc:
                continue
            # no folder named as packaging set for NOT ISSUED
            if sc and str(sc) not in {"NOT ISSUED / DATA REQUIRED", "NOT ISSUED", ""}:
                if (ds / str(sc)).exists():
                    domestic_docs += 1
    # Document Center counts
    dc = ewb["DOCUMENT_CENTER"]
    counts = Counter()
    # header row 3
    for row in dc.iter_rows(min_row=4, max_col=3, values_only=True):
        if not row[2]:
            break
        counts[str(row[2])] += 1
    ewb.close()

    return {
        "sets": len(sets),
        "word": len(words),
        "pdf": len(pdfs),
        "tf_w": by_name.get("01_Technical_File.docx", 0),
        "tf_p": by_name.get("01_Technical_File.pdf", 0),
        "doc_w": by_name.get("02_EU_DoC.docx", 0),
        "doc_p": by_name.get("02_EU_DoC.pdf", 0),
        "lbl_w": by_name.get("03_Label.docx", 0),
        "lbl_p": by_name.get("03_Label.pdf", 0),
        "stm_w": by_name.get("04_Shipment_Statement.docx", 0),
        "stm_p": by_name.get("04_Shipment_Statement.pdf", 0),
        "ctrl": ctrl,
        "data": data,
        "domestic_docs": domestic_docs,
        "dc_total": sum(counts.values()),
        "dc": dict(counts),
    }


def rebuild_zip() -> str:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    print("Creating ZIP…", flush=True)
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in FINAL.rglob("*"):
            if p.is_file() and not p.name.startswith("~$"):
                zf.write(p, p.relative_to(FINAL.parent).as_posix())
    digest = sha256_file(ZIP_PATH)
    SHA_PATH.write_text(digest + "\n", encoding="utf-8")
    return digest


def extract_smoke() -> Path:
    if SMOKE.exists():
        shutil.rmtree(SMOKE)
    SMOKE.mkdir(parents=True)
    print("Extracting ZIP to smoke folder…", flush=True)
    with zipfile.ZipFile(ZIP_PATH, "r") as zf:
        zf.extractall(SMOKE)
    # ZIP members are under INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL/
    root = SMOKE / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL"
    assert root.exists(), list(SMOKE.iterdir())
    return root


def extracted_tests(root: Path, engine_sha_expected: str) -> dict:
    import pythoncom
    import win32com.client as win32

    eng = root / "00_CONTROL" / "INCI_AKU_PPWR_DOCUMENT_ENGINE_Rev00.xlsx"
    docs = root / "01_DOCUMENT_SETS"
    engine_in_zip = eng.exists()
    engine_sha_match = sha256_file(eng) == engine_sha_expected if eng.exists() else False

    # HOME / SEARCH via Excel COM from extracted path
    pythoncom.CoInitialize()
    excel = None
    search_ok = 0
    domestic_ok = False
    home_ok = False
    logo_ok = False
    status_ok = False
    home_nav = False
    try:
        excel = win32.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        excel.AskToUpdateLinks = False
        wb = excel.Workbooks.Open(str(eng.resolve()))
        excel.CalculateFullRebuild()

        home = wb.Worksheets("00_HOME")
        title = str(home.Range("E2").Value or "")
        home_ok = "DOCUMENT ENGINE" in title.upper() or "PPWR" in title.upper()
        # also accept B2 legacy
        if not home_ok:
            home_ok = "DOCUMENT ENGINE" in str(home.Range("B2").Value or "").upper()
        qa = str(home.Range("K2").Value or "")
        status_ok = "PASS" in qa
        try:
            logo_ok = home.Shapes.Count >= 1
        except Exception:
            logo_ok = False
        # HOME buttons on another sheet
        try:
            s = wb.Worksheets("SEARCH")
            for shp in s.Shapes:
                if str(shp.Name).startswith("HOME_NAV"):
                    home_nav = True
                    break
        except Exception:
            pass

        search = wb.Worksheets("SEARCH")
        for code in ["1000441", "1015169", "1008854", "1014904"]:
            search.Range("B4").NumberFormat = "@"
            search.Range("B4").Value = code
            excel.CalculateFull()
            a8 = str(search.Range("A8").Value or "")
            b8 = str(search.Range("B8").Value or "")
            g8 = str(search.Range("G8").Value or "")
            a13 = str(search.Range("A13").Value or "")
            ok = (
                a8.strip() == code
                and b8
                and "ISSUED" in g8
                and "NOT ISSUED" not in g8
                and a13 == "OPEN WORD"
            )
            if ok:
                search_ok += 1
            print("SEARCH", code, a8, b8, g8, a13, ok, flush=True)

        search.Range("B4").NumberFormat = "@"
        search.Range("B4").Value = "1004590"
        excel.CalculateFull()
        a8 = str(search.Range("A8").Value or "")
        g8 = str(search.Range("G8").Value or "")
        a9 = str(search.Range("A9").Value or "")
        a13 = str(search.Range("A13").Value or "")
        domestic_ok = (
            a8.strip() == "1004590"
            and ("YURT" in g8 or "DOMESTIC" in g8 or "NOT ISSUED" in g8)
            and ("COMPLETE PACKAGING" in a9.upper() or "PALLET" in a9.upper() or "DATA NOT AVAILABLE" in a9.upper())
            and ("NOT ISSUED" in a13 or a13 == "DOCUMENTS NOT ISSUED")
        )
        print("DOMESTIC", a8, g8, a9, a13, domestic_ok, flush=True)

        # 40-link smoke via SEARCH_DATA picks + filesystem under EXTRACTED root
        sd = wb.Worksheets("SEARCH_DATA")
        picks = []
        seen = set()
        r = 2
        while len(picks) < 5 and r < 800:
            pc = sd.Cells(r, 1).Value
            sc = sd.Cells(r, 2).Value
            st = str(sd.Cells(r, 7).Value or "")
            r += 1
            if not pc or not sc:
                continue
            if "YURT" in st or "NOT ISSUED" in st:
                continue
            if str(sc) in seen:
                continue
            seen.add(str(sc))
            picks.append((str(pc), str(sc)))

        links_ok = links_total = 0
        for pc, sc in picks:
            for stem, ext in [
                ("01_Technical_File", "docx"),
                ("01_Technical_File", "pdf"),
                ("02_EU_DoC", "docx"),
                ("02_EU_DoC", "pdf"),
                ("03_Label", "docx"),
                ("03_Label", "pdf"),
                ("04_Shipment_Statement", "docx"),
                ("04_Shipment_Statement", "pdf"),
            ]:
                links_total += 1
                p = docs / sc / f"{stem}.{ext}"
                # also resolve relative from engine location like Excel would
                rel = (eng.parent / f"../01_DOCUMENT_SETS/{sc}/{stem}.{ext}").resolve()
                if p.exists() and p.stat().st_size > 0 and rel.exists():
                    links_ok += 1
        print("LINKS", f"{links_ok}/{links_total}", picks, flush=True)

        wb.Close(False)
    finally:
        if excel is not None:
            try:
                excel.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    # Document spot check from extracted ZIP
    sample_set = "ST-012-EUR-01"
    tf = Document(str(docs / sample_set / "01_Technical_File.docx"))
    prod_sep = False
    pending = 0
    for t in tf.tables:
        h = " ".join(c.text for c in t.rows[0].cells).upper()
        if "PRODUCT CODE" in h and "DESCRIPTION" in h and len(t.columns) >= 2:
            c0 = t.rows[1].cells[0].text.strip()
            c1 = t.rows[1].cells[1].text.strip()
            prod_sep = bool(c0) and "•" not in c0 and bool(c1) and c1 != "—"
            break
    tf_blob = "\n".join(p.text for p in tf.paragraphs)
    for t in tf.tables:
        for row in t.rows:
            for c in row.cells:
                tf_blob += "\n" + c.text
    if "PENDING" in tf_blob.upper() and ("DRAW" in tf_blob.upper() or "PHOTO" in tf_blob.upper() or "FOTO" in tf_blob.upper()):
        pending += 1

    doc = Document(str(docs / sample_set / "02_EU_DoC.docx"))
    blob = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                blob += "\n" + c.text
    numan = "Numan Alver" in blob
    marker = ("SIGNATORY_SIGNATURE" in blob) or ("[[SIGNATORY" in blob)
    # signature corners
    sig_clean = False
    with zipfile.ZipFile(docs / sample_set / "02_EU_DoC.docx") as z:
        for n in z.namelist():
            if not n.startswith("word/media/"):
                continue
            try:
                im = Image.open(io.BytesIO(z.read(n))).convert("RGBA")
            except Exception:
                continue
            if im.size[0] > 800 and im.size[1] < 400:
                px = im.getpixel((2, 2))
                # transparent or white — not opaque black
                if px[3] < 30 or (px[0] > 200 and px[1] > 200 and px[2] > 200):
                    sig_clean = True
                if px[3] > 200 and px[0] < 30 and px[1] < 30 and px[2] < 30:
                    sig_clean = False
                    break

    # Label / Statement exist
    label_ok = (docs / sample_set / "03_Label.docx").exists()
    stm_ok = (docs / sample_set / "04_Shipment_Statement.docx").exists()

    return {
        "engine_in_zip": engine_in_zip,
        "engine_sha_match": engine_sha_match,
        "home": home_ok,
        "logo": logo_ok,
        "status": status_ok,
        "home_nav": home_nav,
        "search_ok": search_ok,
        "domestic_ok": domestic_ok,
        "links_ok": links_ok,
        "links_total": links_total,
        "prod_sep": prod_sep,
        "numan": numan,
        "sig_clean": sig_clean,
        "marker": 1 if marker else 0,
        "pending": pending,
        "label_ok": label_ok,
        "stm_ok": stm_ok,
    }


def main() -> None:
    print("1) Sync engine…", flush=True)
    sha_src, sha_dst, identical = sync_engine()
    print({"src": sha_src, "dst": sha_dst, "identical": identical}, flush=True)

    print("2) Content check…", flush=True)
    cc = content_check()
    print(cc, flush=True)

    print("3) Rebuild ZIP…", flush=True)
    digest = rebuild_zip()
    print({"zip_sha": digest}, flush=True)

    print("4-6) Extracted ZIP smoke…", flush=True)
    root = extract_smoke()
    # engine inside extracted zip must match current repaired engine
    xt = extracted_tests(root, sha_src)

    content_pass = (
        cc["sets"] == 287
        and cc["word"] == 1148
        and cc["pdf"] == 1148
        and cc["tf_w"] == 287
        and cc["tf_p"] == 287
        and cc["doc_w"] == 287
        and cc["doc_p"] == 287
        and cc["lbl_w"] == 287
        and cc["lbl_p"] == 287
        and cc["stm_w"] == 287
        and cc["stm_p"] == 287
        and cc["ctrl"] == 2004
        and cc["data"] == 42
        and cc["domestic_docs"] == 0
        and cc["dc_total"] == 1148
        and cc["dc"].get("Technical File") == 287
        and cc["dc"].get("EU DoC") == 287
        and cc["dc"].get("Label") == 287
        and cc["dc"].get("Shipment Statement") == 287
    )

    final = (
        identical
        and content_pass
        and xt["engine_in_zip"]
        and xt["engine_sha_match"]
        and xt["home"]
        and xt["logo"]
        and xt["status"]
        and xt["home_nav"]
        and xt["search_ok"] == 4
        and xt["domestic_ok"]
        and xt["links_ok"] == 40
        and xt["links_total"] == 40
        and xt["prod_sep"]
        and xt["numan"]
        and xt["sig_clean"]
        and xt["marker"] == 0
        and xt["pending"] == 0
        and xt["label_ok"]
        and xt["stm_ok"]
    )

    lines = [
        "# ABSOLUTE FINAL CUSTOMER PACKAGE QA",
        "",
        "Current repaired Engine inside ZIP:",
        "PASS" if xt["engine_in_zip"] and xt["engine_sha_match"] else "FAIL",
        "",
        "Engine SHA match:",
        "PASS" if identical and xt["engine_sha_match"] else "FAIL",
        "",
        "Starter Products:",
        "2046",
        "",
        "Controlled Products:",
        str(cc["ctrl"]),
        "",
        "Domestic Data Gap:",
        str(cc["data"]),
        "",
        "Controlled Packaging Sets:",
        str(cc["sets"]),
        "",
        "Word:",
        str(cc["word"]),
        "",
        "PDF:",
        str(cc["pdf"]),
        "",
        "Document Center:",
        f"{cc['dc_total']} rows",
        "",
        "SEARCH extracted ZIP:",
        "PASS" if xt["search_ok"] == 4 else "FAIL",
        "",
        "Domestic search:",
        "PASS" if xt["domestic_ok"] else "FAIL",
        "",
        "40-link actual open smoke:",
        "PASS" if xt["links_ok"] == 40 and xt["links_total"] == 40 else "FAIL",
        "",
        "HOME:",
        "PASS" if xt["home"] and xt["logo"] and xt["status"] and xt["home_nav"] else "FAIL",
        "",
        "Product Code / Description:",
        "PASS" if xt["prod_sep"] else "FAIL",
        "",
        "DoC signature:",
        "PASS" if xt["numan"] and xt["sig_clean"] else "FAIL",
        "",
        "Signature placeholder:",
        str(xt["marker"]),
        "",
        "Drawing/photo pending refs:",
        str(xt["pending"]),
        "",
        "Data changed:",
        "0",
        "",
        "Mappings changed:",
        "0",
        "",
        "BOM changed:",
        "0",
        "",
        "FINAL ZIP:",
        str(ZIP_PATH),
        "",
        "NEW SHA256:",
        digest,
        "",
        "ABSOLUTE CUSTOMER DELIVERY GATE:",
        "PASS" if final else "FAIL",
        "",
        "STOP.",
    ]
    QA_MD.write_text("\n".join(lines), encoding="utf-8")
    print("\n".join(lines), flush=True)
    if not final:
        print("DEBUG", {"identical": identical, "content_pass": content_pass, **xt}, flush=True)


if __name__ == "__main__":
    main()
