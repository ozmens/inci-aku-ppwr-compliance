"""Fix all EU DoC Word files:
- Place/date of issue: Manisa / Türkiye • 10.08.2026
- Name + Function inside signature table cell (not body paragraphs below)
- Signature image stays in the same cell under İmza
- Normalize dates to 10.08.2026
"""

from __future__ import annotations

import re
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(r"C:\Users\burcu\Documents\YAZILIM\Inci_Aku_PPWR_PIMS")
DOC_SETS = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL" / "01_DOCUMENT_SETS"
SIG = ROOT / "assets" / "signatory" / "numan_alver_signature_transparent.png"
DATE = "10.08.2026"
NAME_LINE = "Ad / Name: Numan Alver"
FUNC_LINE = "Görev / Function: Operasyon Direktörü / Operations Director"
PLACE_LINE = f"Manisa / Türkiye  •  {DATE}"

DATE_PAT = re.compile(r"\b\d{2}\.\d{2}\.\d{4}\b")


def _set_paragraph_text(paragraph, text: str, *, bold: bool = False) -> None:
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].bold = bold
        paragraph.runs[0].font.name = "Tahoma"
        paragraph.runs[0].font.size = Pt(9)
    else:
        r = paragraph.add_run(text)
        r.bold = bold
        r.font.name = "Tahoma"
        r.font.size = Pt(9)


def _para_has_picture(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='drawing' or local-name()='pict']"))


def _clear_cell_keep_structure(cell) -> None:
    # remove all paragraphs except we'll rebuild
    for p in list(cell.paragraphs):
        el = p._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _normalize_dates_in_text(text: str) -> str:
    if not text:
        return text
    # Replace known publish/rev dates with controlled date
    return DATE_PAT.sub(DATE, text)


def _rewrite_runs_dates(paragraph) -> bool:
    changed = False
    for r in paragraph.runs:
        if r.text and DATE_PAT.search(r.text):
            new = DATE_PAT.sub(DATE, r.text)
            if new != r.text:
                r.text = new
                changed = True
    # if dates split across runs, fix full paragraph text once
    full = paragraph.text or ""
    if DATE_PAT.search(full) and DATE not in full.replace(DATE, ""):
        # still has other dates
        pass
    if DATE_PAT.search(full):
        norm = DATE_PAT.sub(DATE, full)
        if norm != full and not _para_has_picture(paragraph):
            _set_paragraph_text(paragraph, norm)
            changed = True
    return changed


def fix_doc(path: Path) -> dict:
    doc = Document(str(path))
    info = {"date_cells": 0, "sig_cell": 0, "body_scrub": 0, "qms": 0, "rev": 0}

    # Body paragraphs: normalize dates + remove orphan name/function lines
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        if t in (NAME_LINE, FUNC_LINE) or t.startswith("Ad / Name: Numan") or (
            t.startswith("Görev / Function:") and "Operasyon" in t
        ):
            el = p._p
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
            info["body_scrub"] += 1
            continue
        if _rewrite_runs_dates(p):
            info["qms"] += 1
        # specifically QMS line
        if "Doküman No" in t and "YS/D/0021" in t:
            fixed = DATE_PAT.sub(DATE, t)
            if "Yayın Trh" in fixed or "Rel. Date" in fixed:
                _set_paragraph_text(p, fixed)
                info["qms"] += 1

    # Tables
    for table in doc.tables:
        for row in table.rows:
            left = row.cells[0].text if row.cells else ""
            right = row.cells[-1] if len(row.cells) > 1 else None
            blob_left = left.upper()

            # Place and date row
            if right is not None and (
                "VERİLİŞ YERİ" in blob_left
                or "PLACE AND DATE" in blob_left
                or "VERILIS YERI" in blob_left
            ):
                _clear_cell_keep_structure(right)
                p = right.add_paragraph()
                _set_paragraph_text(p, PLACE_LINE, bold=True)
                info["date_cells"] += 1

            # Signature name/function/signature row
            if right is not None and (
                ("AD / GÖREV" in blob_left or "AD / GOREV" in blob_left or "NAME / FUNCTION" in blob_left)
                and ("İMZA" in blob_left or "IMZA" in blob_left or "SIGNATURE" in blob_left)
            ):
                _clear_cell_keep_structure(right)
                p1 = right.add_paragraph()
                _set_paragraph_text(p1, NAME_LINE, bold=True)
                p2 = right.add_paragraph()
                _set_paragraph_text(p2, FUNC_LINE, bold=True)
                p3 = right.add_paragraph()
                _set_paragraph_text(p3, "İmza / Signature:")
                p4 = right.add_paragraph()
                run = p4.add_run()
                if SIG.exists():
                    run.add_picture(str(SIG), width=Cm(4.2))
                info["sig_cell"] += 1

            # Revision / issue date row
            if right is not None and ("REVİZYON" in blob_left or "REVISION" in blob_left) and (
                "YAYIN" in blob_left or "ISSUE" in blob_left or "TARİH" in blob_left or "DATE" in blob_left
            ):
                # Keep Rev.00 / DATE
                _clear_cell_keep_structure(right)
                p = right.add_paragraph()
                _set_paragraph_text(p, f"Rev.00 / {DATE}", bold=True)
                info["rev"] += 1

            # Generic date normalize in all cells (text only)
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _para_has_picture(p):
                        continue
                    if _rewrite_runs_dates(p):
                        info["qms"] += 1

    # Ensure no leftover body name lines again
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith("Ad / Name: Numan") or (
            t.startswith("Görev / Function:") and "Operasyon" in t
        ):
            el = p._p
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
            info["body_scrub"] += 1

    doc.save(str(path))
    return info


def main() -> None:
    docs = sorted(DOC_SETS.rglob("02_EU_DoC.docx"))
    assert len(docs) == 287, len(docs)
    print(f"Fixing {len(docs)} DoCs…", flush=True)
    totals = {"date_cells": 0, "sig_cell": 0, "body_scrub": 0, "qms": 0, "rev": 0}
    for i, p in enumerate(docs, 1):
        info = fix_doc(p)
        for k, v in info.items():
            totals[k] += v
        if i % 50 == 0:
            print(f"  {i}/{len(docs)}", flush=True)
    print("totals", totals, flush=True)

    # Verify sample with 1008854
    sample = DOC_SETS / "ST-012-EUR-02" / "02_EU_DoC.docx"
    d = Document(str(sample))
    print("VERIFY paras with Numan:", [p.text for p in d.paragraphs if "Numan" in p.text or "Görev / Function" in p.text])
    for ti, t in enumerate(d.tables):
        for row in t.rows:
            left = row.cells[0].text
            if "Veriliş" in left or "Place and date" in left or "Ad / görev" in left or "Name / function" in left or "Revizyon" in left:
                print("ROW", left[:40].replace("\n", " "), "=>", row.cells[-1].text[:120].replace("\n", " | "))

    # PDF render DoCs only via temp (Word COM once per batch)
    import sys

    sys.path.insert(0, str(ROOT))
    sys.path.insert(0, str(ROOT / "src"))
    from builders.phase_i.render_batch import render_docx_batch

    jobs = [(p, p.with_suffix(".pdf")) for p in docs]
    print(f"PDF {len(jobs)}…", flush=True)
    for i in range(0, len(jobs), 40):
        chunk = jobs[i : i + 40]
        tmp = Path(tempfile.mkdtemp())
        mapped = []
        for j, (dpath, pdf) in enumerate(chunk):
            td = tmp / f"{j}.docx"
            tp = tmp / f"{j}.pdf"
            shutil.copy2(dpath, td)
            mapped.append((td, tp, pdf))
        render_docx_batch([(a, b) for a, b, _ in mapped], progress_every=40, log=[])
        for td, tp, pdf in mapped:
            if tp.exists() and tp.stat().st_size > 0:
                shutil.copy2(tp, pdf)
        print(f"  PDF {min(i+40, len(jobs))}/{len(jobs)}", flush=True)
        shutil.rmtree(tmp, ignore_errors=True)

    # copy one control sample to desktop
    desk = Path(r"C:\Users\burcu\Desktop\İNCİ AKÜ PPWR")
    desk.mkdir(parents=True, exist_ok=True)
    shutil.copy2(sample, desk / "KONTROL_DoC_imza_tarih_ST-012-EUR-02.docx")
    print("SAMPLE", desk / "KONTROL_DoC_imza_tarih_ST-012-EUR-02.docx")
    print("DONE")


if __name__ == "__main__":
    main()
