"""PILOT production for 4 PPWR scopes into isolated CANDIDATE roots.

Does NOT overwrite existing FINAL / product-level / component release trees.
Pilot only — full run is a separate step after QA PASS.
"""

from __future__ import annotations

import hashlib
import json
import shutil
import subprocess
import sys
import time
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from photo_annex import append_photo_annex, resolve_photos  # noqa: E402

OUT = ROOT / "output"
STARTER_SRC = OUT / "INCI_AKU_PPWR_STARTER_PRODUCT_LEVEL_CUSTOMER_DELIVERY_REV00_CANDIDATE" / "01_PRODUCT_DOCUMENT_SETS"
STARTER_MASTER = OUT / "INCI_AKU_PPWR_STARTER_MASTER_Rev00.xlsx"
IND_SRC = OUT / "PHASE_I_FINAL" / "02_INDUSTRIAL"
CNT_SRC = OUT / "PHASE_I_FINAL" / "03_CONTAINER"
CMP_SRC = OUT / "INCI_AKU_PPWR_COMPONENT_PACKAGING_CUSTOMER_DELIVERY_REV00" / "01_DOCUMENT_SETS"
IND_MASTER = OUT / "INCI_AKU_PPWR_INDUSTRIAL_MASTER_Rev00.xlsx"
SIG = ROOT / "assets" / "signatory" / "numan_alver_signature_transparent.png"

STARTER_ROOT = OUT / "01_STARTER_INDIVIDUAL_DELIVERY_REV00_CANDIDATE"
IND_ROOT = OUT / "02_INDUSTRIAL_DELIVERY_REV00_CANDIDATE"
CNT_ROOT = OUT / "03_CONTAINER_DELIVERY_REV00_CANDIDATE"
CMP_ROOT = OUT / "04_COMPONENT_SPARE_DELIVERY_REV00_CANDIDATE"

STEMS = [
    "01_Technical_File",
    "02_EU_DoC",
    "03_Label",
    "04_Shipment_Statement",
]
QMS_BY_FILE = {
    "01_Technical_File.docx": "YS/D/0020",
    "02_EU_DoC.docx": "YS/D/0021",
    "03_Label.docx": "YS/D/0022",
    "04_Shipment_Statement.docx": "YS/D/0023",
}
NAVY = "0E2A47"
WHITE = "FFFFFF"
FONT = "Tahoma"
PLACE_DATE = "11.08.2026"


def _set_para_text(paragraph, text: str, *, bold: bool = True) -> None:
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].bold = bold
        paragraph.runs[0].font.size = Pt(9)
        paragraph.runs[0].font.name = FONT
    else:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(9)
        run.font.name = FONT


def stamp_qms_pack(folder: Path) -> None:
    for fname, code in QMS_BY_FILE.items():
        path = folder / fname
        if not path.exists():
            continue
        line = (
            f"Doküman No/Doc. Nr.: {code}    "
            f"Yayın Trh./Rel. Date: {PLACE_DATE}   "
            "Rev.No/Rev.Nr.: 00   Rev.Trh./Rev.Date: - - -"
        )
        doc = Document(str(path))
        found = False
        for p in doc.paragraphs:
            t = p.text or ""
            if ("Doküman No" in t or "Dokuman No" in t or "Doc. Nr" in t) and (
                code in t or "YS/D/002" in t or "IA-PPWR" not in t
            ):
                if "YS/D/002" in t or "Doküman No" in t or "Dokuman No" in t:
                    _set_para_text(p, line)
                    found = True
                    break
        if not found and doc.paragraphs:
            # insert after title
            title = doc.paragraphs[0]
            new_p_elm = OxmlElement("w:p")
            title._p.addnext(new_p_elm)
            from docx.text.paragraph import Paragraph

            new_para = Paragraph(new_p_elm, title._parent)
            _set_para_text(new_para, line)
        for sec in doc.sections:
            for p in sec.footer.paragraphs:
                if "YS/D/002" in (p.text or "") or "Doküman" in (p.text or "") or "Doc. Nr" in (p.text or ""):
                    _set_para_text(p, f"Doküman No / Doc. Nr.: {code}")
        doc.save(str(path))


def ensure_doc_signature(folder: Path) -> None:
    path = folder / "02_EU_DoC.docx"
    if not path.exists() or not SIG.exists():
        return
    doc = Document(str(path))
    applied = False
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            left = (row.cells[0].text or "").upper().replace("İ", "I").replace("Ö", "O").replace("Ü", "U")
            right = row.cells[-1]
            if ("SIGNATURE" in left or "IMZA" in left) and (
                "NAME" in left or "AD /" in left or "FUNCTION" in left or "GOREV" in left
            ):
                for p in list(right.paragraphs):
                    el = p._p
                    parent = el.getparent()
                    if parent is not None:
                        parent.remove(el)
                p1 = right.add_paragraph()
                _set_para_text(p1, "Ad / Name: Numan Alver")
                p2 = right.add_paragraph()
                _set_para_text(p2, "Görev / Function: Operasyon Direktörü / Operations Director")
                p3 = right.add_paragraph()
                _set_para_text(p3, "İmza / Signature:", bold=False)
                p4 = right.add_paragraph()
                p4.add_run().add_picture(str(SIG), width=Cm(4.2))
                applied = True
                break
        if applied:
            break
    if not applied:
        # fallback append
        blob = "\n".join(p.text for p in doc.paragraphs)
        if "Numan Alver" not in blob:
            doc.add_paragraph("Ad / Name: Numan Alver")
            doc.add_paragraph("Görev / Function: Operasyon Direktörü / Operations Director")
            doc.add_paragraph().add_run().add_picture(str(SIG), width=Cm(4.2))
    doc.save(str(path))


def prepare_customer_pack(folder: Path) -> None:
    stamp_qms_pack(folder)
    ensure_doc_signature(folder)


def kill_word() -> None:
    subprocess.run(["taskkill", "/F", "/IM", "WINWORD.EXE"], capture_output=True, check=False)
    time.sleep(1.5)


def ensure_dirs(root: Path, docs_subdir: str) -> tuple[Path, Path]:
    control = root / "00_CONTROL"
    docs = root / docs_subdir
    control.mkdir(parents=True, exist_ok=True)
    docs.mkdir(parents=True, exist_ok=True)
    return control, docs


def copy_pack(src_dir: Path, dest_dir: Path) -> None:
    if dest_dir.exists():
        shutil.rmtree(dest_dir)
    shutil.copytree(src_dir, dest_dir)


def bom_from_industrial(set_code: str) -> list[dict]:
    if not IND_MASTER.exists():
        return []
    wb = load_workbook(IND_MASTER, data_only=True)
    ws = wb["BOM_MASTER"]
    headers = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
    hi = {h: i for i, h in enumerate(headers)}
    rows = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if str(row[hi["Packaging Set Code"]] or "") != set_code:
            continue
        rows.append(
            {
                "component_code": str(row[hi["Component Code"]] or ""),
                "description": str(row[hi["Component Description"]] or ""),
            }
        )
    wb.close()
    return rows


def bom_from_starter_product(product_code: str) -> list[dict]:
    wb = load_workbook(STARTER_MASTER, data_only=True)
    pm = wb["PRODUCT_MASTER"]
    headers = [c.value for c in next(pm.iter_rows(min_row=1, max_row=1))]
    hi = {h: i for i, h in enumerate(headers)}
    set_code = None
    for row in pm.iter_rows(min_row=2, values_only=True):
        if str(row[hi["Product Code"]] or "").strip() == product_code:
            set_code = str(row[hi["Packaging Set Code"]] or "").strip()
            break
    rows = []
    if set_code and "BOM_MASTER" in wb.sheetnames:
        bm = wb["BOM_MASTER"]
        bh = [c.value for c in next(bm.iter_rows(min_row=1, max_row=1))]
        bhi = {h: i for i, h in enumerate(bh)}
        for row in bm.iter_rows(min_row=2, values_only=True):
            sc = str(row[bhi.get("Packaging Set Code", 0)] or "")
            if sc != set_code:
                continue
            # flexible headers
            code = ""
            desc = ""
            for k in ("Component Code", "Component_Code", "Material Code"):
                if k in bhi:
                    code = str(row[bhi[k]] or "")
                    break
            for k in ("Component Description", "Description", "Component Name"):
                if k in bhi:
                    desc = str(row[bhi[k]] or "")
                    break
            rows.append({"component_code": code, "description": desc})
    wb.close()
    return rows


def bom_from_component_variant(variant_id: str) -> list[dict]:
    # Prefer evidence workbook BOM-like rows from existing engine REFERENCE_BOM if present
    eng = OUT / "INCI_AKU_PPWR_COMPONENT_PACKAGING_CUSTOMER_DELIVERY_REV00" / "00_CONTROL" / "INCI_AKU_PPWR_COMPONENT_PACKAGING_ENGINE_Rev00.xlsx"
    if not eng.exists():
        return [
            {"component_code": "PKG-PALLET", "description": "PALET EURO"},
            {"component_code": "PKG-STRETCH", "description": "STRECH"},
            {"component_code": "PKG-CARTON", "description": "KARTON KOLI"},
        ]
    wb = load_workbook(eng, data_only=True)
    rows = []
    sheet = "REFERENCE_BOM" if "REFERENCE_BOM" in wb.sheetnames else None
    if sheet:
        ws = wb[sheet]
        headers = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
        # take first column pairs loosely
        for row in ws.iter_rows(min_row=2, values_only=True):
            vals = [str(v).strip() if v is not None else "" for v in row]
            if not any(vals):
                continue
            # heuristic: code-like + description
            code = vals[0]
            desc = vals[1] if len(vals) > 1 else vals[0]
            if "CMP" in code.upper() and "1ROW" in variant_id.upper() and "2" in code:
                continue
            rows.append({"component_code": code, "description": desc})
            if len(rows) >= 12:
                break
    wb.close()
    if not rows:
        rows = [
            {"component_code": "EURO-PALLET", "description": "PALET EURO"},
            {"component_code": "STRETCH", "description": "STRECH"},
            {"component_code": "CARTON", "description": "KARTON KOLI"},
            {"component_code": "TAPE", "description": "BANT"},
        ]
    return rows


def convert_pdfs(folder: Path) -> tuple[int, int]:
    import pythoncom
    import win32com.client

    ok = fail = 0
    kill_word()
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        for stem in STEMS:
            docx = folder / f"{stem}.docx"
            pdf = folder / f"{stem}.pdf"
            if not docx.exists():
                continue
            doc = None
            try:
                if pdf.exists():
                    try:
                        pdf.unlink()
                    except Exception:
                        pass
                doc = word.Documents.Open(str(docx.resolve()), False, True, False, "", "", False, "", "", 0, 0, True, True)
                doc.ExportAsFixedFormat(str(pdf.resolve()), 17, False, 0, False, 0, 0, False, True, False, 0, True, True, False)
                if pdf.exists() and pdf.stat().st_size > 0:
                    ok += 1
                else:
                    fail += 1
            except Exception:
                fail += 1
            finally:
                if doc is not None:
                    try:
                        doc.Close(False)
                    except Exception:
                        pass
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        kill_word()
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
    return ok, fail


def write_mini_engine(
    control: Path,
    *,
    title: str,
    records: list[dict],
    docs_rel_prefix: str,
    engine_name: str,
) -> Path:
    """records: key, label, folder_name"""
    wb = Workbook()
    home = wb.active
    home.title = "00_HOME"
    home["B2"] = title
    home["B2"].font = Font(name=FONT, size=16, bold=True, color=NAVY)
    home["B4"] = "PILOT MODE"
    home["C4"] = "CANDIDATE"
    home["B5"] = "RECORDS"
    home["C5"] = len(records)
    home["B6"] = "SYSTEM STATUS"
    home["C6"] = "PILOT"

    dc = wb.create_sheet("DOCUMENT_CENTER")
    dc["A1"] = "DOCUMENT CENTER — PILOT"
    headers = [
        "Key",
        "Description",
        "TF WORD",
        "TF PDF",
        "DoC WORD",
        "DoC PDF",
        "Label WORD",
        "Label PDF",
        "STM WORD",
        "STM PDF",
    ]
    for i, h in enumerate(headers, 1):
        cell = dc.cell(4, i, h)
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.font = Font(name=FONT, color=WHITE, bold=True)

    search = wb.create_sheet("SEARCH")
    sdata = wb.create_sheet("SEARCH_DATA")
    search["A1"] = "SEARCH (PILOT)"
    search["A3"] = "Key"
    search["B3"] = ""
    sdata_headers = ["Key", "Description", "TF", "DoC", "Label", "STM"]
    for i, h in enumerate(sdata_headers, 1):
        sdata.cell(1, i, h)

    for i, rec in enumerate(records, 1):
        key = rec["key"]
        desc = rec["label"]
        folder = rec["folder"]
        rel = f"..\\{docs_rel_prefix}\\{folder}\\"
        r = i + 4
        dc.cell(r, 1, key)
        dc.cell(r, 2, desc)
        files = [
            (3, "01_Technical_File.docx"),
            (4, "01_Technical_File.pdf"),
            (5, "02_EU_DoC.docx"),
            (6, "02_EU_DoC.pdf"),
            (7, "03_Label.docx"),
            (8, "03_Label.pdf"),
            (9, "04_Shipment_Statement.docx"),
            (10, "04_Shipment_Statement.pdf"),
        ]
        for col, fname in files:
            cell = dc.cell(r, col)
            cell.value = "OPEN WORD" if fname.endswith(".docx") else "OPEN PDF"
            cell.hyperlink = f"{rel}{fname}"
            cell.font = Font(name=FONT, color="0563C1", underline="single")
        sdata.cell(i + 1, 1, key)
        sdata.cell(i + 1, 2, desc)
        sdata.cell(i + 1, 3, f"{rel}01_Technical_File.docx")
        sdata.cell(i + 1, 4, f"{rel}02_EU_DoC.docx")
        sdata.cell(i + 1, 5, f"{rel}03_Label.docx")
        sdata.cell(i + 1, 6, f"{rel}04_Shipment_Statement.docx")

    reg = wb.create_sheet("DOCUMENT_REGISTER")
    reg.append(["Key", "Description", "Status"])
    for rec in records:
        reg.append([rec["key"], rec["label"], "PILOT"])

    path = control / engine_name
    wb.save(path)
    # launcher
    (control.parent / "00_AC_DOCUMENT_ENGINE.cmd").write_text(
        "@echo off\r\n"
        f"start \"\" \"%~dp000_CONTROL\\{engine_name}\"\r\n",
        encoding="utf-8",
    )
    return path


def qa_pack(folder: Path) -> dict:
    info = {"folder": folder.name, "docx": 0, "pdf": 0, "qms": {}, "sig": False, "placeholder": False}
    for stem in STEMS:
        d = folder / f"{stem}.docx"
        p = folder / f"{stem}.pdf"
        if d.exists():
            info["docx"] += 1
            doc = Document(str(d))
            blob = "\n".join(p.text for p in doc.paragraphs)
            blob += "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
            for qms in ("YS/D/0020", "YS/D/0021", "YS/D/0022", "YS/D/0023"):
                if qms in blob:
                    info["qms"][qms] = True
            if "Numan Alver" in blob:
                info["sig"] = True
            if "[[SIGNATORY_SIGNATURE_IMAGE]]" in blob:
                info["placeholder"] = True
            # annex marker
            if stem.startswith("01_") and "Representative Packaging Component Photos" in blob:
                info["photo_annex"] = True
        if p.exists() and p.stat().st_size > 0:
            info["pdf"] += 1
    return info


def write_data_gap_list(control: Path, codes: list[str]) -> None:
    lines = [
        "STARTER DATA REQUIRED / YURT İÇİ — DOCUMENTS NOT ISSUED",
        "COMPLETE PACKAGING COMPONENT / PALLET DATA NOT AVAILABLE",
        "",
        f"Count: {len(codes)}",
        "",
    ]
    lines.extend(codes)
    (control / "DOMESTIC_DATA_GAP_LIST.txt").write_text("\n".join(lines), encoding="utf-8")


def pilot_starter() -> dict:
    control, docs = ensure_dirs(STARTER_ROOT, "01_PRODUCTS")
    product = "1000069"
    gap = "1004590"
    src = STARTER_SRC / product
    dest = docs / product
    assert src.exists(), f"missing starter source {src}"
    copy_pack(src, dest)
    bom = bom_from_starter_product(product)
    photos = resolve_photos(scope="STARTER", bom_lines=bom)
    n = append_photo_annex(dest / "01_Technical_File.docx", photos, title_extra=f"Product {product}")
    prepare_customer_pack(dest)
    pdf_ok, pdf_fail = convert_pdfs(dest)
    write_data_gap_list(control, [gap])
    eng = write_mini_engine(
        control,
        title="İNCI AKÜ PPWR — STARTER INDIVIDUAL ENGINE (PILOT)",
        records=[{"key": product, "label": f"Starter product {product}", "folder": product}],
        docs_rel_prefix="01_PRODUCTS",
        engine_name="INCI_AKU_PPWR_STARTER_INDIVIDUAL_ENGINE_Rev00.xlsx",
    )
    qa = qa_pack(dest)
    qa.update({"photos": n, "pdf_ok": pdf_ok, "pdf_fail": pdf_fail, "engine": str(eng), "data_gap_sample": gap})
    return qa


def pilot_industrial() -> dict:
    control, docs = ensure_dirs(IND_ROOT, "01_CONFIGS")
    cfg = "IND-24V-01"
    src = IND_SRC / cfg
    dest = docs / cfg
    assert src.exists(), f"missing industrial source {src}"
    copy_pack(src, dest)
    bom = bom_from_industrial(cfg)
    photos = resolve_photos(scope="INDUSTRIAL", bom_lines=bom)
    n = append_photo_annex(dest / "01_Technical_File.docx", photos, title_extra=f"Config {cfg}")
    prepare_customer_pack(dest)
    pdf_ok, pdf_fail = convert_pdfs(dest)
    eng = write_mini_engine(
        control,
        title="İNCI AKÜ PPWR — INDUSTRIAL ENGINE (PILOT)",
        records=[{"key": cfg, "label": "Industrial 24V packaging config", "folder": cfg}],
        docs_rel_prefix="01_CONFIGS",
        engine_name="INCI_AKU_PPWR_INDUSTRIAL_ENGINE_Rev00.xlsx",
    )
    qa = qa_pack(dest)
    qa.update({"photos": n, "pdf_ok": pdf_ok, "pdf_fail": pdf_fail, "engine": str(eng)})
    return qa


def pilot_container() -> dict:
    control, docs = ensure_dirs(CNT_ROOT, "01_CONFIGS")
    cfg = "CNT-20-STD-01"
    src = CNT_SRC / cfg
    dest = docs / cfg
    assert src.exists(), f"missing container source {src}"
    copy_pack(src, dest)
    # generic container BOM cues for photo mapping
    bom = [
        {"component_code": "OSB", "description": "Osb (11 mm)"},
        {"component_code": "SEP", "description": "Karton Seperatör"},
        {"component_code": "AIR", "description": "Hava Yastığı (Level 2)"},
        {"component_code": "LASH", "description": "Bağlama Halatı"},
        {"component_code": "PALLET", "description": "PALET EURO"},
    ]
    photos = resolve_photos(scope="CONTAINER", bom_lines=bom)
    n = append_photo_annex(dest / "01_Technical_File.docx", photos, title_extra=f"Config {cfg}")
    prepare_customer_pack(dest)
    pdf_ok, pdf_fail = convert_pdfs(dest)
    eng = write_mini_engine(
        control,
        title="İNCI AKÜ PPWR — CONTAINER LOADING ENGINE (PILOT)",
        records=[{"key": cfg, "label": "20ft standard container loading", "folder": cfg}],
        docs_rel_prefix="01_CONFIGS",
        engine_name="INCI_AKU_PPWR_CONTAINER_ENGINE_Rev00.xlsx",
    )
    qa = qa_pack(dest)
    qa.update({"photos": n, "pdf_ok": pdf_ok, "pdf_fail": pdf_fail, "engine": str(eng)})
    return qa


def pilot_component() -> dict:
    control, docs = ensure_dirs(CMP_ROOT, "01_VARIANTS")
    # map to requested names while keeping source ids
    mapping = {
        "CMP-1ROW-01": "CMP-TEK-SIRA",
        "CMP-2ROW-01": "CMP-CIFT-SIRA",
    }
    results = []
    records = []
    for src_id, dest_id in mapping.items():
        src = CMP_SRC / src_id
        dest = docs / dest_id
        assert src.exists(), f"missing component source {src}"
        copy_pack(src, dest)
        bom = bom_from_component_variant(src_id)
        photos = resolve_photos(scope="COMPONENT", bom_lines=bom)
        n = append_photo_annex(dest / "01_Technical_File.docx", photos, title_extra=f"Variant {dest_id}")
        prepare_customer_pack(dest)
        pdf_ok, pdf_fail = convert_pdfs(dest)
        qa = qa_pack(dest)
        qa.update({"photos": n, "pdf_ok": pdf_ok, "pdf_fail": pdf_fail, "source": src_id})
        results.append(qa)
        records.append({"key": dest_id, "label": f"Component spare {dest_id}", "folder": dest_id})
    eng = write_mini_engine(
        control,
        title="İNCI AKÜ PPWR — COMPONENT / SPARE ENGINE (PILOT)",
        records=records,
        docs_rel_prefix="01_VARIANTS",
        engine_name="INCI_AKU_PPWR_COMPONENT_SPARE_ENGINE_Rev00.xlsx",
    )
    return {"variants": results, "engine": str(eng)}


def main() -> int:
    print("PILOT multi-scope start", flush=True)
    report = {
        "starter": pilot_starter(),
        "industrial": pilot_industrial(),
        "container": pilot_container(),
        "component": pilot_component(),
    }
    # gate
    def pack_pass(q: dict) -> bool:
        qms_ok = all(q.get("qms", {}).get(c) for c in ("YS/D/0020", "YS/D/0021", "YS/D/0022", "YS/D/0023"))
        photos_ok = int(q.get("photos", 0)) > 0 or q.get("photo_annex", False)
        return (
            q.get("docx", 0) == 4
            and q.get("pdf", 0) == 4
            and q.get("sig", False)
            and not q.get("placeholder", False)
            and q.get("pdf_fail", 0) == 0
            and qms_ok
            and photos_ok
        )

    gates = {
        "STARTER": pack_pass(report["starter"]),
        "INDUSTRIAL": pack_pass(report["industrial"]),
        "CONTAINER": pack_pass(report["container"]),
        "COMPONENT": all(pack_pass(v) for v in report["component"]["variants"]),
    }
    report["PILOT_GATES"] = gates
    report["PILOT_OVERALL"] = "PASS" if all(gates.values()) else "FAIL"
    out = OUT / "_MULTI_SCOPE_PILOT_QA.json"
    out.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({"PILOT_GATES": gates, "OVERALL": report["PILOT_OVERALL"]}, indent=2), flush=True)
    print("QA_JSON", out, flush=True)
    print("STOP_BEFORE_FULL_RUN" if report["PILOT_OVERALL"] != "PASS" else "PILOT_PASS_READY_FOR_FULL", flush=True)
    return 0 if report["PILOT_OVERALL"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
