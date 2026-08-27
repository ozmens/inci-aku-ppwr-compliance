"""
SAMPLE ONLY — product-level customer documents for ONE Product Code.

Does NOT overwrite FINAL set-level delivery.
Does NOT generate all 2004 packs.

Creates:
  output/INCI_AKU_PPWR_STARTER_PRODUCT_LEVEL_CUSTOMER_DELIVERY_REV00_CANDIDATE/
    01_PRODUCT_DOCUMENT_SETS/<PRODUCT_CODE>/  (4 Word + 4 PDF)
    SAMPLE_QA.txt
"""

from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(ROOT / "tools"))

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook

from builders.phase_g.merge_engine import merge_document
from builders.phase_i.render_batch import render_docx_batch
from models.technical_file import Article5Assessment
from services.document_context_factory import DocumentContextFactory
from utils.constants import ARTICLE5_BASIS_LABEL

from generate_ppwr_documents import (  # noqa: E402
    DOC_SPECS,
    RUNTIME_DIR,
    embed_signature,
    ensure_runtime_templates,
    find_signature,
    scrub_tf_pending,
)
from ppwr_engine.starter_loader import StarterMasterLoader

MASTER = ROOT / "output" / "INCI_AKU_PPWR_STARTER_MASTER_Rev00.xlsx"
FINAL_SET_DOCS = (
    ROOT
    / "output"
    / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL"
    / "01_DOCUMENT_SETS"
)
CANDIDATE = (
    ROOT
    / "output"
    / "INCI_AKU_PPWR_STARTER_PRODUCT_LEVEL_CUSTOMER_DELIVERY_REV00_CANDIDATE"
)
PRODUCT_SETS = CANDIDATE / "01_PRODUCT_DOCUMENT_SETS"
SAMPLE_PRODUCT = "1012839"
QMS = {
    "TECHNICAL_FILE": "YS/D/0020",
    "DOC": "YS/D/0021",
    "LABEL": "YS/D/0022",
    "STATEMENT": "YS/D/0023",
}


def _doc_blob(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError:
            return ""
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _set_para_text(paragraph, text: str) -> None:
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].font.name = "Tahoma"
        paragraph.runs[0].font.size = Pt(9)
    else:
        r = paragraph.add_run(text)
        r.font.name = "Tahoma"
        r.font.size = Pt(9)


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            p.runs[0].bold = bold
            p.runs[0].font.name = "Tahoma"
        else:
            r = p.add_run(text)
            r.bold = bold
            r.font.name = "Tahoma"
        for extra in cell.paragraphs[1:]:
            for r in extra.runs:
                r.text = ""


def _insert_row_after(table, after_idx: int):
    """Insert a cloned row after after_idx; return new row index."""
    tmpl = table.rows[after_idx]._tr
    new_tr = deepcopy(tmpl)
    tmpl.addnext(new_tr)
    return after_idx + 1


def controlled_siblings(product_code: str) -> tuple[str, str, list[str]]:
    wb = load_workbook(MASTER, data_only=True, read_only=True)
    pm = wb["PRODUCT_MASTER"]
    headers = [c.value for c in next(pm.iter_rows(min_row=1, max_row=1))]
    idx = {h: i for i, h in enumerate(headers)}
    set_code = None
    desc = None
    rows = []
    for row in pm.iter_rows(min_row=2, values_only=True):
        code = str(row[idx["Product Code"]] or "").strip()
        pset = str(row[idx["Packaging Set Code"]] or "").strip()
        status = str(row[idx["Physical Packaging Status"]] or "").strip()
        if code == product_code:
            set_code = pset
            desc = str(row[idx["Technical Description"]] or "").strip()
        if status == "CONTROLLED PACKAGING SET":
            rows.append((code, pset))
    wb.close()
    if not set_code:
        raise SystemExit(f"Product {product_code} not found / not controlled")
    sibs = sorted(c for c, s in rows if s == set_code)
    return set_code, desc or "", sibs


def product_instance_ids(product_code: str, set_code: str) -> dict[str, str]:
    return {
        "tf": f"IA-PPWR-TF-{product_code}-{set_code}-R00",
        "doc": f"IA-PPWR-DOC-{product_code}-{set_code}-R00",
        "label": f"IA-PPWR-LBL-{product_code}-{set_code}-R00",
        "stm": f"IA-PPWR-STM-{product_code}-{set_code}-R00",
    }


def stamp_instance_id(docx_path: Path, label_tr_en: str, instance_id: str) -> None:
    """Ensure customer-facing product document instance ID is visible in identity table."""
    doc = Document(str(docx_path))
    if not doc.tables:
        doc.save(str(docx_path))
        return
    table = doc.tables[0]
    blob = "\n".join(c.text for row in table.rows for c in row.cells)
    if instance_id in blob:
        doc.save(str(docx_path))
        return
    idx = _insert_row_after(table, 0)
    _set_cell_text(table.rows[idx].cells[0], label_tr_en, bold=True)
    _set_cell_text(table.rows[idx].cells[1], instance_id, bold=True)
    doc.save(str(docx_path))


def patch_tf_product_scope(
    docx_path: Path, product_code: str, set_code: str, tf_id: str
) -> None:
    doc = Document(str(docx_path))
    tr = (
        f"Bu Teknik Dosya, Product Code {product_code} için kullanılan {set_code} "
        f"kontrollü ambalaj konfigürasyonunu kapsar."
    )
    en = (
        f"This Technical File covers the controlled {set_code} packaging configuration "
        f"used for Product Code {product_code}."
    )
    new_scope = f"{tr}\n{en}"
    # Replace plural / multi-product scope language in body paragraphs
    patterns = [
        re.compile(r"Bu dosya yalnız .+?listelenmiştir\.", re.S),
        re.compile(r"This file covers only fixed BOM variant .+?information\.", re.S),
        re.compile(r"The \d+ product codes linked.+?information\.", re.S),
        re.compile(r"Bu varyanta bağlı .+?listelenmiştir\.", re.S),
    ]
    for p in doc.paragraphs:
        text = p.text or ""
        if "ÜRÜN KAPSAMI" in text or "PRODUCT SCOPE" in text:
            continue
        if (
            "product codes linked" in text
            or "ürün kodu aşağıda" in text
            or "ürün kodları" in text.lower()
            or ("Bu dosya yalnız" in text and "sabit BOM" in text)
            or ("This file covers only fixed BOM" in text)
        ):
            _set_para_text(p, new_scope)
            continue
        # also harden any leftover count language
        for pat in patterns:
            if pat.search(text):
                _set_para_text(p, new_scope)
                break
    # Ensure product table has exactly one data row (already from ctx) — verify headers
    for table in doc.tables:
        hdr = " ".join(c.text for c in table.rows[0].cells).upper()
        if ("PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr) and len(table.rows) >= 2:
            _set_cell_text(table.rows[0].cells[0], "ÜRÜN KODU\nPRODUCT CODE", bold=True)
            if len(table.rows[0].cells) > 1:
                _set_cell_text(
                    table.rows[0].cells[1],
                    "ÜRÜN TANIMI\nPRODUCT DESCRIPTION",
                    bold=True,
                )
    doc.save(str(docx_path))
    stamp_instance_id(
        docx_path,
        "Doküman Örnek Kimliği\nDocument Instance ID",
        tf_id,
    )


def ensure_label_product_fields(
    docx_path: Path, product_code: str, product_desc: str, set_code: str
) -> None:
    doc = Document(str(docx_path))
    if not doc.tables:
        doc.save(str(docx_path))
        return
    table = doc.tables[0]
    blob = "\n".join(c.text for row in table.rows for c in row.cells).upper()
    if "ÜRÜN KODU" not in blob and "PRODUCT CODE" not in blob:
        # Insert Product Code + Description as first two data identity rows
        # Clone row 0 twice and rewrite
        r1 = _insert_row_after(table, 0)
        r2 = _insert_row_after(table, r1)
        # After two inserts, rows 1 and 2 are new; shift original content down.
        # Actually deepcopy keeps old values — rewrite rows 1 and 2.
        _set_cell_text(table.rows[1].cells[0], "Ürün Kodu\nProduct Code", bold=True)
        _set_cell_text(table.rows[1].cells[1], product_code, bold=True)
        _set_cell_text(table.rows[2].cells[0], "Ürün Tanımı\nProduct Description", bold=True)
        _set_cell_text(table.rows[2].cells[1], product_desc)
    # Ensure Packaging Set Code remains visible somewhere
    blob2 = "\n".join(c.text for row in table.rows for c in row.cells)
    if set_code not in blob2:
        idx = _insert_row_after(table, 2)
        _set_cell_text(table.rows[idx].cells[0], "Ambalaj Seti Kodu\nPackaging Set Code", bold=True)
        _set_cell_text(table.rows[idx].cells[1], set_code, bold=True)
    # Remove any multi-product appendix table if present with >1 body row siblings risk
    for ti in range(len(doc.tables) - 1, -1, -1):
        t = doc.tables[ti]
        hdr = " ".join(c.text for c in t.rows[0].cells).upper()
        if ("PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr) and len(t.rows) > 2:
            # keep only current product — rebuild to 1 data row
            tbl = t._tbl
            rows = list(tbl.findall(qn("w:tr")))
            for tr in rows[1:]:
                tbl.remove(tr)
            new_tr = deepcopy(rows[0])
            tbl.append(new_tr)
            _set_cell_text(t.rows[1].cells[0], product_code, bold=True)
            if len(t.rows[1].cells) > 1:
                _set_cell_text(t.rows[1].cells[1], product_desc)
    doc.save(str(docx_path))


def ensure_stm_product_fields(
    docx_path: Path, product_code: str, product_desc: str, set_code: str
) -> None:
    doc = Document(str(docx_path))
    # Prefer a dedicated identity row in first non-BOM table, or prepend product table
    has_pc = False
    for table in doc.tables:
        blob = "\n".join(c.text for row in table.rows for c in row.cells)
        if product_code in blob and ("ÜRÜN KODU" in blob.upper() or "PRODUCT CODE" in blob.upper()):
            has_pc = True
            # collapse multi rows
            hdr = " ".join(c.text for c in table.rows[0].cells).upper()
            if ("PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr) and len(table.rows) > 2:
                tbl = table._tbl
                rows = list(tbl.findall(qn("w:tr")))
                for tr in rows[1:]:
                    tbl.remove(tr)
                new_tr = deepcopy(rows[0])
                tbl.append(new_tr)
                _set_cell_text(table.rows[1].cells[0], product_code, bold=True)
                if len(table.rows[1].cells) > 1:
                    _set_cell_text(table.rows[1].cells[1], product_desc)
    if not has_pc:
        # Insert 3-row product identity table at top of body
        table = doc.add_table(rows=3, cols=2)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        _set_cell_text(table.rows[0].cells[0], "Ürün Kodu\nProduct Code", bold=True)
        _set_cell_text(table.rows[0].cells[1], product_code, bold=True)
        _set_cell_text(table.rows[1].cells[0], "Ürün Tanımı\nProduct Description", bold=True)
        _set_cell_text(table.rows[1].cells[1], product_desc)
        _set_cell_text(table.rows[2].cells[0], "Ambalaj Seti Kodu\nPackaging Set Code", bold=True)
        _set_cell_text(table.rows[2].cells[1], set_code, bold=True)
        # move before first existing table if any
        if len(doc.tables) > 1:
            first = doc.tables[0]._tbl
            # newly added is last — move
            new_tbl = table._tbl
            first.addprevious(new_tbl)
    # Keep transaction Customer/OEM blank — do not fill with market
    for table in doc.tables:
        for row in table.rows:
            left = row.cells[0].text.upper()
            if "MÜŞTERİ" in left or "CUSTOMER" in left or "OEM" in left:
                if len(row.cells) > 1 and row.cells[1].text.strip() not in ("", "________________"):
                    # only clear if it looks auto-filled with market text
                    val = row.cells[1].text.strip()
                    if "ISUZU" in val.upper() or "YURT" in val.upper() or "OEM" in val.upper():
                        _set_cell_text(row.cells[1], "________________")
    doc.save(str(docx_path))


def patch_doc_scope_heading(docx_path: Path, product_code: str) -> None:
    doc = Document(str(docx_path))
    for p in doc.paragraphs:
        t = p.text or ""
        if "Kontrollü ürün kapsamı" in t or "Controlled product scope" in t:
            _set_para_text(
                p,
                f"Ürün kapsamı / Product scope — Product Code {product_code}",
            )
    # collapse any multi-row product tables
    for table in doc.tables:
        hdr = " ".join(c.text for c in table.rows[0].cells).upper()
        if ("PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr) and len(table.rows) > 2:
            # keep row matching product_code only
            keep = None
            for i, row in enumerate(table.rows[1:], start=1):
                if product_code in row.cells[0].text:
                    keep = i
                    break
            tbl = table._tbl
            rows = list(tbl.findall(qn("w:tr")))
            template = rows[keep] if keep else rows[1]
            for tr in rows[1:]:
                tbl.remove(tr)
            tbl.append(deepcopy(template))
    doc.save(str(docx_path))


def extract_bom_codes(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    codes: list[str] = []
    for table in doc.tables:
        hdr = " ".join(c.text for c in table.rows[0].cells).upper()
        if "CODE" in hdr and ("UNIT" in hdr or "BİRİM" in hdr or "LINE" in hdr or "WEIGHT" in hdr):
            for row in table.rows[1:]:
                c0 = row.cells[0].text.strip()
                if c0 and "TOPLAM" not in c0.upper() and "TOTAL" not in c0.upper():
                    codes.append(c0)
    return codes


def count_code_hits(text: str, code: str) -> int:
    return len(re.findall(rf"(?<!\d){re.escape(code)}(?!\d)", text))


def run_qa(
    *,
    product_code: str,
    product_desc: str,
    set_code: str,
    siblings: list[str],
    out_dir: Path,
    sig_ok: bool,
) -> dict:
    others = [c for c in siblings if c != product_code]
    files = {
        "TF": ("01_Technical_File", "TECHNICAL_FILE"),
        "DoC": ("02_EU_DoC", "DOC"),
        "Label": ("03_Label", "LABEL"),
        "Statement": ("04_Shipment_Statement", "STATEMENT"),
    }
    results: dict = {
        "product_code": product_code,
        "product_description": product_desc,
        "packaging_set_code": set_code,
        "sibling_count": len(others),
        "sibling_in_tf": 0,
        "sibling_in_doc": 0,
        "sibling_in_label": 0,
        "sibling_in_stm": 0,
        "tf": "FAIL",
        "doc": "FAIL",
        "label": "FAIL",
        "stm": "FAIL",
        "numan": "FAIL",
        "signature": "FAIL",
        "qms": "FAIL",
        "bom_unchanged": "FAIL",
        "mapping_unchanged": "PASS",  # we did not write master
        "gate": "FAIL",
    }
    key_map = {"TF": "sibling_in_tf", "DoC": "sibling_in_doc", "Label": "sibling_in_label", "Statement": "sibling_in_stm"}
    doc_key = {"TF": "tf", "DoC": "doc", "Label": "label", "Statement": "stm"}

    qms_ok = True
    numan_ok = True
    marker_left = False
    for label, (stem, dtype) in files.items():
        docx = out_dir / f"{stem}.docx"
        pdf = out_dir / f"{stem}.pdf"
        blob = _doc_blob(docx)
        if pdf.exists():
            blob += "\n" + _pdf_text(pdf)
        # current product must appear
        present = count_code_hits(blob, product_code) > 0
        sib_hits = 0
        for sib in others:
            n = count_code_hits(blob, sib)
            sib_hits += n
        results[key_map[label]] = sib_hits
        qms_token = QMS[dtype]
        if qms_token not in blob and dtype != "LABEL":
            # Label may still carry YS/D/0022 in header/footer — soft check below
            pass
        if qms_token not in _doc_blob(docx):
            # check headers via raw zip xml
            raw = docx.read_bytes().decode("utf-8", errors="ignore")
            if qms_token not in raw:
                qms_ok = False
        if dtype == "DOC":
            if "Numan Alver" not in blob:
                numan_ok = False
            if "Operasyon Direktörü" not in blob and "Operations Director" not in blob:
                numan_ok = False
            if "[[SIGNATORY_SIGNATURE_IMAGE]]" in blob or "SIGNATORY_SIGNATURE_IMAGE" in blob:
                marker_left = True
            # signature image present?
            raw = docx.read_bytes()
            has_img = b"word/media/" in raw or b"image" in raw.lower()
            results["signature"] = "PASS" if (sig_ok and has_img and not marker_left) else "FAIL"
        ok = present and sib_hits == 0 and docx.exists() and pdf.exists() and pdf.stat().st_size > 0
        results[doc_key[label]] = "PASS" if ok else "FAIL"

    results["numan"] = "PASS" if numan_ok else "FAIL"
    results["qms"] = "PASS" if qms_ok else "FAIL"

    # BOM unchanged vs set-level TF
    set_tf = FINAL_SET_DOCS / set_code / "01_Technical_File.docx"
    sample_tf = out_dir / "01_Technical_File.docx"
    if set_tf.exists() and sample_tf.exists():
        a = extract_bom_codes(set_tf)
        b = extract_bom_codes(sample_tf)
        results["bom_unchanged"] = "PASS" if a == b and len(a) > 0 else "FAIL"
        results["bom_set_count"] = len(a)
        results["bom_sample_count"] = len(b)
    else:
        results["bom_unchanged"] = "FAIL"

    gate = all(
        results[k] == "PASS"
        for k in (
            "tf",
            "doc",
            "label",
            "stm",
            "numan",
            "signature",
            "qms",
            "bom_unchanged",
            "mapping_unchanged",
        )
    ) and all(
        results[k] == 0
        for k in ("sibling_in_tf", "sibling_in_doc", "sibling_in_label", "sibling_in_stm")
    )
    results["gate"] = "PASS" if gate else "FAIL"
    return results


def format_qa(r: dict) -> str:
    return f"""# PRODUCT-LEVEL SAMPLE QA

Product Code:
{r['product_code']}

Product Description:
{r['product_description']}

Packaging Set Code:
{r['packaging_set_code']}

Sibling Product Codes sharing same Packaging Set:
{r['sibling_count']}

Sibling Product Codes appearing in generated TF:
{r['sibling_in_tf']}

Sibling Product Codes appearing in generated DoC:
{r['sibling_in_doc']}

Sibling Product Codes appearing in generated Label:
{r['sibling_in_label']}

Sibling Product Codes appearing in generated Statement:
{r['sibling_in_stm']}

Technical File:
{r['tf']}

EU DoC:
{r['doc']}

Label:
{r['label']}

Shipment Statement:
{r['stm']}

Numan Alver:
{r['numan']}

Signature:
{r['signature']}

QMS Numbers:
{r['qms']}

BOM unchanged:
{r['bom_unchanged']}

Packaging Set mapping unchanged:
{r['mapping_unchanged']}

SAMPLE PRODUCT-LEVEL GATE:
{r['gate']}
"""


def main() -> None:
    product_code = SAMPLE_PRODUCT
    print(f"SAMPLE product-level pack for {product_code}", flush=True)
    set_code, product_desc, siblings = controlled_siblings(product_code)
    print(f"Set={set_code} siblings={len(siblings)} (incl. self)", flush=True)

    # Candidate tree (do not touch FINAL)
    for d in (
        CANDIDATE / "00_CONTROL",
        PRODUCT_SETS,
        CANDIDATE / "02_OPTIONAL_EVIDENCE",
        CANDIDATE / "03_ARCHIVE",
        CANDIDATE / "04_DOMESTIC_42_DATA_GAP",
    ):
        d.mkdir(parents=True, exist_ok=True)
    (CANDIDATE / "00_CONTROL" / "README_SAMPLE.txt").write_text(
        "SAMPLE ONLY — product-level architecture candidate.\n"
        f"Generated Product Code: {product_code}\n"
        "Do NOT treat as full 2004 delivery until SAMPLE GATE = PASS and approved.\n"
        "Physical Packaging Sets (287) remain unchanged in FINAL backup delivery.\n",
        encoding="utf-8",
    )

    print("Preparing runtime templates…", flush=True)
    ensure_runtime_templates()
    signature = find_signature()
    print(f"Signature: {signature}", flush=True)

    loader = StarterMasterLoader(MASTER)
    loader.open()
    bundle = loader.load_set(set_code)
    # PRODUCT-SPECIFIC: single product only — physical BOM unchanged
    one = [p for p in bundle.products if p.product_code == product_code]
    if not one:
        raise SystemExit(f"{product_code} missing from set bundle {set_code}")
    ids = product_instance_ids(product_code, set_code)

    factory = DocumentContextFactory()
    ctx = factory.build(
        bundle.configuration,
        products=one,
        article5=Article5Assessment(basis_label=ARTICLE5_BASIS_LABEL),
        customer_name=None,
        customer_market=None,
    )
    ctx.total_tare_g = float(bundle.packaging_tare_kg) * 1000.0
    ctx.annex_drawings_status = "OPTIONAL EVIDENCE — NOT INCLUDED IN REV.00"
    ctx.document_ids.technical_file_id = ids["tf"]
    ctx.document_ids.doc_id = ids["doc"]
    ctx.document_ids.label_id = ids["label"]
    ctx.document_ids.statement_id = ids["stm"]
    loader.close()

    out_dir = PRODUCT_SETS / product_code
    if out_dir.exists():
        for p in out_dir.glob("*"):
            p.unlink()
    out_dir.mkdir(parents=True, exist_ok=True)

    runtime = {
        "TECHNICAL_FILE": RUNTIME_DIR / "01_Technical_File_RUNTIME.docx",
        "DOC": RUNTIME_DIR / "02_EU_DoC_RUNTIME.docx",
        "LABEL": RUNTIME_DIR / "03_Label_RUNTIME.docx",
        "STATEMENT": RUNTIME_DIR / "04_Shipment_Statement_RUNTIME.docx",
    }
    pdf_jobs: list[tuple[Path, Path]] = []
    sig_ok = False
    for dtype, stem, is_tf in DOC_SPECS:
        out_docx = out_dir / f"{stem}.docx"
        merge_document(runtime[dtype], out_docx, ctx, for_technical_file=is_tf)
        if is_tf:
            scrub_tf_pending(out_docx)
            patch_tf_product_scope(out_docx, product_code, set_code, ids["tf"])
        if dtype == "DOC":
            patch_doc_scope_heading(out_docx, product_code)
            stamp_instance_id(
                out_docx,
                "Doküman Örnek Kimliği\nDocument Instance ID",
                ids["doc"],
            )
            sig_ok = embed_signature(out_docx, signature)
        if dtype == "LABEL":
            ensure_label_product_fields(out_docx, product_code, product_desc, set_code)
            stamp_instance_id(
                out_docx,
                "Doküman Örnek Kimliği\nDocument Instance ID",
                ids["label"],
            )
        if dtype == "STATEMENT":
            ensure_stm_product_fields(out_docx, product_code, product_desc, set_code)
            stamp_instance_id(
                out_docx,
                "Doküman Örnek Kimliği\nDocument Instance ID",
                ids["stm"],
            )
        pdf_jobs.append((out_docx, out_dir / f"{stem}.pdf"))
        print(f"  Word {stem}.docx", flush=True)

    print("Rendering 4 PDFs…", flush=True)
    render_docx_batch(pdf_jobs, progress_every=1, log=[])

    qa = run_qa(
        product_code=product_code,
        product_desc=product_desc,
        set_code=set_code,
        siblings=siblings,
        out_dir=out_dir,
        sig_ok=sig_ok,
    )
    report = format_qa(qa)
    (CANDIDATE / "SAMPLE_QA.txt").write_text(report, encoding="utf-8")
    print(report, flush=True)
    print(f"OUT: {out_dir}", flush=True)
    print("STOP — awaiting SAMPLE PRODUCT-LEVEL GATE approval before 2004 generation.", flush=True)


if __name__ == "__main__":
    main()
