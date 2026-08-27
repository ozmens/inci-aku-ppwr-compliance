"""Stamp QMS management control line on all 1148 Word docs, then re-render PDFs + ZIP.

Does not change mappings/BOM/IDs. Keeps IA-PPWR configuration IDs.
"""

from __future__ import annotations

import hashlib
import shutil
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(r"C:\Users\burcu\Documents\YAZILIM\Inci_Aku_PPWR_PIMS")
FINAL = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL"
DOC_SETS = FINAL / "01_DOCUMENT_SETS"
ZIP_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL.zip"
SHA_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL_SHA256.txt"

QMS_LINE = {
    "01_Technical_File.docx": (
        "Doküman No/Doc. Nr.: YS/D/0020    "
        "Yayın Trh./Rel. Date: 11.08.2026   "
        "Rev.No/Rev.Nr.: 00   Rev.Trh./Rev.Date: - - -"
    ),
    "02_EU_DoC.docx": (
        "Doküman No/Doc. Nr.: YS/D/0021    "
        "Yayın Trh./Rel. Date: 11.08.2026   "
        "Rev.No/Rev.Nr.: 00   Rev.Trh./Rev.Date: - - -"
    ),
    "03_Label.docx": (
        "Doküman No/Doc. Nr.: YS/D/0022    "
        "Yayın Trh./Rel. Date: 11.08.2026   "
        "Rev.No/Rev.Nr.: 00   Rev.Trh./Rev.Date: - - -"
    ),
    "04_Shipment_Statement.docx": (
        "Doküman No/Doc. Nr.: YS/D/0023    "
        "Yayın Trh./Rel. Date: 11.08.2026   "
        "Rev.No/Rev.Nr.: 00   Rev.Trh./Rev.Date: - - -"
    ),
}
QMS_CODE = {
    "01_Technical_File.docx": "YS/D/0020",
    "02_EU_DoC.docx": "YS/D/0021",
    "03_Label.docx": "YS/D/0022",
    "04_Shipment_Statement.docx": "YS/D/0023",
}


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def _set_para_text(paragraph, text: str) -> None:
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].font.size = Pt(9)
        paragraph.runs[0].font.bold = True
    else:
        run = paragraph.add_run(text)
        run.font.size = Pt(9)
        run.font.bold = True


def stamp_doc(path: Path) -> bool:
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    line = QMS_LINE[path.name]
    code = QMS_CODE[path.name]
    doc = Document(str(path))
    changed = False
    found = False

    for p in doc.paragraphs:
        t = p.text or ""
        if ("Doküman No" in t or "Doküman Tip No" in t or "Document Type No" in t) and (
            code in t or "YS/D/002" in t
        ):
            if t.strip() != line:
                _set_para_text(p, line)
                changed = True
            found = True
            break

    if not found and doc.paragraphs:
        # Insert QMS line immediately after title (paragraph 0), before IA-PPWR id line
        title = doc.paragraphs[0]
        new_p_elm = OxmlElement("w:p")
        title._p.addnext(new_p_elm)
        new_para = Paragraph(new_p_elm, title._parent)
        _set_para_text(new_para, line)
        changed = True

    for sec in doc.sections:
        for p in sec.footer.paragraphs:
            t = p.text or ""
            if "YS/D/002" in t or "Doküman" in t or "Doc. Nr" in t:
                short = f"Doküman No / Doc. Nr.: {code}"
                if t.strip() != short:
                    _set_para_text(p, short)
                    changed = True

    if changed or found:
        doc.save(str(path))
    return True


def main() -> None:
    words = sorted(
        p for p in DOC_SETS.rglob("*.docx") if not p.name.startswith("~$") and p.name in QMS_LINE
    )
    assert len(words) == 1148, len(words)
    print(f"Stamping {len(words)} Word…", flush=True)
    n = 0
    for p in words:
        stamp_doc(p)
        n += 1
        if n % 200 == 0:
            print(f"  {n}/{len(words)}", flush=True)
    print("Word stamp done", flush=True)

    # QA stamp presence
    qms = Counter()
    for p in words:
        doc = Document(str(p))
        blob = "\n".join(x.text for x in doc.paragraphs)
        code = QMS_CODE[p.name]
        line = QMS_LINE[p.name]
        if line in blob or (code in blob and "Yayın Trh" in blob and "Doküman No" in blob):
            qms[code] += 1
    print("QMS counts", dict(qms), flush=True)

    # Re-render PDFs
    from builders.phase_i.render_batch import render_docx_batch

    jobs = [(p, p.with_suffix(".pdf")) for p in words]
    print(f"Rendering {len(jobs)} PDFs…", flush=True)
    for i in range(0, len(jobs), 40):
        chunk = jobs[i : i + 40]
        render_docx_batch(chunk, progress_every=20, log=[])
        done = min(i + 40, len(jobs))
        ok = sum(1 for _, pdf in jobs[:done] if pdf.exists() and pdf.stat().st_size > 0)
        print(f"  PDF {done}/{len(jobs)} ok={ok}", flush=True)
    retry = [(d, p) for d, p in jobs if not p.exists() or p.stat().st_size == 0]
    if retry:
        print(f"Retry {len(retry)}", flush=True)
        for i in range(0, len(retry), 20):
            render_docx_batch(retry[i : i + 20], progress_every=10, log=[])

    pdfs = [p for p in DOC_SETS.rglob("*.pdf") if p.stat().st_size > 0]
    gate = (
        all(qms.get(c, 0) == 287 for c in ("YS/D/0020", "YS/D/0021", "YS/D/0022", "YS/D/0023"))
        and len(pdfs) == 1148
    )

    print("Rebuild ZIP…", flush=True)
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
        for p in FINAL.rglob("*"):
            if p.is_file() and not p.name.startswith("~$") and "PRE_YS_D" not in p.name:
                zf.write(p, p.relative_to(FINAL).as_posix())
    digest = sha256_file(ZIP_PATH)
    SHA_PATH.write_text(digest + "\n", encoding="utf-8")

    print(f"TF YS/D/0020: {qms.get('YS/D/0020',0)}/287")
    print(f"DoC YS/D/0021: {qms.get('YS/D/0021',0)}/287")
    print(f"Label YS/D/0022: {qms.get('YS/D/0022',0)}/287")
    print(f"STM YS/D/0023: {qms.get('YS/D/0023',0)}/287")
    print(f"PDF: {len(pdfs)}")
    print(f"SHA256: {digest}")
    print(f"GATE: {'PASS' if gate else 'FAIL'}")
    if not gate:
        raise SystemExit(1)


if __name__ == "__main__":
    # ensure src on path for PDF render
    import sys

    sys.path.insert(0, str(ROOT))
    sys.path.insert(0, str(ROOT / "src"))
    main()
