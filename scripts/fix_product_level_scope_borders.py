"""Add borders + 1 blank line after product-scope tables on all product-level DOCX."""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from generate_product_level_delivery import (  # noqa: E402
    PRODUCT_SETS,
    _ensure_blank_after_table,
    _table_borders,
    style_product_scope_table,
)


def _is_scope_table(table) -> bool:
    if not table.rows or len(table.columns) < 2:
        return False
    hdr = " ".join(c.text for c in table.rows[0].cells).upper()
    return ("PRODUCT CODE" in hdr or "ÜRÜN KODU" in hdr) and (
        "DESCRIPTION" in hdr or "TANIM" in hdr
    )


def fix_file(path: Path) -> int:
    doc = Document(str(path))
    n = 0
    # collect product code/desc from existing data row if present
    for table in list(doc.tables):
        if not _is_scope_table(table):
            continue
        code = table.rows[1].cells[0].text.strip() if len(table.rows) > 1 else ""
        desc = table.rows[1].cells[1].text.strip() if len(table.rows) > 1 else ""
        if not code:
            # try folder name
            code = path.parent.name
        style_product_scope_table(table, code, desc)
        _ensure_blank_after_table(doc, table)
        n += 1
    if n:
        doc.save(str(path))
    return n


def main() -> int:
    only = set(sys.argv[1:]) if len(sys.argv) > 1 else None
    folders = sorted(p for p in PRODUCT_SETS.iterdir() if p.is_dir())
    if only:
        folders = [p for p in folders if p.name in only]
    files = fixed = 0
    for i, folder in enumerate(folders, 1):
        for stem in (
            "01_Technical_File.docx",
            "02_EU_DoC.docx",
            "03_Label.docx",
            "04_Shipment_Statement.docx",
        ):
            path = folder / stem
            if not path.exists():
                continue
            files += 1
            fixed += fix_file(path)
        if i % 200 == 0:
            print(f"… {i}/{len(folders)}", flush=True)
    print(f"DONE files={files} scope_tables_fixed={fixed}", flush=True)
    # sample verify
    for code in ("1003150", "1000441", "1002873"):
        sample = PRODUCT_SETS / code / "02_EU_DoC.docx"
        if not sample.exists():
            continue
        d = Document(str(sample))
        for t in d.tables:
            if _is_scope_table(t):
                tblPr = t._tbl.tblPr
                has_b = tblPr is not None and tblPr.find(qn("w:tblBorders")) is not None
                nxt = t._tbl.getnext()
                blank = nxt is not None and nxt.tag == qn("w:p")
                print(f"sample {code}: borders={has_b} blank_after={blank}", flush=True)
                break
        break
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
