"""Replace messy QMS paragraph stamps with a clean control TABLE matching Desktop fields.

Desktop line content kept:
  Doküman No / Doc. Nr. | Yayın Trh. / Rel. Date | Rev.No / Rev.Nr. | Rev.Trh. / Rev.Date
  YS/D/002x             | 11.08.2026             | 00              | - - -

Removes floating IA-PPWR paragraph under title (IDs remain in body tables).
Cleans footer to a single short Doc. Nr. line.
"""

from __future__ import annotations

import hashlib
import re
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:\Users\burcu\Documents\YAZILIM\Inci_Aku_PPWR_PIMS")
FINAL = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL"
DOC_SETS = FINAL / "01_DOCUMENT_SETS"
ZIP_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL.zip"
SHA_PATH = ROOT / "output" / "INCI_AKU_PPWR_STARTER_CUSTOMER_DELIVERY_REV00_FINAL_SHA256.txt"

QMS = {
    "01_Technical_File.docx": "YS/D/0020",
    "02_EU_DoC.docx": "YS/D/0021",
    "03_Label.docx": "YS/D/0022",
    "04_Shipment_Statement.docx": "YS/D/0023",
}
REL_DATE = "11.08.2026"
REV_NO = "00"
REV_DATE = "- - -"

HDR = [
    "Doküman No / Doc. Nr.",
    "Yayın Trh. / Rel. Date",
    "Rev.No / Rev.Nr.",
    "Rev.Trh. / Rev.Date",
]


def sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    with p.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def _set_run(cell, text: str, *, bold: bool = False, size: int = 8) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    if not p.runs:
        r = p.add_run(text)
    else:
        p.runs[0].text = text
        r = p.runs[0]
        for extra in p.runs[1:]:
            extra.text = ""
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Tahoma"
    r.font.color.rgb = RGBColor(0x0E, 0x2A, 0x47)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "C8A24A")
        borders.append(el)
    tblPr.append(borders)


def _clear_para(p) -> None:
    for r in p.runs:
        r.text = ""


def _is_qms_para(text: str) -> bool:
    t = text or ""
    u = t.upper()
    if "YS/D/002" in t and ("DOKÜMAN" in u or "DOKUMAN" in u or "DOC. NR" in u or "DOC.NR" in u):
        return True
    if "YAYIN TRH" in u or "REL. DATE" in u:
        return True
    if "REV.TRH" in u or "REV.DATE" in u:
        return True
    return False


def _is_floating_config_id(text: str) -> bool:
    t = (text or "").strip()
    return bool(re.match(r"^IA-PPWR-(TF|DOC|LBL|STM)-", t))


def fix_doc(path: Path) -> None:
    code = QMS[path.name]
    doc = Document(str(path))

    # Remove existing QMS control tables we may have inserted (4-col with Doküman No header)
    for table in list(doc.tables):
        if not table.rows:
            continue
        hdr = " ".join(c.text for c in table.rows[0].cells)
        if "Doküman No" in hdr and "Yayın Trh" in hdr and len(table.columns) == 4:
            table._tbl.getparent().remove(table._tbl)

    # Clean body paragraphs after title
    title = doc.paragraphs[0] if doc.paragraphs else None
    to_clear = []
    for i, p in enumerate(doc.paragraphs):
        if i == 0:
            continue
        t = p.text or ""
        if _is_qms_para(t) or _is_floating_config_id(t):
            to_clear.append(p)
            continue
        # stop at first real section heading
        if t.strip().startswith("01 ") or t.strip().startswith("İZLENEBİLİRLİK") or t.strip().startswith("01\t"):
            break
        if not t.strip():
            # leave one blank later
            to_clear.append(p)

    for p in to_clear:
        _clear_para(p)
        # remove empty paragraph element entirely
        el = p._p
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    # Insert control table right after title
    if title is not None:
        # create table via document then move XML after title
        table = doc.add_table(rows=2, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_borders(table)
        vals = [code, REL_DATE, REV_NO, REV_DATE]
        for ci, h in enumerate(HDR):
            _set_run(table.rows[0].cells[ci], h, bold=True, size=7)
            _shade(table.rows[0].cells[ci], "0E2A47")
            for r in table.rows[0].cells[ci].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_run(table.rows[1].cells[ci], vals[ci], bold=True, size=9)
            _shade(table.rows[1].cells[ci], "F7F3EA")
        # move table after title
        title._p.addnext(table._tbl)
        # remove the table from end-of-body position (add_table appended it)
        # After addnext, the same tbl element is moved — good (addnext relocates)

        # ensure one blank paragraph after table before content
        blank = OxmlElement("w:p")
        table._tbl.addnext(blank)

    # Footer: single clean line (match short form)
    short = f"Doküman No / Doc. Nr.: {code}"
    for sec in doc.sections:
        footer = sec.footer
        # clear all footer paragraphs then set first
        paras = list(footer.paragraphs)
        if not paras:
            footer.add_paragraph(short)
        else:
            _clear_para(paras[0])
            run = paras[0].add_run(short)
            run.font.size = Pt(8)
            run.font.name = "Tahoma"
            for p in paras[1:]:
                _clear_para(p)
                el = p._p
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)

    doc.save(str(path))


def main() -> None:
    words = sorted(
        p for p in DOC_SETS.rglob("*.docx") if not p.name.startswith("~$") and p.name in QMS
    )
    assert len(words) == 1148, len(words)
    print(f"Fixing QMS table on {len(words)} docs…", flush=True)
    for i, p in enumerate(words, 1):
        fix_doc(p)
        if i % 200 == 0:
            print(f"  {i}/{len(words)}", flush=True)

    # spot check
    sample = DOC_SETS / "ST-012-EUR-01" / "02_EU_DoC.docx"
    d = Document(str(sample))
    print("paras0-4:", [repr(p.text[:80]) for p in d.paragraphs[:5]])
    print("table0:", [[c.text for c in r.cells] for r in d.tables[0].rows] if d.tables else None)

    qms = Counter()
    for p in words:
        doc = Document(str(p))
        code = QMS[p.name]
        ok = False
        for t in doc.tables:
            blob = " | ".join(c.text for r in t.rows for c in r.cells)
            if code in blob and "Doküman No" in blob and "Yayın Trh" in blob:
                ok = True
                break
        if ok:
            qms[code] += 1
    print("QMS table counts", dict(qms), flush=True)

    import sys

    sys.path.insert(0, str(ROOT))
    sys.path.insert(0, str(ROOT / "src"))
    from builders.phase_i.render_batch import render_docx_batch

    jobs = [(p, p.with_suffix(".pdf")) for p in words]
    print(f"PDF {len(jobs)}…", flush=True)
    for i in range(0, len(jobs), 40):
        render_docx_batch(jobs[i : i + 40], progress_every=40, log=[])
        done = min(i + 40, len(jobs))
        print(f"  PDF {done}/{len(jobs)}", flush=True)
    retry = [(d, p) for d, p in jobs if not p.exists() or p.stat().st_size == 0]
    if retry:
        for i in range(0, len(retry), 20):
            render_docx_batch(retry[i : i + 20], progress_every=20, log=[])

    pdf_n = sum(1 for _, p in jobs if p.exists() and p.stat().st_size > 0)
    gate = all(qms.get(c, 0) == 287 for c in QMS.values()) and pdf_n == 1148

    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
        for p in FINAL.rglob("*"):
            if p.is_file() and not p.name.startswith("~$"):
                zf.write(p, p.relative_to(FINAL).as_posix())
    digest = sha256_file(ZIP_PATH)
    SHA_PATH.write_text(digest + "\n", encoding="utf-8")
    print(f"PDF {pdf_n} GATE {'PASS' if gate else 'FAIL'} SHA {digest}")
    if not gate:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
