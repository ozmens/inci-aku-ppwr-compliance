"""Phase O3 Part A — remove Nominal Load from TF Golden + regenerate 247 TFs."""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Any, Callable

from docx import Document

from builders.phase_g.merge_engine import merge_document
from builders.phase_g.pims_loader import ProductionDocumentLoader
from builders.phase_g.runtime_template_builder import build_runtime_templates, sha256_file
from builders.phase_g.tokens import GOLDEN_FILES, RUNTIME_FILES
from builders.phase_i.service import FAMILY_FOLDER
from models.technical_file import Article5Assessment
from services.document_context_factory import DocumentContextFactory
from utils.constants import ARTICLE5_BASIS_LABEL

NOMINAL_PATTERNS = (
    re.compile(r"Nominal\s+yük", re.IGNORECASE),
    re.compile(r"Nominal\s+load", re.IGNORECASE),
)


def remove_nominal_load_row_from_docx(path: Path) -> bool:
    """Delete the dedicated Nominal yük / Nominal load table row."""
    doc = Document(str(path))
    removed = False
    for table in doc.tables:
        for row in list(table.rows):
            label = (row.cells[0].text or "").replace("\n", " ")
            if "Nominal yük" in label or "Nominal load" in label:
                tr = row._tr
                tr.getparent().remove(tr)
                removed = True
                break
        if removed:
            break
    if removed:
        doc.save(str(path))
    return removed


def update_golden_technical_file(golden_dir: Path) -> dict[str, Any]:
    """Update Golden TF (+ alias twin). DoC/Label/STM untouched. Idempotent."""
    primary = golden_dir / GOLDEN_FILES["TECHNICAL_FILE"]
    alias = golden_dir / "01_IA-PPWR-TF-ST-012-EUR-01-R00.docx"
    before = sha256_file(primary)
    hits_before = scan_nominal_in_docx(primary)
    removed = False
    if hits_before > 0:
        removed = remove_nominal_load_row_from_docx(primary)
        if not removed:
            raise RuntimeError(f"Nominal Load row not found in {primary.name}")
    if alias.exists():
        shutil.copy2(primary, alias)
    after = sha256_file(primary)
    hits = scan_nominal_in_docx(primary)
    return {
        "golden_file": primary.name,
        "removed": removed,
        "already_clean": hits_before == 0,
        "sha256_before": before,
        "sha256_after": after,
        "nominal_hits_in_golden": hits,
        "alias_synced": alias.exists(),
    }


def rebuild_tf_runtime_only(golden_dir: Path, runtime_dir: Path) -> dict[str, Any]:
    """Rebuild runtimes from Golden (non-TF goldens unchanged → stable hashes)."""
    inv = build_runtime_templates(golden_dir, runtime_dir)
    tf_rt = runtime_dir / RUNTIME_FILES["TECHNICAL_FILE"]
    text = Document(str(tf_rt)).element.xml
    return {
        "inventory": {
            k: {
                "golden_sha256": v["golden_sha256"],
                "runtime_sha256": v["runtime_sha256"],
                "sample_leak_count": v["sample_leak_count"],
            }
            for k, v in inv.items()
        },
        "tf_runtime_nominal_xml_hits": len(
            re.findall(r"Nominal\s+(yük|load)", text, flags=re.I)
        ),
        "tf_golden_sha256": inv["TECHNICAL_FILE"]["golden_sha256"],
    }


def scan_nominal_in_docx(path: Path) -> int:
    hits = 0
    try:
        import zipfile

        with zipfile.ZipFile(path) as z:
            for name in z.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    raw = z.read(name).decode("utf-8", errors="ignore")
                    for p in NOMINAL_PATTERNS:
                        hits += len(p.findall(raw))
    except Exception:
        doc = Document(str(path))
        blob = "\n".join(p.text or "" for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                blob += "\n" + " | ".join(c.text or "" for c in row.cells)
        for p in NOMINAL_PATTERNS:
            hits += len(p.findall(blob))
    return hits


def smoke_render_docx_batch(
    paths: list[Path],
    *,
    progress_every: int = 50,
    log: Callable[[str], None] | None = None,
) -> dict[str, Any]:
    """Open each DOCX in Word COM (no PDF) — validates renderability."""
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    ok = fail = 0
    errors: list[str] = []
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        for i, docx_path in enumerate(paths, start=1):
            doc = None
            try:
                doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
                _ = int(doc.ComputeStatistics(2))
                ok += 1
            except Exception as exc:  # noqa: BLE001
                fail += 1
                errors.append(f"{docx_path.name}: {exc}")
            finally:
                if doc is not None:
                    try:
                        doc.Close(False)
                    except Exception:
                        pass
            if log and (i % progress_every == 0 or i == len(paths)):
                log(f"  TF smoke render {i}/{len(paths)} ok={ok} fail={fail}")
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
    return {"ok": ok, "fail": fail, "errors": errors[:20]}


def regenerate_technical_files_only(
    project_root: Path,
    *,
    log: Callable[[str], None] | None = None,
) -> dict[str, Any]:
    """Overwrite only 01_Technical_File.docx under PHASE_I_FINAL packs."""

    def _log(msg: str) -> None:
        if log:
            log(msg)

    production = project_root / "output" / "INCI_AKU_PPWR_PIMS_Rev00_FINAL.xlsx"
    if not production.exists():
        production = project_root / "output" / "INCI_AKU_PPWR_PIMS_Rev00_PRODUCTION.xlsx"
    runtime_dir = project_root / "templates" / "word_runtime"
    out = project_root / "output" / "PHASE_I_FINAL"
    runtime_tf = runtime_dir / RUNTIME_FILES["TECHNICAL_FILE"]

    loader = ProductionDocumentLoader(production)
    loader.open()
    factory = DocumentContextFactory()
    tf_paths: list[Path] = []
    hash_map: dict[str, str] = {}
    errors: list[str] = []

    try:
        packs: list[tuple[str, str, Path]] = []
        for family, folder in FAMILY_FOLDER.items():
            root = out / folder
            if not root.exists():
                continue
            for d in sorted(p for p in root.iterdir() if p.is_dir()):
                packs.append((family, d.name, d))
        if len(packs) != 247:
            raise RuntimeError(f"Expected 247 packs, found {len(packs)}")

        _log(f"Regenerating {len(packs)} Technical Files…")
        for idx, (family, set_code, pack_dir) in enumerate(packs, start=1):
            try:
                cfg, products = loader.load_configuration(set_code)
                ctx = factory.build(
                    cfg,
                    products=products,
                    article5=Article5Assessment(
                        basis_label=ARTICLE5_BASIS_LABEL,
                        evidence_references=[],
                        numerical_claim=None,
                    ),
                )
                out_path = pack_dir / "01_Technical_File.docx"
                merge_document(
                    runtime_tf,
                    out_path,
                    ctx.for_technical_file(),
                    for_technical_file=True,
                )
                digest = sha256_file(out_path)
                rel = f"{FAMILY_FOLDER[family]}/{set_code}/01_Technical_File.docx"
                hash_map[rel.replace("\\", "/")] = digest
                tf_paths.append(out_path)
            except Exception as exc:  # noqa: BLE001
                errors.append(f"{set_code}: {exc}")
            if idx % 25 == 0 or idx == len(packs):
                _log(f"  TF merge {idx}/{len(packs)}")
    finally:
        loader.close()

    _log(f"Smoke-rendering {len(tf_paths)} TFs via Word COM…")
    render = smoke_render_docx_batch(tf_paths, progress_every=50, log=_log)

    nominal_hits = sum(scan_nominal_in_docx(p) for p in tf_paths)
    sample_checks: dict[str, Any] = {}
    for label, rel in (
        ("STARTER", "01_STARTER/ST-051-STD-01/01_Technical_File.docx"),
        ("INDUSTRIAL", "02_INDUSTRIAL/IND-24V-01/01_Technical_File.docx"),
        ("CONTAINER", "03_CONTAINER/CNT-20-STD-01/01_Technical_File.docx"),
    ):
        p = out / rel
        sample_checks[label] = {
            "path": str(p),
            "exists": p.exists(),
            "nominal_hits": scan_nominal_in_docx(p) if p.exists() else -1,
            "sha256": sha256_file(p) if p.exists() else None,
            "row_labels_table0": _table0_labels(p) if p.exists() else [],
        }

    return {
        "tf_count": len(tf_paths),
        "nominal_hits_total": nominal_hits,
        "render_ok": render["ok"],
        "render_failures": render["fail"],
        "render_errors": render["errors"],
        "hash_map": hash_map,
        "sample_checks": sample_checks,
        "errors": errors,
        "runtime_tf": str(runtime_tf),
    }


def _table0_labels(path: Path) -> list[str]:
    doc = Document(str(path))
    if not doc.tables:
        return []
    return [(r.cells[0].text or "").replace("\n", " / ")[:60] for r in doc.tables[0].rows]


def update_document_library_tf_hashes(
    workbook_path: Path, hash_map: dict[str, str]
) -> dict[str, int]:
    """Update FILE_HASH only for Technical File rows."""
    from openpyxl import load_workbook

    wb = load_workbook(workbook_path)
    ws = wb["DOCUMENT_LIBRARY"]
    headers = [ws.cell(1, c).value for c in range(1, (ws.max_column or 1) + 1)]
    uri_col = headers.index("FILE_URI") + 1
    hash_col = headers.index("FILE_HASH") + 1

    updated = same = non_tf = skipped = 0
    for r in range(2, (ws.max_row or 1) + 1):
        uri = str(ws.cell(r, uri_col).value or "").replace("\\", "/")
        if not uri.lower().endswith("01_technical_file.docx"):
            non_tf += 1
            continue
        key = uri.lstrip("./")
        new_hash = hash_map.get(key)
        if new_hash is None:
            for k, v in hash_map.items():
                if key.endswith(k) or k.endswith(key):
                    new_hash = v
                    break
        if new_hash is None:
            skipped += 1
            continue
        old = ws.cell(r, hash_col).value
        ws.cell(r, hash_col).value = new_hash
        if old != new_hash:
            updated += 1
        else:
            same += 1
    wb.save(workbook_path)
    wb.close()
    return {
        "tf_hashes_updated": updated,
        "tf_hashes_same": same,
        "non_tf_rows_untouched": non_tf,
        "tf_rows_skipped_missing_map": skipped,
    }


def update_phase_i_manifest_tf_hashes(phase_i: Path, hash_map: dict[str, str]) -> int:
    """Best-effort update of manifest xlsx if present."""
    manifest_xlsx = phase_i / "90_MANIFEST" / "INCI_AKU_PPWR_DOCUMENT_MANIFEST.xlsx"
    if not manifest_xlsx.exists():
        return 0
    from openpyxl import load_workbook

    wb = load_workbook(manifest_xlsx)
    ws = wb.active
    headers = [ws.cell(1, c).value for c in range(1, (ws.max_column or 1) + 1)]

    def find(*names: str) -> int | None:
        for n in names:
            if n in headers:
                return headers.index(n) + 1
        return None

    path_col = find("Relative Path", "Document Path", "File Path", "Path", "FILE_URI")
    hash_col = find("SHA-256", "SHA256", "FILE_HASH", "Hash")
    type_col = find("Document Type", "Type", "DOCUMENT_TYPE")
    if not path_col or not hash_col:
        wb.close()
        return 0
    updated = 0
    for r in range(2, (ws.max_row or 1) + 1):
        p = str(ws.cell(r, path_col).value or "").replace("\\", "/")
        is_tf = p.lower().endswith("01_technical_file.docx")
        if type_col:
            t = str(ws.cell(r, type_col).value or "").upper()
            is_tf = is_tf or "TECHNICAL" in t
        if not is_tf:
            continue
        for k, v in hash_map.items():
            if p.endswith(k) or k in p or p.endswith(k.split("/")[-1]):
                # prefer full path match
                if k in p or p.endswith(k):
                    ws.cell(r, hash_col).value = v
                    updated += 1
                    break
    wb.save(manifest_xlsx)
    wb.close()
    return updated


def patch_phase_g_golden_hash_constant(service_py: Path, new_hash: str) -> None:
    """Update PHASE_G_GOLDEN_HASHES TECHNICAL_FILE entry in phase_i/service.py."""
    text = service_py.read_text(encoding="utf-8")
    pat = re.compile(
        r'("TECHNICAL_FILE":\s*")([0-9a-f]{64})(")',
        re.IGNORECASE,
    )
    new_text, n = pat.subn(rf"\g<1>{new_hash}\g<3>", text, count=1)
    if n != 1:
        raise RuntimeError("Could not patch PHASE_G_GOLDEN_HASHES TECHNICAL_FILE")
    service_py.write_text(new_text, encoding="utf-8")
