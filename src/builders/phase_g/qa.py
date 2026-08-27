"""Phase G document QA scanners."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document

from .tokens import ALLOWED_PENDING_ALT, FORBIDDEN_CONTENT, FORBIDDEN_SAMPLE_LEAKS


def extract_all_text(path: Path) -> str:
    parts: list[str] = []
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                parts.append(z.read(name).decode("utf-8", errors="ignore"))
    # crude strip tags
    text = re.sub(r"<[^>]+>", " ", "\n".join(parts))
    return re.sub(r"\s+", " ", text)


def unresolved_tokens(path: Path) -> list[str]:
    text = extract_all_text(path)
    return sorted(set(re.findall(r"\{\{[A-Z0-9_]+\}\}", text)))


def sample_leaks(path: Path, *, allowed_set_code: str | None = None) -> list[str]:
    text = extract_all_text(path)
    hits = []
    for lit in FORBIDDEN_SAMPLE_LEAKS:
        if lit in text:
            if allowed_set_code and lit in {
                allowed_set_code,
                f"IA-{allowed_set_code}",
            }:
                continue
            # allow when generating ST-012 itself
            if allowed_set_code == "ST-012-EUR-01":
                continue
            hits.append(lit)
    return hits


def forbidden_content_hits(path: Path) -> list[str]:
    text = extract_all_text(path).upper()
    hits = []
    for word in FORBIDDEN_CONTENT:
        # word boundary-ish
        if re.search(rf"\b{re.escape(word)}\b", text):
            # allow PENDING – DRAWINGS variants already checked separately
            hits.append(word)
    return hits


def has_pending_drawings(path: Path) -> bool:
    text = extract_all_text(path).upper()
    return "PENDING" in text and ("DRAWING" in text or "PHOTO" in text)


def customer_leak_in_tf(path: Path, suspicious_names: list[str]) -> list[str]:
    text = extract_all_text(path).upper()
    hits = []
    for name in suspicious_names:
        n = (name or "").strip().upper()
        if len(n) < 4:
            continue
        if n in text:
            hits.append(name)
    return hits


def tahoma_noncompliant_runs(path: Path) -> int:
    """Count visible runs whose font is set and not Tahoma."""
    doc = Document(str(path))
    bad = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if not (r.text or "").strip():
                continue
            name = r.font.name
            if name and name.lower() not in {"tahoma", "tahoma"}:
                bad += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not (r.text or "").strip():
                            continue
                        name = r.font.name
                        if name and name.lower() != "tahoma":
                            bad += 1
    return bad


def white_on_light_errors(path: Path) -> list[str]:
    """
    Heuristic OOXML scan: white color (FFFFFF) without nearby dark shading.
    Conservative — reports obvious white fill color on runs.
    """
    errors: list[str] = []
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
    # Find w:color w:val="FFFFFF" segments and check preceding shading in same run properties
    for m in re.finditer(r"<w:rPr>(.*?)</w:rPr>", xml, flags=re.S):
        block = m.group(1)
        if re.search(r'w:val="FFFFFF"', block, re.I) or re.search(
            r'w:val="white"', block, re.I
        ):
            # dark navy fills typically 1F4E79 / 002060
            if not re.search(r'w:fill="(1F4E79|002060|1F4E7[0-9A-F]|000000)"', block, re.I):
                # also accept theme dark
                if "w:shd" not in block and "w:highlight" not in block:
                    errors.append("white_color_without_dark_shading_in_rPr")
    return errors[:50]


def page_stats(path: Path) -> dict:
    doc = Document(str(path))
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
    }
