"""Build runtime Word templates from populated Golden masters (tokenize sample data)."""

from __future__ import annotations

import hashlib
import io
import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from lxml import etree

from .tokens import (
    FORBIDDEN_SAMPLE_LEAKS,
    GOLDEN_FILES,
    RUNTIME_FILES,
    SAMPLE_LITERAL_MAP,
    token,
)


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _replace_in_xml(xml: str) -> str:
    # longest-first already ordered in SAMPLE_LITERAL_MAP
    out = xml
    for lit, tok in SAMPLE_LITERAL_MAP:
        # Skip overly broad 12 adet/pcs in XML global replace — handled in cell-aware pass
        if lit in {"12 adet", "12 pcs"}:
            continue
        out = out.replace(lit, tok)
    return out


def _tokenize_zip_copy(src: Path, dst: Path) -> None:
    buf = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    text = data.decode("utf-8")
                except UnicodeDecodeError:
                    zout.writestr(item, data)
                    continue
                if "document.xml" in item.filename or "header" in item.filename or "footer" in item.filename:
                    text = _replace_in_xml(text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    dst.write_bytes(buf.getvalue())


def _is_bom_table(table: Table) -> bool:
    if len(table.columns) < 6 or len(table.rows) < 3:
        return False
    header = " ".join(c.text for c in table.rows[0].cells).upper()
    return "CODE" in header and ("LINE WEIGHT" in header or "SATIR" in header or "BİRİM" in header or "UNIT" in header)


def _is_product_table(table: Table) -> bool:
    if len(table.rows) < 1:
        return False
    header = " ".join(c.text for c in table.rows[0].cells).upper()
    if "CONTROLLED PRODUCTS" in header or "KONTROLL" in header:
        return True
    if ("PRODUCT CODE" in header or "ÜRÜN KODU" in header or "URUN KODU" in header) and (
        "DESCRIPTION" in header or "TANIM" in header
    ):
        return True
    return False


def _clear_bom_data_keep_total(table: Table) -> None:
    """Keep header + TOTAL; remove sample BOM lines; leave one template data row."""
    # Identify total row
    total_idx = None
    for i, row in enumerate(table.rows):
        t0 = row.cells[0].text.upper()
        if "TOPLAM" in t0 or "TOTAL" in t0:
            total_idx = i
            break
    if total_idx is None or total_idx <= 1:
        return
    # Keep row 0 header, row 1 as clone template (will clear content to tokens), delete 2..total-1
    # First set row 1 as empty template markers
    tmpl = table.rows[1]
    markers = [
        "{{BOM_CODE}}",
        "{{BOM_DESC}}",
        "{{BOM_QTY}}",
        "{{BOM_UOM}}",
        "{{BOM_UNIT_WT}}",
        "{{BOM_LINE_WT}}",
    ]
    for ci, cell in enumerate(tmpl.cells):
        if ci < len(markers):
            _set_cell_text(cell, markers[ci])
    # Delete rows between template and total (from end)
    # python-docx: delete row via XML
    tbl = table._tbl
    rows = list(tbl.findall(qn("w:tr")))
    # After edits, re-find total
    # Delete indices 2 .. total_idx-1
    for idx in range(total_idx - 1, 1, -1):
        rows = list(tbl.findall(qn("w:tr")))
        if idx < len(rows):
            tbl.remove(rows[idx])
    # Ensure total tare token
    rows = list(tbl.findall(qn("w:tr")))
    last = table.rows[-1]
    _set_cell_text(last.cells[-1], "{{TOTAL_TARE_KG}}")


def _clear_product_sample(table: Table) -> None:
    if len(table.rows) < 2:
        return
    # Keep header; replace body with single token row (two columns: code | description)
    tbl = table._tbl
    rows = list(tbl.findall(qn("w:tr")))
    for idx in range(len(rows) - 1, 0, -1):
        tbl.remove(rows[idx])
    header_tr = tbl.findall(qn("w:tr"))[0]
    new_tr = deepcopy(header_tr)
    tbl.append(new_tr)
    hdr = table.rows[0]
    if len(hdr.cells) >= 2:
        _set_cell_text(hdr.cells[0], "ÜRÜN KODU\nPRODUCT CODE")
        _set_cell_text(hdr.cells[1], "ÜRÜN TANIMI\nPRODUCT DESCRIPTION")
    data_row = table.rows[1]
    _set_cell_text(data_row.cells[0], "{{PRODUCT_CODE}}")
    if len(data_row.cells) > 1:
        _set_cell_text(data_row.cells[1], "{{PRODUCT_DESCRIPTION}}")


def _set_cell_text(cell, text: str) -> None:
    # Clear paragraphs and set one run preserving first paragraph style if possible
    for p in cell.paragraphs:
        for r in list(p.runs):
            r.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
        for extra in cell.paragraphs[1:]:
            for r in extra.runs:
                r.text = ""


def _postprocess_docx_tables(path: Path) -> None:
    doc = Document(str(path))
    for table in doc.tables:
        if _is_bom_table(table):
            _clear_bom_data_keep_total(table)
        elif _is_product_table(table):
            _clear_product_sample(table)
    doc.save(str(path))


def build_runtime_templates(
    golden_dir: Path,
    runtime_dir: Path,
) -> dict[str, dict[str, str]]:
    """
    Create tokenized runtime templates from Golden masters.
    Returns hash inventory for golden + runtime.
    """
    runtime_dir.mkdir(parents=True, exist_ok=True)
    inventory: dict[str, dict[str, str]] = {}
    for kind, golden_name in GOLDEN_FILES.items():
        src = golden_dir / golden_name
        if not src.exists():
            raise FileNotFoundError(f"Missing Golden template: {src}")
        dst = runtime_dir / RUNTIME_FILES[kind]
        # Step 1: ZIP-level literal tokenization
        _tokenize_zip_copy(src, dst)
        # Step 2: table structure cleanup via python-docx
        _postprocess_docx_tables(dst)
        # Step 3: leakage scan on runtime (sample identity must be gone)
        leaks = scan_sample_leaks(dst)
        inventory[kind] = {
            "golden_file": golden_name,
            "golden_sha256": sha256_file(src),
            "runtime_file": RUNTIME_FILES[kind],
            "runtime_sha256": sha256_file(dst),
            "sample_leak_count": str(len(leaks)),
            "sample_leaks": "; ".join(leaks[:20]),
        }
        if leaks:
            raise RuntimeError(
                f"Runtime template {dst.name} still contains sample identity: {leaks[:10]}"
            )
    return inventory


def scan_sample_leaks(path: Path) -> list[str]:
    hits: list[str] = []
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            text = z.read(name).decode("utf-8", errors="ignore")
            # strip tokens so we don't false-positive
            plain = re.sub(r"\{\{[^}]+\}\}", "", text)
            for lit in FORBIDDEN_SAMPLE_LEAKS:
                if lit in plain:
                    hits.append(f"{name}:{lit}")
    return hits
