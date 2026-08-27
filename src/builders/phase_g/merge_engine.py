"""Merge DocumentContext into runtime Word templates."""

from __future__ import annotations

import io
import re
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

from models.document_context import DocumentContext

from .tokens import RUNTIME_FILES


def context_token_values(ctx: DocumentContext, *, for_technical_file: bool = False) -> dict[str, str]:
    cfg = ctx.configuration
    ids = ctx.document_ids
    tare_kg = round(ctx.total_tare_g / 1000.0, 4)
    tare_str = f"{tare_kg:.4f}"
    qty = cfg.extras.get("nominal_product_qty")
    if qty is None:
        # try parse from name
        qty = cfg.extras.get("product_qty", "")
    qty_s = str(qty) if qty not in (None, "") else ""
    product_lines = []
    for p in ctx.products:
        bit = p.product_code
        if getattr(p, "product_name", None):
            bit = f"{p.product_code} • {p.product_name}"
        product_lines.append(bit)
    if for_technical_file:
        # strip commercial from product display — code + technical desc only
        product_lines = []
        for p in ctx.products:
            name = getattr(p, "product_name", "") or ""
            product_lines.append(f"{p.product_code} • {name}".strip(" •"))

    return {
        "CONFIG_SET_CODE": cfg.packaging_set_code,
        "CONFIG_ID": cfg.final_configuration_id,
        "SOURCE_CONFIG_ID": (cfg.lineage.source_configuration_id or ""),
        "VARIANT_BASIS_TR": cfg.variant_basis_tr,
        "VARIANT_BASIS_EN": cfg.variant_basis_en,
        "VARIANT_BASIS_PAIR": f"{cfg.variant_basis_tr}\n{cfg.variant_basis_en}",
        "PACKAGING_DESCRIPTION_TR": cfg.extras.get("name_tr") or cfg.name,
        "PACKAGING_DESCRIPTION_EN": cfg.extras.get("name_en") or cfg.name,
        "PACKAGING_DESCRIPTION_PAIR": (
            f"{cfg.extras.get('name_tr') or cfg.name}\n{cfg.extras.get('name_en') or cfg.name}"
        ),
        "NOMINAL_LOAD_TR": cfg.extras.get("nominal_load_tr")
        or (f"{qty_s} adet / palet" if qty_s else ""),
        "NOMINAL_LOAD_EN": cfg.extras.get("nominal_load_en")
        or (f"{qty_s} units / pallet" if qty_s else ""),
        "NOMINAL_LOAD_PAIR": "",
        "PRODUCT_QTY_PAIR": f"{qty_s} adet / {qty_s} pcs" if qty_s else "",
        "TOTAL_TARE_KG": f"{tare_str} kg",
        "TOTAL_TARE_KG_PLAIN": tare_str,
        "TF_ID": ids.technical_file_id,
        "DOC_ID": ids.doc_id,
        "LABEL_ID": ids.label_id,
        "STM_ID": ids.statement_id,
        "REVISION": ids.revision_display,
        "REVISION_DATE": cfg.extras.get("revision_date", "08.08.2026"),
        "REVISION_PAIR": f"{ids.revision_display} / {cfg.extras.get('revision_date', '08.08.2026')}",
        "LEGAL_NAME": ctx.legal_entity.legal_name,
        "LEGAL_ADDRESS": ctx.legal_entity.address,
        "LEGAL_EMAIL": ctx.legal_entity.email,
        "LEGAL_WEBSITE": ctx.legal_entity.website,
        "LEGAL_PHONE": cfg.extras.get("phone", ""),
        "ARTICLE5_BASIS": (
            ctx.article5.basis_label
            if ctx.article5
            else "REV00 CURRENT EVIDENCE BASIS — ARTICLE 5 ASSESSMENT BASIS"
        ),
        "ANNEX_DRAWINGS_STATUS": ctx.annex_drawings_status,
        "TRACEABILITY_KEY": (
            f"{cfg.packaging_set_code} → {cfg.final_configuration_id} → "
            f"{cfg.lineage.source_configuration_id or ''}"
        ),
        "PRODUCT_LINES": "\n".join(product_lines) if product_lines else "—",
        "SHIPMENT_NO": "",
        "SHIPMENT_DATE": "",
        "CUSTOMER_OEM": "" if for_technical_file else (ctx.customer_name or ""),
        "DESTINATION": "",
        "INCOTERM": "",
        "PACKING_LIST_REF": "",
        "PACKAGING_LOT": "",
    }


def _apply_token_map_xml(xml: str, values: dict[str, str]) -> str:
    out = xml
    # longest token names first
    for key in sorted(values.keys(), key=len, reverse=True):
        out = out.replace("{{" + key + "}}", _xml_escape(values[key]))
    return out


def _xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def merge_document(
    runtime_path: Path,
    output_path: Path,
    ctx: DocumentContext,
    *,
    for_technical_file: bool = False,
) -> Path:
    values = context_token_values(ctx, for_technical_file=for_technical_file)
    if not values["NOMINAL_LOAD_PAIR"]:
        values["NOMINAL_LOAD_PAIR"] = (
            f"{values['NOMINAL_LOAD_TR']}\n{values['NOMINAL_LOAD_EN']}"
        ).strip()

    # First: XML token replace in zip copy
    buf = io.BytesIO()
    with zipfile.ZipFile(runtime_path, "r") as zin, zipfile.ZipFile(
        buf, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    text = data.decode("utf-8")
                    text = _apply_token_map_xml(text, values)
                    data = text.encode("utf-8")
                except UnicodeDecodeError:
                    pass
            zout.writestr(item, data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_bytes(buf.getvalue())

    # Second: expand BOM / product tables
    doc = Document(str(output_path))
    found_product = False
    for table in doc.tables:
        if _is_bom_table(table):
            _fill_bom_table(table, ctx)
        elif _is_product_table(table):
            _fill_product_table(table, ctx, for_technical_file=for_technical_file)
            found_product = True
    # DoC: ensure a 2-column product-scope table exists when products are linked
    if (not for_technical_file) and (not found_product) and ctx.products:
        _insert_doc_product_scope_table(doc, ctx)
    doc.save(str(output_path))
    return output_path


def _is_bom_table(table: Table) -> bool:
    if len(table.columns) < 6 or len(table.rows) < 2:
        return False
    header = " ".join(c.text for c in table.rows[0].cells).upper()
    return "CODE" in header and ("UNIT" in header or "BİRİM" in header or "LINE" in header or "SATIR" in header or "WEIGHT" in header)


def _is_product_table(table: Table) -> bool:
    if len(table.rows) < 1:
        return False
    header = " ".join(c.text for c in table.rows[0].cells).upper()
    if "CONTROLLED PRODUCTS" in header or "KONTROLL" in header:
        return True
    # two-column product scope after premium fix
    if ("PRODUCT CODE" in header or "ÜRÜN KODU" in header or "URUN KODU" in header) and (
        "PRODUCT DESCRIPTION" in header or "ÜRÜN TANIMI" in header or "URUN TANIMI" in header or "DESCRIPTION" in header
    ):
        return True
    return False


def _set_cell_text(cell, text: str) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
        else:
            run = p.add_run(text)
            run.font.name = "Tahoma"
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
            except Exception:
                pass
        for extra in cell.paragraphs[1:]:
            for r in extra.runs:
                r.text = ""


def _fill_bom_table(table: Table, ctx: DocumentContext) -> None:
    tbl = table._tbl
    rows = list(tbl.findall(qn("w:tr")))
    if len(rows) < 3:
        return
    header_tr = rows[0]
    template_tr = rows[1]
    total_tr = rows[-1]
    # remove middle rows except we'll rebuild from template
    for tr in rows[1:-1]:
        tbl.remove(tr)
    # insert one row per BOM line before total
    total_tr = tbl.findall(qn("w:tr"))[-1]
    for line in ctx.bom_lines:
        new_tr = deepcopy(template_tr)
        total_tr.addprevious(new_tr)
    # remove leftover template if still present as first data? we removed all middle then added
    # Actually template was removed with middle; we deepcopied before remove - good.

    # Refill via python-docx API
    # rows now: header, N data, total
    data_rows = table.rows[1:-1]
    for row, line in zip(data_rows, ctx.bom_lines):
        group = line.extras.get("component_group", "") if hasattr(line, "extras") else ""
        # PackagingConfigurationLine may not have extras - use notes parsing
        code = line.component_erp_code
        desc = line.component_name or ""
        # bilingual desc if stored
        desc_tr = getattr(line, "name_tr", None) or desc
        desc_en = getattr(line, "name_en", None) or desc
        if desc_tr != desc_en:
            desc_cell = f"{desc_tr}\n{desc_en}"
        else:
            desc_cell = desc
        qty = line.quantity
        uom = (line.uom or "PCS").upper()
        uom_pair = {"ADT": "adet / pcs", "PCS": "adet / pcs", "M": "m / m", "KG": "kg / kg"}.get(
            uom, f"{uom.lower()} / {uom.lower()}"
        )
        unit_kg = None
        line_kg = None
        if line.weight_g is not None:
            # weight_g is per unit for WeightService; line = qty * weight_g
            unit_kg = line.weight_g / 1000.0
            line_kg = (line.weight_g * float(line.quantity)) / 1000.0
            if uom == "KG" and abs(line.weight_g - 1000.0) < 1e-6:
                unit_disp = "Kütle esaslı / Mass-based"
            else:
                unit_disp = f"{unit_kg:.4f} kg"
        else:
            unit_disp = "—"
            line_kg = 0.0
        vals = [
            code,
            desc_cell,
            f"{qty:g}",
            uom_pair,
            unit_disp,
            f"{line_kg:.4f} kg",
        ]
        for ci, val in enumerate(vals):
            if ci < len(row.cells):
                _set_cell_text(row.cells[ci], val)
    # total
    tare = f"{round(ctx.total_tare_g/1000.0, 4):.4f} kg"
    _set_cell_text(table.rows[-1].cells[-1], tare)


def _set_run_bold(cell, bold: bool) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold
            r.font.name = "Tahoma"
            try:
                r.font.size = r.font.size  # preserve
            except Exception:
                pass


def _set_product_table_widths(table: Table) -> None:
    """Product Code ~25%, Product Description ~75% of usable width."""
    try:
        from docx.shared import Cm

        table.autofit = False
        code_w = Cm(3.8)
        desc_w = Cm(12.2)
        for row in table.rows:
            if len(row.cells) >= 2:
                row.cells[0].width = code_w
                row.cells[1].width = desc_w
    except Exception:
        pass


def _fill_product_table(table: Table, ctx: DocumentContext, *, for_technical_file: bool) -> None:
    """Two real columns: Product Code | Product Description (never concatenated)."""
    tbl = table._tbl
    rows = list(tbl.findall(qn("w:tr")))
    if not rows:
        return
    header_tr = rows[0]
    # remove body
    for tr in rows[1:]:
        tbl.remove(tr)

    # Ensure two-column bilingual headers (no duplicate "Controlled products")
    products = ctx.products
    # rebuild header text on first row after body clear
    hdr = table.rows[0]
    if len(hdr.cells) < 2:
        # cannot split columns safely here — fall back to code-only
        if not products:
            new_tr = deepcopy(header_tr)
            tbl.append(new_tr)
            _set_cell_text(table.rows[1].cells[0], "—")
            return
        for p in products:
            new_tr = deepcopy(header_tr)
            tbl.append(new_tr)
        for i, p in enumerate(products):
            _set_cell_text(table.rows[i + 1].cells[0], str(p.product_code))
            _set_run_bold(table.rows[i + 1].cells[0], True)
        return

    _set_cell_text(hdr.cells[0], "ÜRÜN KODU\nPRODUCT CODE")
    _set_cell_text(hdr.cells[1], "ÜRÜN TANIMI\nPRODUCT DESCRIPTION")
    _set_run_bold(hdr.cells[0], True)
    _set_run_bold(hdr.cells[1], True)

    if not products:
        new_tr = deepcopy(header_tr)
        tbl.append(new_tr)
        _set_cell_text(table.rows[1].cells[0], "—")
        _set_cell_text(table.rows[1].cells[1], "—")
        _set_product_table_widths(table)
        return

    for _p in products:
        new_tr = deepcopy(header_tr)
        tbl.append(new_tr)
    for i, p in enumerate(products):
        row = table.rows[i + 1]
        name = (getattr(p, "product_name", "") or "").strip()
        _set_cell_text(row.cells[0], str(p.product_code))
        _set_run_bold(row.cells[0], True)
        _set_cell_text(row.cells[1], name if name else "—")
        _set_run_bold(row.cells[1], False)
        # enable wrap on description
        for para in row.cells[1].paragraphs:
            pf = para.paragraph_format
            try:
                pf.space_after = None
            except Exception:
                pass
    _set_product_table_widths(table)


def _insert_doc_product_scope_table(doc: Document, ctx: DocumentContext) -> None:
    """Add DoC Controlled Product Scope table (code | description) before signature block."""
    # Heading paragraph
    heading = doc.add_paragraph()
    run = heading.add_run(
        "Kontrollü ürün kapsamı / Controlled product scope"
    )
    run.bold = True
    run.font.name = "Tahoma"

    n = max(len(ctx.products), 1)
    table = doc.add_table(rows=1 + n, cols=2)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _set_cell_text(table.rows[0].cells[0], "ÜRÜN KODU\nPRODUCT CODE")
    _set_cell_text(table.rows[0].cells[1], "ÜRÜN TANIMI\nPRODUCT DESCRIPTION")
    _set_run_bold(table.rows[0].cells[0], True)
    _set_run_bold(table.rows[0].cells[1], True)
    if not ctx.products:
        _set_cell_text(table.rows[1].cells[0], "—")
        _set_cell_text(table.rows[1].cells[1], "—")
    else:
        for i, p in enumerate(ctx.products):
            _set_cell_text(table.rows[i + 1].cells[0], str(p.product_code))
            _set_run_bold(table.rows[i + 1].cells[0], True)
            _set_cell_text(
                table.rows[i + 1].cells[1],
                (getattr(p, "product_name", "") or "").strip() or "—",
            )
            _set_run_bold(table.rows[i + 1].cells[1], False)
    _set_product_table_widths(table)

    # Move heading + table before signature table if present
    sig_tbl = None
    for t in doc.tables:
        blob = " ".join(c.text for row in t.rows for c in row.cells)
        if "Numan Alver" in blob or ("İmza" in blob and "Signature" in blob):
            sig_tbl = t._tbl
            break
    if sig_tbl is not None:
        sig_tbl.addprevious(heading._p)
        sig_tbl.addprevious(table._tbl)


def ensure_doc_product_scope_table(doc: Document, ctx: DocumentContext) -> None:
    """Compatibility wrapper."""
    for table in doc.tables:
        if _is_product_table(table):
            _fill_product_table(table, ctx, for_technical_file=False)
            return
    _insert_doc_product_scope_table(doc, ctx)
