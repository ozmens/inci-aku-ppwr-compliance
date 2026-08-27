"""PPWR Supplier management — TDS / analysis document vault.

Stored under workspace/suppliers/ only. Never touches frozen deliveries.
"""

from __future__ import annotations

import json
import re
import shutil
import uuid
from datetime import datetime, timezone
from pathlib import Path

from fastapi import HTTPException, UploadFile

from web_mode import is_web_mode
from workspace_store import WORKSPACE, _assert_workspace, log_activity

SUPPLIERS = WORKSPACE / "suppliers"

DOC_TYPES = ("TDS", "ANALYSIS", "CERTIFICATE", "HEAVY_METALS", "SVHC", "PFAS", "OTHER")
ALLOWED_EXT = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt"}

MATERIAL_FAMILIES = ("", "PE", "PP", "PET", "PAPER", "CARDBOARD", "WOOD", "STEEL", "MIXED", "OTHER")
HM_STATUSES = ("unknown", "compliant", "non_compliant", "no_evidence")
SVHC_STATUSES = ("unknown", "none", "present", "no_declaration")
PFAS_STATUSES = ("unknown", "not_added", "present", "not_applicable")

PPWR_KEYWORDS = [
    ("recycled", "recycled_content"),
    ("geri dönüşüm", "recycled_content"),
    ("geri donusum", "recycled_content"),
    ("recycl", "recyclability"),
    ("geri dönüştürülebilir", "recyclability"),
    ("polymer", "polymer"),
    ("polimer", "polymer"),
    ("polyethylene", "pe"),
    ("polypropylen", "pp"),
    (" polyethylene", "pe"),
    ("hdpe", "pe"),
    ("ldpe", "pe"),
    ("pp ", "pp"),
    ("pet ", "pet"),
    ("cardboard", "paper"),
    ("karton", "paper"),
    ("paper", "paper"),
    ("kağıt", "paper"),
    ("kagit", "paper"),
    ("wood", "wood"),
    ("ahşap", "wood"),
    ("ahsap", "wood"),
    ("reach", "reach"),
    ("rohs", "rohs"),
    ("food contact", "food_contact"),
    ("gıda teması", "food_contact"),
    ("gida temasi", "food_contact"),
    ("heavy metal", "heavy_metals"),
    ("ağır metal", "heavy_metals"),
    ("agir metal", "heavy_metals"),
    ("pfas", "pfas"),
    ("svhc", "svhc"),
    ("reach", "reach"),
    ("substance", "substances"),
    ("madde", "substances"),
    ("density", "density"),
    ("yoğunluk", "density"),
    ("yogunluk", "density"),
    ("thickness", "thickness"),
    ("kalınlık", "kalinlik"),
    ("kalinlik", "kalinlik"),
]


def _clean_decl(raw: object, allowed: tuple[str, ...], extra_keys: tuple[str, ...] = ()) -> dict:
    src = raw if isinstance(raw, dict) else {}
    status = str(src.get("status") or "unknown").strip().lower()
    if status not in allowed:
        status = "unknown"
    out = {
        "status": status,
        "evidence_date": str(src.get("evidence_date") or "").strip()[:32],
        "evidence_doc_id": str(src.get("evidence_doc_id") or "").strip()[:80],
        "note": str(src.get("note") or "").strip()[:500],
    }
    for key in extra_keys:
        out[key] = str(src.get(key) or "").strip()[:200]
    return out


def normalize_link(link: dict) -> dict:
    """Fill PPWR declaration fields on a component–supplier link. Never writes TF/DoC."""
    row = dict(link or {})
    family = str(row.get("material_family") or "").strip().upper()
    row["material_family"] = family if family in MATERIAL_FAMILIES else ""
    row["heavy_metals"] = _clean_decl(row.get("heavy_metals"), HM_STATUSES)
    row["svhc"] = _clean_decl(
        row.get("svhc"),
        SVHC_STATUSES,
        extra_keys=("substance_name", "candidate_list_date"),
    )
    row["pfas"] = _clean_decl(row.get("pfas"), PFAS_STATUSES)
    pct = row.get("recycled_content_pct")
    try:
        row["recycled_content_pct"] = None if pct in (None, "") else float(pct)
    except (TypeError, ValueError):
        row["recycled_content_pct"] = None
    row["recyclability_note"] = str(row.get("recyclability_note") or "").strip()[:500]
    return row


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _ensure() -> None:
    SUPPLIERS.mkdir(parents=True, exist_ok=True)


def _slug(name: str) -> str:
    s = re.sub(r"[^\w\-]+", "_", (name or "").strip(), flags=re.UNICODE)
    s = re.sub(r"_+", "_", s).strip("_")
    return (s or "TEDARIKCI")[:48]


def _safe_id(raw: str) -> str:
    s = re.sub(r"[^\w\-]+", "_", (raw or "").strip())
    s = re.sub(r"_+", "_", s).strip("_")
    if not s:
        raise HTTPException(400, "Invalid supplier id")
    return s[:64]


def _dir(supplier_id: str) -> Path:
    return SUPPLIERS / _safe_id(supplier_id)


def _meta_path(supplier_id: str) -> Path:
    return _dir(supplier_id) / "SUPPLIER.json"


def _docs_dir(supplier_id: str) -> Path:
    return _dir(supplier_id) / "documents"


def _analyses_dir(supplier_id: str) -> Path:
    return _dir(supplier_id) / "analyses"


def _links_path(supplier_id: str) -> Path:
    return _dir(supplier_id) / "LINKS.json"


def _load(supplier_id: str) -> dict:
    path = _meta_path(supplier_id)
    if not path.exists():
        raise HTTPException(404, f"Supplier not found: {supplier_id}")
    return json.loads(path.read_text(encoding="utf-8"))


def _save_meta(data: dict) -> dict:
    sid = data["id"]
    root = _dir(sid)
    root.mkdir(parents=True, exist_ok=True)
    _docs_dir(sid).mkdir(parents=True, exist_ok=True)
    _analyses_dir(sid).mkdir(parents=True, exist_ok=True)
    path = _meta_path(sid)
    _assert_workspace(SUPPLIERS)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return data


def _doc_summary(supplier_id: str) -> dict:
    docs = list_documents(supplier_id).get("documents") or []
    types = {d["doc_type"] for d in docs}
    return {
        "doc_count": len(docs),
        "has_tds": "TDS" in types,
        "has_analysis": "ANALYSIS" in types,
        "has_certificate": "CERTIFICATE" in types,
        "readiness": (
            "ready"
            if "TDS" in types and "ANALYSIS" in types
            else "partial"
            if docs
            else "empty"
        ),
    }


def list_suppliers(q: str = "") -> dict:
    _ensure()
    needle = (q or "").strip().lower()
    items = []
    for folder in sorted(p for p in SUPPLIERS.iterdir() if p.is_dir()):
        meta = folder / "SUPPLIER.json"
        if not meta.exists():
            continue
        data = json.loads(meta.read_text(encoding="utf-8"))
        sid = data.get("id") or folder.name
        summary = _doc_summary(sid)
        links = list_links(sid).get("links") or []
        row = {
            "id": sid,
            "code": data.get("code") or sid,
            "name": data.get("name") or sid,
            "country": data.get("country") or "",
            "status": data.get("status") or "ACTIVE",
            "external_ref": data.get("external_ref") or "",
            "note": data.get("note") or "",
            "updated_at": data.get("updated_at"),
            "link_count": len(links),
            **summary,
        }
        if needle:
            blob = " ".join(
                [
                    row["code"],
                    row["name"],
                    row["country"],
                    row["external_ref"],
                    row["note"],
                ]
            ).lower()
            if needle not in blob:
                continue
        items.append(row)
    items.sort(key=lambda x: (x.get("name") or "").lower())
    return {"suppliers_root": str(SUPPLIERS), "count": len(items), "suppliers": items}


def get_supplier(supplier_id: str) -> dict:
    data = _load(supplier_id)
    docs = list_documents(supplier_id)
    links = list_links(supplier_id).get("links") or []
    return {
        **data,
        **_doc_summary(supplier_id),
        "link_count": len(links),
        "links": links,
        "documents": docs.get("documents") or [],
        "folder": str(_dir(supplier_id)),
    }


def save_supplier(
    *,
    name: str,
    code: str = "",
    country: str = "",
    status: str = "ACTIVE",
    external_ref: str = "",
    note: str = "",
    contact: str = "",
    materials: str = "",
    supplier_id: str | None = None,
) -> dict:
    _ensure()
    name = (name or "").strip()
    if len(name) < 2:
        raise HTTPException(400, "Supplier name required")
    code = (code or "").strip() or _slug(name)
    status = (status or "ACTIVE").strip().upper()
    if status not in ("ACTIVE", "INACTIVE", "PENDING"):
        status = "ACTIVE"

    sid = (supplier_id or "").strip() or f"{_slug(code)}_{uuid.uuid4().hex[:6]}"
    sid = _safe_id(sid)

    prev = {}
    if _meta_path(sid).exists():
        prev = _load(sid)

    data = {
        "id": sid,
        "code": code,
        "name": name,
        "country": (country or "").strip(),
        "status": status,
        "external_ref": (external_ref or "").strip(),
        "contact": (contact or "").strip(),
        "materials": (materials or "").strip(),
        "note": (note or "").strip(),
        "created_at": prev.get("created_at") or _now(),
        "updated_at": _now(),
    }
    _save_meta(data)
    log_activity("supplier_save", supplier_id=sid, name=name)
    return get_supplier(sid)


def delete_supplier(supplier_id: str) -> dict:
    root = _dir(supplier_id)
    if not root.exists():
        raise HTTPException(404, "Supplier not found")
    _assert_workspace(root)
    shutil.rmtree(root)
    log_activity("supplier_delete", supplier_id=supplier_id)
    return {"deleted": supplier_id}


def list_documents(supplier_id: str) -> dict:
    _load(supplier_id)
    docs_root = _docs_dir(supplier_id)
    docs_root.mkdir(parents=True, exist_ok=True)
    index_path = docs_root / "INDEX.json"
    if index_path.exists():
        index = json.loads(index_path.read_text(encoding="utf-8"))
    else:
        index = {"documents": []}
    # drop missing files
    kept = []
    for d in index.get("documents") or []:
        p = docs_root / d.get("stored_name", "")
        if p.exists():
            d["exists"] = True
            d["size"] = p.stat().st_size
            kept.append(d)
    index["documents"] = kept
    index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"supplier_id": supplier_id, "documents": kept}


def _write_index(supplier_id: str, documents: list[dict]) -> None:
    docs_root = _docs_dir(supplier_id)
    path = docs_root / "INDEX.json"
    _assert_workspace(docs_root)
    path.write_text(
        json.dumps({"documents": documents, "updated_at": _now()}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def upload_document(
    supplier_id: str,
    file: UploadFile,
    *,
    doc_type: str = "TDS",
    title: str = "",
    note: str = "",
    analyze: bool = True,
) -> dict:
    data = _load(supplier_id)
    doc_type = (doc_type or "TDS").strip().upper()
    if doc_type not in DOC_TYPES:
        raise HTTPException(400, f"doc_type must be one of {DOC_TYPES}")

    original = Path(file.filename or "upload.bin").name
    ext = Path(original).suffix.lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(400, f"Allowed extensions: {sorted(ALLOWED_EXT)}")

    docs_root = _docs_dir(supplier_id)
    docs_root.mkdir(parents=True, exist_ok=True)
    _assert_workspace(docs_root)

    doc_id = f"{doc_type}_{uuid.uuid4().hex[:10]}"
    stored = f"{doc_id}{ext}"
    dest = docs_root / stored
    with dest.open("wb") as out:
        shutil.copyfileobj(file.file, out)

    entry = {
        "id": doc_id,
        "doc_type": doc_type,
        "title": (title or "").strip() or Path(original).stem,
        "original_name": original,
        "stored_name": stored,
        "note": (note or "").strip(),
        "uploaded_at": _now(),
        "size": dest.stat().st_size,
        "exists": True,
        "analysis_id": None,
    }

    idx = list_documents(supplier_id)
    docs = idx.get("documents") or []
    docs.append(entry)
    _write_index(supplier_id, docs)

    analysis = None
    if analyze:
        analysis = analyze_document(supplier_id, doc_id)
        entry["analysis_id"] = analysis.get("id")
        # refresh index with analysis_id
        docs = [entry if d["id"] == doc_id else d for d in docs]
        _write_index(supplier_id, docs)

    data["updated_at"] = _now()
    _save_meta(data)
    log_activity(
        "supplier_upload",
        supplier_id=supplier_id,
        doc_id=doc_id,
        doc_type=doc_type,
        name=original,
    )
    return {"document": entry, "analysis": analysis, "supplier": get_supplier(supplier_id)}


def _extract_text(path: Path, limit: int = 40000) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")[:limit]
    if ext == ".docx":
        try:
            from docx import Document

            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)[:limit]
        except Exception as ex:  # noqa: BLE001
            return f"[docx extract error: {ex}]"
    if ext == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            parts = []
            for page in reader.pages[:30]:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
            return "\n".join(parts)[:limit]
        except Exception as ex:  # noqa: BLE001
            return f"[pdf extract error: {ex}]"
    if ext in {".xlsx", ".xls"}:
        try:
            from openpyxl import load_workbook

            wb = load_workbook(str(path), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets[:3]:
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i > 80:
                        break
                    cells = [str(c) for c in row if c is not None and str(c).strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)[:limit]
        except Exception as ex:  # noqa: BLE001
            return f"[xlsx extract error: {ex}]"
    return ""


def _scan_keywords(text: str) -> dict:
    low = (text or "").lower()
    hits: dict[str, int] = {}
    found_terms: list[str] = []
    for term, bucket in PPWR_KEYWORDS:
        if term in low:
            hits[bucket] = hits.get(bucket, 0) + 1
            if term not in found_terms:
                found_terms.append(term)
    return {"topics": hits, "matched_terms": found_terms[:40]}


def analyze_document(supplier_id: str, doc_id: str) -> dict:
    docs = list_documents(supplier_id).get("documents") or []
    entry = next((d for d in docs if d.get("id") == doc_id), None)
    if not entry:
        raise HTTPException(404, f"Document not found: {doc_id}")
    path = _docs_dir(supplier_id) / entry["stored_name"]
    if not path.exists():
        raise HTTPException(404, "Document file missing")

    text = _extract_text(path)
    scan = _scan_keywords(text)
    preview = re.sub(r"\s+", " ", text).strip()[:800]
    lang = "tr" if any(x in text.lower() for x in ("ürün", "ambalaj", "geri", "madde")) else "en"

    analysis = {
        "id": f"AN_{doc_id}",
        "supplier_id": supplier_id,
        "document_id": doc_id,
        "doc_type": entry.get("doc_type"),
        "title": entry.get("title"),
        "analyzed_at": _now(),
        "char_count": len(text),
        "language_guess": lang,
        "topics": scan["topics"],
        "matched_terms": scan["matched_terms"],
        "ppwr_signals": sorted(scan["topics"].keys()),
        "text_preview": preview,
        "extract_ok": bool(text) and not text.startswith("["),
        "note": "",
    }

    out_dir = _analyses_dir(supplier_id)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{analysis['id']}.json"
    _assert_workspace(out_dir)
    out_path.write_text(json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8")

    # link on document
    for d in docs:
        if d["id"] == doc_id:
            d["analysis_id"] = analysis["id"]
    _write_index(supplier_id, docs)
    log_activity("supplier_analyze", supplier_id=supplier_id, doc_id=doc_id)
    return analysis


def get_analysis(supplier_id: str, analysis_id: str) -> dict:
    path = _analyses_dir(supplier_id) / f"{analysis_id}.json"
    if not path.exists():
        raise HTTPException(404, "Analysis not found")
    return json.loads(path.read_text(encoding="utf-8"))


def open_document(supplier_id: str, doc_id: str) -> dict:
    docs = list_documents(supplier_id).get("documents") or []
    entry = next((d for d in docs if d.get("id") == doc_id), None)
    if not entry:
        raise HTTPException(404, "Document not found")
    path = resolve_document_path(supplier_id, doc_id)
    import os

    if is_web_mode() or not hasattr(os, "startfile"):
        return {
            "opened": str(path),
            "document": entry,
            "download": True,
            "download_url": (
                f"/api/suppliers/{supplier_id}/documents/{doc_id}/file"
            ),
        }
    os.startfile(str(path))
    return {"opened": str(path), "document": entry, "download": False}


def resolve_document_path(supplier_id: str, doc_id: str) -> Path:
    docs = list_documents(supplier_id).get("documents") or []
    entry = next((d for d in docs if d.get("id") == doc_id), None)
    if not entry:
        raise HTTPException(404, "Document not found")
    path = _docs_dir(supplier_id) / entry["stored_name"]
    if not path.exists():
        raise HTTPException(404, "File missing")
    return path


def open_supplier_folder(supplier_id: str | None = None) -> dict:
    _ensure()
    import os

    if is_web_mode() or not hasattr(os, "startfile"):
        raise HTTPException(400, "Klasör açma yalnızca masaüstü kurulumunda")
    target = _dir(supplier_id) if supplier_id else SUPPLIERS
    if supplier_id and not target.exists():
        raise HTTPException(404, "Supplier folder missing")
    target.mkdir(parents=True, exist_ok=True)
    os.startfile(str(target))
    return {"opened": str(target)}


def delete_document(supplier_id: str, doc_id: str) -> dict:
    docs = list_documents(supplier_id).get("documents") or []
    entry = next((d for d in docs if d.get("id") == doc_id), None)
    if not entry:
        raise HTTPException(404, "Document not found")
    docs_root = _docs_dir(supplier_id)
    path = docs_root / entry["stored_name"]
    _assert_workspace(docs_root)
    if path.exists():
        path.unlink()
    aid = entry.get("analysis_id")
    if aid:
        ap = _analyses_dir(supplier_id) / f"{aid}.json"
        if ap.exists():
            ap.unlink()
    kept = [d for d in docs if d["id"] != doc_id]
    _write_index(supplier_id, kept)
    log_activity("supplier_doc_delete", supplier_id=supplier_id, doc_id=doc_id)
    return {"deleted": doc_id}


def list_links(supplier_id: str) -> dict:
    _load(supplier_id)
    path = _links_path(supplier_id)
    if not path.exists():
        return {"supplier_id": supplier_id, "links": []}
    data = json.loads(path.read_text(encoding="utf-8"))
    links = [normalize_link(l) for l in (data.get("links") or [])]
    links.sort(key=lambda x: (x.get("component_code") or "").lower())
    return {"supplier_id": supplier_id, "links": links}


def _save_links(supplier_id: str, links: list[dict]) -> None:
    root = _dir(supplier_id)
    root.mkdir(parents=True, exist_ok=True)
    path = _links_path(supplier_id)
    _assert_workspace(root)
    path.write_text(
        json.dumps({"links": links, "updated_at": _now()}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def link_component(
    supplier_id: str,
    *,
    component_code: str,
    description: str = "",
    set_code: str = "",
    note: str = "",
    preferred: bool = False,
    scope: str = "starter",
) -> dict:
    """Link a packaging component (from master BOM) to this supplier."""
    from masters import search_components

    code = (component_code or "").strip()
    if not code:
        raise HTTPException(400, "component_code required")

    desc = (description or "").strip()
    set_code = (set_code or "").strip()
    if not desc:
        found = search_components(scope, code, limit=20).get("components") or []
        exact = next((c for c in found if c["component_code"].upper() == code.upper()), None)
        if exact:
            code = exact["component_code"]
            desc = exact.get("description") or ""
            if not set_code and exact.get("set_codes"):
                set_code = exact["set_codes"][0]

    links = list_links(supplier_id).get("links") or []
    existing = next((l for l in links if (l.get("component_code") or "").upper() == code.upper()), None)
    if existing:
        existing["description"] = desc or existing.get("description") or ""
        existing["set_code"] = set_code or existing.get("set_code") or ""
        existing["note"] = (note or "").strip() or existing.get("note") or ""
        existing["preferred"] = bool(preferred)
        existing["scope"] = scope
        existing["updated_at"] = _now()
        link = existing
    else:
        link = {
            "id": f"L_{uuid.uuid4().hex[:8]}",
            "component_code": code,
            "description": desc,
            "set_code": set_code,
            "note": (note or "").strip(),
            "preferred": bool(preferred),
            "scope": scope,
            "linked_at": _now(),
            "updated_at": _now(),
        }
        links.append(link)

    links = [normalize_link(l) for l in links]
    link = next(l for l in links if (l.get("component_code") or "").upper() == code.upper())
    _save_links(supplier_id, links)
    log_activity("supplier_link", supplier_id=supplier_id, component_code=code)
    return {"link": link, "supplier": get_supplier(supplier_id)}


def unlink_component(supplier_id: str, link_id: str) -> dict:
    links = list_links(supplier_id).get("links") or []
    kept = [l for l in links if l.get("id") != link_id]
    if len(kept) == len(links):
        # also allow unlink by component code
        kept = [l for l in links if (l.get("component_code") or "") != link_id]
    if len(kept) == len(links):
        raise HTTPException(404, "Link not found")
    _save_links(supplier_id, kept)
    log_activity("supplier_unlink", supplier_id=supplier_id, link_id=link_id)
    return {"deleted": link_id, "supplier": get_supplier(supplier_id)}


def update_link(supplier_id: str, link_id: str, patch: dict) -> dict:
    """Update PPWR declaration fields on a component link. Does not touch TF/DoC files."""
    links = list_links(supplier_id).get("links") or []
    found = next((l for l in links if l.get("id") == link_id), None)
    if not found:
        found = next((l for l in links if (l.get("component_code") or "") == link_id), None)
    if not found:
        raise HTTPException(404, "Link not found")
    src = patch or {}
    if "preferred" in src and src["preferred"] is not None:
        found["preferred"] = bool(src["preferred"])
    if "note" in src and src["note"] is not None:
        found["note"] = str(src["note"]).strip()[:500]
    if "material_family" in src and src["material_family"] is not None:
        found["material_family"] = src["material_family"]
    if "recyclability_note" in src and src["recyclability_note"] is not None:
        found["recyclability_note"] = src["recyclability_note"]
    if "recycled_content_pct" in src:
        found["recycled_content_pct"] = src["recycled_content_pct"]
    if isinstance(src.get("heavy_metals"), dict):
        found["heavy_metals"] = {**(found.get("heavy_metals") or {}), **src["heavy_metals"]}
    if isinstance(src.get("svhc"), dict):
        found["svhc"] = {**(found.get("svhc") or {}), **src["svhc"]}
    if isinstance(src.get("pfas"), dict):
        found["pfas"] = {**(found.get("pfas") or {}), **src["pfas"]}
    found["updated_at"] = _now()
    links = [normalize_link(l) if l.get("id") != found.get("id") else normalize_link(found) for l in links]
    saved = next(l for l in links if l.get("id") == found.get("id"))
    _save_links(supplier_id, links)
    log_activity(
        "supplier_link_update",
        supplier_id=supplier_id,
        component_code=saved.get("component_code"),
        link_id=saved.get("id"),
    )
    return {"link": saved, "supplier": get_supplier(supplier_id)}


def suppliers_for_component(component_code: str) -> dict:
    """Reverse lookup: which suppliers are linked to a component."""
    code = (component_code or "").strip().upper()
    if not code:
        raise HTTPException(400, "component_code required")
    _ensure()
    hits = []
    for folder in sorted(p for p in SUPPLIERS.iterdir() if p.is_dir()):
        meta = folder / "SUPPLIER.json"
        if not meta.exists():
            continue
        data = json.loads(meta.read_text(encoding="utf-8"))
        sid = data.get("id") or folder.name
        for link in list_links(sid).get("links") or []:
            if (link.get("component_code") or "").upper() == code:
                summary = _doc_summary(sid)
                hits.append(
                    {
                        "supplier_id": sid,
                        "supplier_code": data.get("code"),
                        "supplier_name": data.get("name"),
                        "preferred": bool(link.get("preferred")),
                        "link": link,
                        "has_tds": summary["has_tds"],
                        "has_analysis": summary["has_analysis"],
                        "readiness": summary["readiness"],
                        "material_family": link.get("material_family") or "",
                        "heavy_metals_status": (link.get("heavy_metals") or {}).get("status") or "unknown",
                        "svhc_status": (link.get("svhc") or {}).get("status") or "unknown",
                        "pfas_status": (link.get("pfas") or {}).get("status") or "unknown",
                    }
                )
    hits.sort(key=lambda x: (not x["preferred"], (x.get("supplier_name") or "").lower()))
    return {"component_code": component_code, "count": len(hits), "suppliers": hits}


def _all_link_index() -> dict[str, list[dict]]:
    """component_code(upper) -> supplier link summaries."""
    _ensure()
    index: dict[str, list[dict]] = {}
    for folder in sorted(p for p in SUPPLIERS.iterdir() if p.is_dir()):
        meta = folder / "SUPPLIER.json"
        if not meta.exists():
            continue
        data = json.loads(meta.read_text(encoding="utf-8"))
        sid = data.get("id") or folder.name
        summary = _doc_summary(sid)
        for link in list_links(sid).get("links") or []:
            code = (link.get("component_code") or "").strip().upper()
            if not code:
                continue
            index.setdefault(code, []).append(
                {
                    "supplier_id": sid,
                    "supplier_code": data.get("code"),
                    "supplier_name": data.get("name"),
                    "preferred": bool(link.get("preferred")),
                    "has_tds": summary["has_tds"],
                    "has_analysis": summary["has_analysis"],
                    "readiness": summary["readiness"],
                    "set_code": link.get("set_code") or "",
                    "material_family": link.get("material_family") or "",
                    "heavy_metals_status": (link.get("heavy_metals") or {}).get("status") or "unknown",
                    "svhc_status": (link.get("svhc") or {}).get("status") or "unknown",
                    "pfas_status": (link.get("pfas") or {}).get("status") or "unknown",
                }
            )
    for code in index:
        index[code].sort(key=lambda x: (not x["preferred"], (x.get("supplier_name") or "").lower()))
    return index


def component_matrix(*, q: str = "", kind: str = "starter", limit: int = 80, linked_only: bool = False) -> dict:
    """Packaging components with supplier coverage matrix."""
    from masters import search_components

    link_index = _all_link_index()
    if linked_only and not (q or "").strip():
        # show only components that already have supplier links
        rows = []
        for code, suppliers in sorted(link_index.items()):
            rows.append(
                {
                    "component_code": code,
                    "description": "",
                    "set_codes": [s["set_code"] for s in suppliers if s.get("set_code")][:8],
                    "set_count": 0,
                    "supplier_count": len(suppliers),
                    "has_tds_supplier": any(s.get("has_tds") for s in suppliers),
                    "preferred_supplier": next(
                        (s["supplier_name"] for s in suppliers if s.get("preferred")),
                        (suppliers[0]["supplier_name"] if suppliers else None),
                    ),
                    "suppliers": suppliers,
                    "coverage": (
                        "covered"
                        if any(s.get("has_tds") for s in suppliers)
                        else "linked"
                        if suppliers
                        else "gap"
                    ),
                }
            )
        rows = rows[: max(1, min(limit, 300))]
        return {
            "kind": kind,
            "q": q,
            "linked_only": True,
            "total": len(rows),
            "components": rows,
            "linked_components": len(link_index),
        }

    base = search_components(kind, q, limit=max(1, min(limit, 300)))
    out = []
    for c in base.get("components") or []:
        code = (c.get("component_code") or "").upper()
        suppliers = link_index.get(code) or []
        if linked_only and not suppliers:
            continue
        coverage = (
            "covered"
            if any(s.get("has_tds") for s in suppliers)
            else "linked"
            if suppliers
            else "gap"
        )
        out.append(
            {
                **c,
                "supplier_count": len(suppliers),
                "has_tds_supplier": any(s.get("has_tds") for s in suppliers),
                "preferred_supplier": next(
                    (s["supplier_name"] for s in suppliers if s.get("preferred")),
                    (suppliers[0]["supplier_name"] if suppliers else None),
                ),
                "suppliers": suppliers,
                "coverage": coverage,
            }
        )
    return {
        "kind": base.get("kind") or kind,
        "q": q,
        "linked_only": linked_only,
        "total": len(out),
        "returned": len(out),
        "components": out,
        "linked_components": len(link_index),
    }
